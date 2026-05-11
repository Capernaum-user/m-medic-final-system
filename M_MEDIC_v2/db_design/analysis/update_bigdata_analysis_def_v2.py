"""
빅데이터 분석 정의서 전면 수정 (의료사고 대응 및 상처 부위 식별 중심)
양식 유지 및 프로젝트 목적 정렬
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path("db_design/analysis/05_빅데이터분석정의서.docx")

# ─── 양식 설정 (상수) ────────────────────────────────────────────────────
CONTENT_W = 9639
COL1 = 1843
COL2 = CONTENT_W - COL1
TOTAL = COL1 + COL2
BG_SECTION = "D9D9D9"
BG_LABEL   = "F2F2F2"
BG_WHITE   = "FFFFFF"
BG_INNER_H = "E8E8E8"
KR_FONT = "맑은 고딕"
EN_FONT = "맑은 고딕"

def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")): tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_width(cell, dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")): tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def cell_pad(cell, top=80, bottom=80, left=113, right=113):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:cantSplit")): trPr.remove(old)
    cs = OxmlElement("w:cantSplit"); trPr.append(cs)

def row_header(row):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:tblHeader")): trPr.remove(old)
    th = OxmlElement("w:tblHeader"); trPr.append(th)

def para_line_bottom(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single"); btm.set(qn("w:sz"), "6"); btm.set(qn("w:space"), "1"); btm.set(qn("w:color"), "000000")
    pBdr.append(btm); pPr.append(pBdr)

def tbl_borders(table):
    tbl = table._tbl; tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}"); tag.set(qn("w:val"), "single"); tag.set(qn("w:sz"), "4"); tag.set(qn("w:space"), "0"); tag.set(qn("w:color"), "000000")
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def tbl_layout_fixed(table):
    tbl = table._tbl; tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblLayout")): tblPr.remove(old)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)

def tbl_width_dxa(table, dxa):
    tbl = table._tbl; tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")): tblPr.remove(old)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(dxa)); w.set(qn("w:type"), "dxa"); tblPr.append(w)

def tbl_grid(table, widths_dxa):
    tbl = table._tbl; old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None: tbl.remove(old_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None: tblPr.addnext(tblGrid)
    else: tbl.insert(0, tblGrid)

def set_run_font(run, size=10, bold=False):
    run.bold = bold; run.font.size = Pt(size); run.font.name = EN_FONT
    rPr = run._r.get_or_add_rPr(); rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None: rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT); rFonts.set(qn("w:hAnsi"), EN_FONT); rFonts.set(qn("w:eastAsia"), KR_FONT); rFonts.set(qn("w:cs"), EN_FONT)

def R(para, text, bold=False, size=10):
    run = para.add_run(text); set_run_font(run, size=size, bold=bold); return run

def row_funcname(table, name):
    row = table.add_row(); row_cant_split(row); row_header(row)
    lc, rc = row.cells[0], row.cells[1]; shade(lc, BG_SECTION); shade(rc, BG_WHITE); cell_width(lc, COL1);  cell_width(rc, COL2)
    cell_pad(lc, top=100, bottom=100, left=80, right=80); cell_pad(rc, top=90, bottom=90, left=130, right=113)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER; R(lp, "기능명", bold=True, size=10)
    rp = rc.paragraphs[0]; R(rp, name, bold=False, size=10)

def row_section(table, text):
    row = table.add_row(); row_cant_split(row); cell = row.cells[0].merge(row.cells[1]); shade(cell, BG_SECTION); cell_width(cell, TOTAL); cell_pad(cell, top=90, bottom=90); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; R(p, text, bold=True, size=10)

def row_data(table, label, split=False):
    row = table.add_row()
    if not split: row_cant_split(row)
    lc, rc = row.cells[0], row.cells[1]; shade(lc, BG_LABEL); shade(rc, BG_WHITE); cell_width(lc, COL1); cell_width(rc, COL2); cell_pad(lc, top=110, bottom=110, left=80, right=80); cell_pad(rc, top=90, bottom=90, left=130, right=113)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP; lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER; R(lp, label, bold=True, size=10); return lc, rc

def fill(rc, lines):
    first = True
    for item in lines:
        if isinstance(item, str): text, bold, ind = item, False, 0.0
        else: text = item[0]; bold = item[1] if len(item) > 1 else False; ind  = item[2] if len(item) > 2 else 0.0
        if first: p = rc.paragraphs[0]; first = False
        else: p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if ind: p.paragraph_format.left_indent = Cm(ind)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2); R(p, text, bold=bold, size=10)

def inner_table(rc, headers, rows_data, col_widths=None):
    inner_total = COL2 - 260
    if col_widths is None:
        each = inner_total // len(headers); col_widths = [each] * (len(headers) - 1) + [inner_total - each * (len(headers) - 1)]
    else:
        s = sum(col_widths); col_widths = [int(w * inner_total / s) for w in col_widths]; diff = inner_total - sum(col_widths); col_widths[-1] += diff
    tbl = rc.add_table(rows=1, cols=len(headers)); tbl_borders(tbl); tbl_layout_fixed(tbl); tbl_width_dxa(tbl, inner_total); tbl_grid(tbl, col_widths)
    hrow = tbl.rows[0]; row_cant_split(hrow)
    for cell, h, cw in zip(hrow.cells, headers, col_widths):
        shade(cell, BG_INNER_H); cell_width(cell, cw); cell_pad(cell, top=50, bottom=50, left=60, right=60); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; R(p, h, bold=True, size=9)
    for rd in rows_data:
        row = tbl.add_row(); row_cant_split(row)
        for cell, val, cw in zip(row.cells, rd, col_widths):
            cell_width(cell, cw); cell_pad(cell, top=40, bottom=40, left=60, right=60); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; R(p, val, size=9)
    return tbl

def setup_document(doc):
    section = doc.sections[0]; section.page_width = Cm(21.0); section.page_height = Cm(29.7); section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0); section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    style = doc.styles["Normal"]; style.font.name = EN_FONT; rPr = style.element.get_or_add_rPr(); rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None: rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT); rFonts.set(qn("w:hAnsi"), EN_FONT); rFonts.set(qn("w:eastAsia"), KR_FONT); rFonts.set(qn("w:cs"), EN_FONT)
    header = section.header; hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; R(hp, "MDTS 프로젝트  빅데이터 분석 정의서", size=9); para_line_bottom(hp)

def build_cover(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; R(p, "MDTS 프로젝트 성과발표회", size=9); para_line_bottom(p)
    for _ in range(6): doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1); tbl_borders(t); tbl_layout_fixed(t); tbl_width_dxa(t, CONTENT_W); tbl_grid(t, [CONTENT_W])
    tc = t.rows[0].cells[0]; row_cant_split(t.rows[0]); shade(tc, BG_WHITE); cell_pad(tc, top=400, bottom=400)
    p = tc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("빅데이터  분석  정의서"); set_run_font(r, size=30, bold=True)
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(2.5); r1 = p.add_run("과제명 :  "); set_run_font(r1, size=13, bold=True)
    r2 = p.add_run('MobileNetV3 기반 선박 외상 분류 시스템,  "MDTS"'); set_run_font(r2, size=13, bold=True)
    for _ in range(6): doc.add_paragraph()
    for text in ("2026.04.20.", "팀명:  MDTS (빅데이터분석팀)"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(text); set_run_font(r, size=12, bold=True)
    doc.add_page_break()

def build_overview(doc):
    p = doc.add_paragraph(); R(p, "Ⅰ.  개요", bold=True, size=13)
    def item(num, title, detail=None):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); R(p, f"{num}. {title}", size=11)
        if detail:
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(1.5); R(p2, f": {detail}", size=11)
    item("1", "아이디어 주제", "선박 내 외상 이미지 자동 분류 및 신체 부위별 응급처치 가이드 제공")
    item("2", "개발 목표", "의료 공백 해소를 위한 외상 및 부위 식별 AI 모델 기반 스마트 의료 지원")
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); R(p, "3. 개발 내용", size=11)
    for detail in (": MobileNetV3-Small을 활용한 외상 유형 및 상처 면적 분석", ": 신체 부위별 응급처치 도구 매칭 및 가이드 제공", ": 의료 사고 대응을 위한 중증도 자동 판정 시스템"):
        p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(1.5); R(p2, detail, size=11)
    doc.add_page_break()

def build_func1(doc):
    p = doc.add_paragraph(); R(p, "Ⅱ.  기능별 빅데이터 분석 명세서", bold=True, size=13); doc.add_paragraph()
    table = create_main_table(doc)
    row_funcname(table, "외상 유형 및 신체 부위 식별 (의료 대응)")
    row_section(table, "1. 데이터 준비")
    _, rc = row_data(table, "데이터 정의", split=True)
    fill(rc, ["선박 내 의료 사고 대응용 외상 이미지 데이터", "- 외상 유형: 찰과상, 타박상, 화상, 절상, 열상 등", "- 신체 부위: 상처가 발생한 위치 정보 포함", "- 치료 도구: First Aid Kit 관련 이미지"])
    _, rc = row_data(table, "데이터\n획득 방법", split=True)
    fill(rc, ["1. Kaggle - Wound Classification 및 Segmentation 데이터셋", "2. Kaggle - Skin Burn (화상 심도 분석용 특화 데이터)", "3. 의료용 해부학 이미지 (신체 부위 학습용 대체 데이터)"])
    _, rc = row_data(table, "수집 데이터", split=True)
    inner_table(rc, ["데이터셋", "규모", "핵심 용도"], [("Wound Classification", "2,000장+", "상처 종류 식별"), ("Wound Segmentation", "2,760세트", "상처 면적 측정"), ("Burn Extra", "600장+", "화상 중증도 판정")], col_widths=[4, 3, 3])
    row_section(table, "2. 전처리")
    _, rc = row_data(table, "전처리 과정"); fill(rc, ["1) 리사이징 (224x224)", "2) 상처 영역 마스킹 (Segmentation)", "3) 신체 부위별 라벨링 정규화"])
    row_section(table, "3. 데이터 분석")
    _, rc = row_data(table, "데이터\n분석 목표"); fill(rc, ["상처 유형과 부위를 동시에 분석하여 적절한 응급처치 도구(Kit) 권고"])
    _, rc = row_data(table, "데이터\n분석\n알고리즘"); fill(rc, ["Multi-Task Learning 기반 MobileNetV3 (Type + Location)"])
    row_section(table, "4. 모델 생성 및 학습")
    _, rc = row_data(table, "학습 모델"); fill(rc, ["MobileNetV3-Small (경량화 추론 최적화)"])
    row_section(table, "5. 검증")
    _, rc = row_data(table, "모델링\n검증 방안"); fill(rc, ["의료사고 대응 적절성 평가 (처치 도구 매칭 정확도)"])

def create_main_table(doc):
    table = doc.add_table(rows=0, cols=2); tbl_borders(table); tbl_layout_fixed(table); tbl_width_dxa(table, TOTAL); tbl_grid(table, [COL1, COL2]); return table

def main():
    doc = Document(); setup_document(doc); build_cover(doc); build_overview(doc); build_func1(doc); doc.save(str(OUT)); print("수정 완료")

if __name__ == "__main__": main()
