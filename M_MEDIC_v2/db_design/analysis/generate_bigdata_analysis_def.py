"""
빅데이터 분석 정의서 생성기
서식 기준: 3-3 빅데이터분석정의서_샘플1_ver2.pdf
프로젝트: M-MEDIC v2 (선박 외상 분류 및 해양사고 위험도 예측)
출력: 05_빅데이터분석정의서.docx

수정:
- 표 너비를 본문 너비(9639 twips ≈ 17cm)에 맞춤 (기존 11622 → 9639)
- 한글 폰트(맑은 고딕) eastAsia 설정으로 한글 깨짐 방지
- tblLayout fixed + cantSplit 적용 (페이지 분할 시 셀 잘림 방지)
- 표 행 헤더 자동 반복 (다중 페이지 표)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "05_빅데이터분석정의서.docx"

# ─── 페이지 및 컬럼 너비 ────────────────────────────────────────────────────
# A4 21 × 29.7 cm, 좌우 여백 2 cm씩 → 본문 폭 17 cm ≈ 9639 twips
CONTENT_W = 9639
COL1 = 1843         # 라벨 열 ≈ 3.25 cm
COL2 = CONTENT_W - COL1   # 7796 (약 13.75 cm)
TOTAL = COL1 + COL2       # = 9639

# ─── 배경색 ──────────────────────────────────────────────────────────────────
BG_SECTION = "D9D9D9"   # 섹션 헤더 (진회색)
BG_LABEL   = "F2F2F2"   # 라벨 셀 (연회색)
BG_WHITE   = "FFFFFF"
BG_INNER_H = "E8E8E8"   # 내부 테이블 헤더

# ─── 폰트 ─────────────────────────────────────────────────────────────────────
KR_FONT = "맑은 고딕"
EN_FONT = "맑은 고딕"


# ═══════════════════════════════ 유틸 함수 ═══════════════════════════════════

def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_width(cell, dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def cell_pad(cell, top=80, bottom=80, left=113, right=113):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def row_cant_split(row):
    """행이 페이지 분할 시 쪼개지지 않도록"""
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:cantSplit")):
        trPr.remove(old)
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)


def row_header(row):
    """행을 반복 헤더 행으로 지정"""
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:tblHeader")):
        trPr.remove(old)
    th = OxmlElement("w:tblHeader")
    trPr.append(th)


def para_line_bottom(para):
    """단락 아래 구분선"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single")
    btm.set(qn("w:sz"), "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "000000")
    pBdr.append(btm)
    pPr.append(pBdr)


def tbl_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def tbl_layout_fixed(table):
    """셀 너비 고정 (내용이 많아도 컬럼 안 밀림)"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def tbl_width_dxa(table, dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(dxa))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)


def tbl_grid(table, widths_dxa):
    tbl = table._tbl
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def set_run_font(run, size=10, bold=False):
    """폰트 / 한글 폰트 / 크기 / 굵기 설정"""
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = EN_FONT
    # eastAsia(한글) 폰트 명시
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    rFonts.set(qn("w:eastAsia"), KR_FONT)
    rFonts.set(qn("w:cs"), EN_FONT)


def R(para, text, bold=False, size=10):
    """단락에 Run 추가"""
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


# ═══════════════════════════════ 테이블 행 빌더 ═══════════════════════════════

def row_funcname(table, name):
    row = table.add_row()
    row_cant_split(row)
    row_header(row)
    lc, rc = row.cells[0], row.cells[1]
    shade(lc, BG_SECTION); shade(rc, BG_WHITE)
    cell_width(lc, COL1);  cell_width(rc, COL2)
    cell_pad(lc, top=100, bottom=100, left=80, right=80)
    cell_pad(rc, top=90, bottom=90, left=130, right=113)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(lp, "기능명", bold=True, size=10)
    rp = rc.paragraphs[0]
    R(rp, name, bold=False, size=10)


def row_section(table, text):
    row = table.add_row()
    row_cant_split(row)
    cell = row.cells[0].merge(row.cells[1])
    shade(cell, BG_SECTION)
    cell_width(cell, TOTAL)
    cell_pad(cell, top=90, bottom=90)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, text, bold=True, size=10)


def row_data(table, label, split=False):
    """
    split=True 이면 행 내용이 많아도 페이지 분할 허용 (cantSplit 미적용).
    기본적으로 대부분의 행은 합리적인 크기라서 cantSplit 적용.
    """
    row = table.add_row()
    if not split:
        row_cant_split(row)
    lc, rc = row.cells[0], row.cells[1]
    shade(lc, BG_LABEL); shade(rc, BG_WHITE)
    cell_width(lc, COL1); cell_width(rc, COL2)
    cell_pad(lc, top=110, bottom=110, left=80, right=80)
    cell_pad(rc, top=90, bottom=90, left=130, right=113)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(lp, label, bold=True, size=10)
    return lc, rc


def fill(rc, lines):
    first = True
    for item in lines:
        if isinstance(item, str):
            text, bold, ind = item, False, 0.0
        else:
            text = item[0]
            bold = item[1] if len(item) > 1 else False
            ind  = item[2] if len(item) > 2 else 0.0
        if first:
            p = rc.paragraphs[0]; first = False
        else:
            p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if ind:
            p.paragraph_format.left_indent = Cm(ind)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        R(p, text, bold=bold, size=10)


def append_to(rc, lines):
    """inner_table 등 이후 rc에 단락을 추가 (항상 add_paragraph)"""
    for item in lines:
        if isinstance(item, str):
            text, bold, ind = item, False, 0.0
        else:
            text = item[0]
            bold = item[1] if len(item) > 1 else False
            ind  = item[2] if len(item) > 2 else 0.0
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if ind:
            p.paragraph_format.left_indent = Cm(ind)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        R(p, text, bold=bold, size=10)


def inner_table(rc, headers, rows_data, col_widths=None):
    """셀 내부 표"""
    # 부모 컨텐트 너비 내로 맞춤 (COL2 - 좌우 padding)
    inner_total = COL2 - 260   # 좌우 padding 고려
    if col_widths is None:
        each = inner_total // len(headers)
        col_widths = [each] * (len(headers) - 1) + [inner_total - each * (len(headers) - 1)]
    else:
        # 주어진 비율로 스케일
        s = sum(col_widths)
        col_widths = [int(w * inner_total / s) for w in col_widths]
        # 합계 보정
        diff = inner_total - sum(col_widths)
        col_widths[-1] += diff

    tbl = rc.add_table(rows=1, cols=len(headers))
    tbl_borders(tbl)
    tbl_layout_fixed(tbl)
    tbl_width_dxa(tbl, inner_total)
    tbl_grid(tbl, col_widths)

    hrow = tbl.rows[0]
    row_cant_split(hrow)
    for cell, h, cw in zip(hrow.cells, headers, col_widths):
        shade(cell, BG_INNER_H)
        cell_width(cell, cw)
        cell_pad(cell, top=50, bottom=50, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9)
    for rd in rows_data:
        row = tbl.add_row()
        row_cant_split(row)
        for cell, val, cw in zip(row.cells, rd, col_widths):
            cell_width(cell, cw)
            cell_pad(cell, top=40, bottom=40, left=60, right=60)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            R(p, val, size=9)
    return tbl


# ═══════════════════════════════ 문서 설정 ═══════════════════════════════════

def setup_document(doc):
    section = doc.sections[0]
    section.page_width        = Cm(21.0)
    section.page_height       = Cm(29.7)
    section.top_margin        = Cm(2.0)
    section.bottom_margin     = Cm(2.0)
    section.left_margin       = Cm(2.0)
    section.right_margin      = Cm(2.0)
    section.different_first_page_header_footer = True

    # Normal 스타일에 맑은 고딕 기본 설정
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    rFonts.set(qn("w:eastAsia"), KR_FONT)
    rFonts.set(qn("w:cs"), EN_FONT)

    # 2페이지 이후 헤더
    header = section.header
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    R(hp, "MDTS 프로젝트  빅데이터 분석 정의서", size=9)
    para_line_bottom(hp)

    fh = section.first_page_header
    for p in fh.paragraphs:
        p.clear()


# ═══════════════════════════════ 표지 ═══════════════════════════════════════

def build_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    R(p, "MDTS 프로젝트 성과발표회", size=9)
    para_line_bottom(p)

    for _ in range(6):
        doc.add_paragraph()

    # 제목 박스 (표 1 × 1, 본문 폭에 맞춤)
    t = doc.add_table(rows=1, cols=1)
    tbl_borders(t)
    tbl_layout_fixed(t)
    tbl_width_dxa(t, CONTENT_W)
    tbl_grid(t, [CONTENT_W])
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc = t.rows[0].cells[0]
    row_cant_split(t.rows[0])
    cell_width(tc, CONTENT_W)
    shade(tc, BG_WHITE)
    cell_pad(tc, top=400, bottom=400, left=200, right=200)
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("빅데이터  분석  정의서")
    set_run_font(r, size=30, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    # 과제명
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2.5)
    r1 = p.add_run("과제명 :  ")
    set_run_font(r1, size=13, bold=True)
    r2 = p.add_run('MobileNetV3 기반 선박 외상 분류 시스템,  "MDTS"')
    set_run_font(r2, size=13, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    # 날짜 / 팀명
    for text in ("2026.04.17.", "팀명:  MDTS"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=12, bold=True)

    doc.add_page_break()


# ═══════════════════════════════ I. 개요 ════════════════════════════════════

def build_overview(doc):
    p = doc.add_paragraph()
    R(p, "Ⅰ.  개요", bold=True, size=13)

    def item(num, title, detail=None):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        R(p, f"{num}. {title}", size=11)
        if detail:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.5)
            p2.paragraph_format.space_after = Pt(2)
            R(p2, f": {detail}", size=11)

    item("1", "아이디어 주제",
         "MobileNetV3을 활용한 선박 외상 이미지 자동 분류 및 응급처치 가이드 제공 시스템")
    item("2", "개발 목표",
         "선박 응급처치 지원을 위한 외상 분류 AI 모델 개발 및 처치 가이드 서비스 구현")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    R(p, "3. 개발 내용", size=11)

    for detail in (
        ": MobileNetV3-Small을 활용한 6종 외상 이미지 분류 모델 개발",
        ": 외상 분류 결과 기반 선박 내 응급처치 가이드 자동 제공",
    ):
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.5)
        p2.paragraph_format.space_after = Pt(2)
        R(p2, detail, size=11)

    doc.add_paragraph()


# ═══════════════════════════════ II. 제목 ════════════════════════════════════

def build_section2_heading(doc):
    p = doc.add_paragraph()
    R(p, "Ⅱ.  기능별 빅데이터 분석 명세서", bold=True, size=13)
    doc.add_paragraph()


# ═══════════════════════════════ 공통 표 생성 ═══════════════════════════════

def create_main_table(doc):
    table = doc.add_table(rows=0, cols=2)
    tbl_borders(table)
    tbl_layout_fixed(table)
    tbl_width_dxa(table, TOTAL)
    tbl_grid(table, [COL1, COL2])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    return table


# ═══════════════════════════════ 기능 1: 외상 분류 ═══════════════════════════

def build_func1(doc):
    table = create_main_table(doc)

    row_funcname(table, "외상 이미지 분류  (CNN)")

    row_section(table, "1. 데이터 준비")

    _, rc = row_data(table, "데이터 정의", split=True)
    fill(rc, [
        "선박 사고 시 발생하는 외상 이미지  (6종 클래스)",
        ("- Abrasions   (찰과상) : 갑판·구조물에 쓸린 상처  ·  249장", False, 0.3),
        ("- Bruises     (타박상) : 충돌·낙하 시 타박상      ·  364장", False, 0.3),
        ("- Burns       (화상)   : 엔진실 화재·증기 화상    ·  193장", False, 0.3),
        ("- Cut         (절창)   : 날카로운 도구에 의한 절상·  150장", False, 0.3),
        ("- Laceration  (열창)   : 거친 표면에 의한 찢김    ·  183장", False, 0.3),
        ("- Stab_wound  (자창)   : 날카로운 물체 관통       ·   23장", False, 0.3),
        "합계 : 1,162장",
    ])

    _, rc = row_data(table, "데이터\n획득 방법", split=True)
    fill(rc, [
        "1.  Kaggle  —  wound-classification-system",
        ("URL : https://www.kaggle.com/datasets/wound-classification-system", False, 0.4),
        ("파일 prefix : wcs_*  (6개 클래스 전체 포함)", False, 0.4),
        "",
        "2.  Kaggle  —  wound_dataset_yasin",
        ("URL : https://www.kaggle.com/datasets/yasin/wound-dataset", False, 0.4),
        ("파일 prefix : yasin_*  (7개 클래스 중 선박 외상 관련 6개 선별,  원본 431장)", False, 0.4),
        ("※ Ingrown_nails (내성발톱)은 선박 외상 무관 → deprecated/ 제외", False, 0.4),
        ("※ yasin 431장은 wound_classification 1,162장에 병합 포함됨", False, 0.4),
    ])

    _, rc = row_data(table, "수집 데이터", split=True)
    fill(rc, ["[CNN 분류 학습용]  통합 경로 :  01_data/raw/wound_classification/", ""])
    inner_table(rc,
        ["클래스 (한글명)", "장수", "위험도"],
        [
            ("Abrasions  (찰과상)", "249", "LOW"),
            ("Bruises    (타박상)", "364", "LOW"),
            ("Burns      (화상)",   "193", "HIGH"),
            ("Cut        (절창)",   "150", "MEDIUM"),
            ("Laceration (열창)",   "183", "HIGH"),
            ("Stab_wound (자창)",   " 23", "HIGH"),
            ("합  계",              "1,162", "—"),
        ],
        col_widths=[5, 2, 2]
    )
    append_to(rc, [
        "",
        ("[세그멘테이션 별도 모델용]  01_data/raw/wound_segmentation/", False, 0),
        ("- train_images 2,208장 + test_images 552장 = 이미지 2,760장  (마스크 동수 포함)", False, 0.3),
        ("- CNN 분류기와 독립적인 별도 파이프라인 (상처 영역 검출용)", False, 0.3),
    ])

    row_section(table, "2. 전처리")

    _, rc = row_data(table, "전처리 과정", split=True)
    fill(rc, [
        "1)  이미지 리사이즈 : 224 × 224 픽셀 통일",
        "2)  데이터 증강  (훈련셋 한정)",
        ("- RandomHorizontalFlip,  RandomVerticalFlip", False, 0.4),
        ("- RandomRotation (±25°),  ColorJitter (밝기·대비·채도·색조)", False, 0.4),
        ("- RandomAffine (translate ±5%)", False, 0.4),
        "3)  정규화 : ImageNet 기준",
        ("mean = [0.485, 0.456, 0.406]   /   std = [0.229, 0.224, 0.225]", False, 0.4),
        "4)  WeightedRandomSampler : 소수 클래스 (Stab_wound 23장) 오버샘플링",
        "5)  Train / Val 분할 : 8 : 2  (Stratified Split,  random_state=42)",
        ("→ 훈련 929장  /  검증 233장", False, 0.4),
    ])

    row_section(table, "3. 데이터 분석")

    _, rc = row_data(table, "데이터\n분석 목표")
    fill(rc, ["선박 사고 외상 6종 자동 분류 및 위험도  (HIGH / MEDIUM / LOW)  판정"])

    _, rc = row_data(table, "데이터\n분석\n알고리즘", split=True)
    fill(rc, [
        "MobileNetV3-Small  (Transfer Learning,  ImageNet 사전학습 가중치 활용)",
        ("- 파라미터 수  :  약 2.5 M", False, 0.4),
        ("- 모델 크기    :  약 10 MB  →  선박 엣지 디바이스 배포 최적화", False, 0.4),
        ("- 추론 속도    :  < 200 ms  (CPU 기준)", False, 0.4),
        ("- 분류기 교체  :  Classifier 마지막 레이어 → Linear (in_features → 6)", False, 0.4),
    ])

    row_section(table, "4. 모델 생성 및 학습")

    _, rc = row_data(table, "모델링 목표")
    fill(rc, [
        "MobileNetV3-Small 을 이용한 선박 사고 외상 자동 분류 모델 생성",
        ("- 찰과상, 타박상, 화상, 절창, 열창, 자창  6종 분류", False, 0.4),
    ])

    _, rc = row_data(table, "학습 모델")
    fill(rc, ["MobileNetV3-Small  (torchvision.models.mobilenet_v3_small)"])

    _, rc = row_data(table, "학습 방법", split=True)
    fill(rc, [
        "Batch-size :  32",
        "Epochs     :  20",
        "Optimizer  :  AdamW  (lr = 1e-3,  weight_decay = 1e-4)",
        "Scheduler  :  CosineAnnealingLR  (T_max = 20,  eta_min = 1e-6)",
        "Loss       :  CrossEntropyLoss",
        "장치       :  CUDA (GPU 우선)  /  CPU 폴백",
    ])

    row_section(table, "5. 검증")

    _, rc = row_data(table, "모델링\n검증 방안", split=True)
    fill(rc, [
        "- Accuracy        :  전체 외상 분류 정확도",
        "- Precision (정밀도) :  해당 클래스로 분류한 것 중 실제 해당 클래스인 비율",
        "- Recall  (재현율)   :  실제 해당 클래스 중 모델이 올바르게 예측한 비율",
        "- F1 Score        :  Precision 과 Recall 의 조화평균",
        "- Confusion Matrix :  클래스별 분류 오류 패턴 시각화",
    ])

    _, rc = row_data(table, "모델\n평가 결과", split=True)
    fill(rc, [
        "(학습 실행 후 업데이트 예정)",
        "- Accuracy          :  —",
        "- Precision (정밀도) :  —",
        "- Recall  (재현율)   :  —",
        "- F1 Score          :  —",
    ])


# ═══════════════════════════════ 기능 2: 해양사고 위험도 예측 ════════════════

def build_func2(doc):
    doc.add_paragraph()   # 기능 간 여백

    table = create_main_table(doc)

    row_funcname(table, "해양사고 위험도 예측  (ML)")

    row_section(table, "1. 데이터 준비")

    _, rc = row_data(table, "데이터 정의", split=True)
    fill(rc, [
        "해양사고 통계 데이터  (6개 컬럼)",
        ("- 연도, 사고유형, 선박종류, 사고원인, 발생해역, 사상자수", False, 0.4),
        "사고유형 : 충돌, 침몰, 화재, 전복, 좌초, 침수, 기관고장, 안전사고",
        "선박종류 : 어선, 화물선, 유조선, 상선, 낚시어선, 레저기구, 레저보트",
        "사고원인 : 운항과실, 기상악화, 기계결함, 정비불량, 전기단락, 적재불량 등",
        "발생해역 : 남해, 서해, 동해, 제주",
    ])

    _, rc = row_data(table, "데이터\n획득 방법", split=True)
    fill(rc, [
        "1.  GICOMS 해양사고 정보시스템 재구성 데이터",
        ("출처 : 중앙해양안전심판원 통계 기반 재구성  (원본 데이터 아님)", False, 0.4),
        ("파일 : raw/gicoms_marine_accidents_2014_2024.csv", False, 0.4),
        ("규모 : 48건 × 6컬럼  (2014 ~ 2024년)", False, 0.4),
        "",
        "2.  합성 해양사고 데이터",
        ("출처 : 한국 해양안전심판원 통계 패턴 기반 합성 생성", False, 0.4),
        ("파일 : processed/marine_accidents_augmented.csv", False, 0.4),
        ("규모 : 1,500건 × 6컬럼  (01_generate_accident_data.py 로 생성 완료)", False, 0.4),
        ("저장 위치 : processed/marine_accidents_augmented.csv", False, 0.4),
    ])

    _, rc = row_data(table, "수집 데이터", split=True)
    fill(rc, ["통합 활용 데이터 :  GICOMS 재구성 (48건) + 합성 생성 (1,500건) = 1,548건", ""])
    inner_table(rc,
        ["데이터 종류", "건수", "기간 / 비고"],
        [
            ("GICOMS 재구성", "48건",    "2014 ~ 2024"),
            ("합성 데이터",   "1,500건", "통계 패턴 기반"),
            ("합  계",        "1,548건", "—"),
        ],
        col_widths=[3, 2, 4]
    )

    row_section(table, "2. 전처리")

    _, rc = row_data(table, "전처리 과정", split=True)
    fill(rc, [
        "1)  범주형 변수 인코딩  (Label Encoding)",
        ("- 사고유형, 선박종류, 사고원인, 발생해역  →  정수 코드 변환", False, 0.4),
        "2)  위험도 라벨 생성  (사상자수 기준 3단계)",
        ("- HIGH   : 사상자 10명 이상", False, 0.4),
        ("- MEDIUM : 사상자 1 ~ 9명", False, 0.4),
        ("- LOW    : 사상자 없음  (0명)", False, 0.4),
        "3)  결측치 확인 및 처리",
        "4)  Train / Test 분할 : 8 : 2  (random_state = 42)",
    ])

    row_section(table, "3. 데이터 분석")

    _, rc = row_data(table, "데이터\n분석 목표")
    fill(rc, ["사고 유형·선박 종류·원인·발생 해역 기반 해양사고 위험도  (HIGH / MEDIUM / LOW)  자동 예측"])

    _, rc = row_data(table, "데이터\n분석\n알고리즘", split=True)
    fill(rc, [
        "1)  Random Forest Classifier  (주 분류 모델)",
        ("- 앙상블 기법으로 과적합 방지 및 안정적인 성능 확보", False, 0.4),
        ("- Feature Importance 분석을 통한 사고 위험 요인 식별", False, 0.4),
        "",
        "2)  특성 중요도  (Feature Importance)  분석",
        ("- 위험도 예측에 영향이 큰 특성 순위 도출", False, 0.4),
        ("  ex)  사고유형 > 발생해역 > 선박종류 > 사고원인  (예시)", False, 0.4),
    ])

    row_section(table, "4. 모델 생성 및 학습")

    _, rc = row_data(table, "모델링 목표")
    fill(rc, [
        "해양사고 발생 시 위험도  (HIGH / MEDIUM / LOW)  자동 예측 모델 생성",
        ("- 입력 : 사고유형, 선박종류, 사고원인, 발생해역, 연도", False, 0.4),
        ("- 출력 : 위험도 등급 3단계", False, 0.4),
    ])

    _, rc = row_data(table, "학습 모델")
    fill(rc, ["Random Forest Classifier  (scikit-learn)"])

    _, rc = row_data(table, "학습 방법", split=True)
    fill(rc, [
        "n_estimators      :  100  (결정 트리 수)",
        "max_depth         :  None  (자동 결정)",
        "min_samples_split :  2",
        "Cross-validation  :  5-fold CV  (과적합 방지)",
    ])

    row_section(table, "5. 검증")

    _, rc = row_data(table, "모델링\n검증 방안", split=True)
    fill(rc, [
        "- Accuracy           :  전체 위험도 예측 정확도",
        "- Precision / Recall / F1 Score  (클래스별 HIGH·MEDIUM·LOW)",
        "- Feature Importance :  위험도 예측에 영향을 미치는 특성 순위",
        "- Confusion Matrix   :  위험도 등급별 오분류 패턴 분석",
    ])

    _, rc = row_data(table, "모델\n평가 결과", split=True)
    fill(rc, [
        "(학습 실행 후 업데이트 예정)",
        "- Accuracy          :  —",
        "- Precision (정밀도) :  —",
        "- Recall  (재현율)   :  —",
        "- F1 Score          :  —",
    ])


# ═══════════════════════════════ 메인 ════════════════════════════════════════

def main():
    doc = Document()
    setup_document(doc)

    build_cover(doc)
    build_overview(doc)
    build_section2_heading(doc)
    build_func1(doc)

    doc.save(str(OUT))
    print(f"저장 완료: {OUT}")
    print(f"경로: {OUT.resolve()}")


if __name__ == "__main__":
    main()
