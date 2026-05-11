"""
분석목표서 + 분석용 컬럼 사전을 Word 문서(.docx)로 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── 공통 스타일 헬퍼 ─────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 강제 지정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        set_font(run, size=13, bold=True, color=(68, 114, 196))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        set_font(run, size=11, bold=True, color=(0, 0, 0))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    run = p.add_run(text)
    set_font(run, size=10)
    p.paragraph_format.space_after = Pt(1)

def shade_cell(cell, hex_color="D9E1F2"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, bg=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        shade_cell(cell, bg)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 70)
    set_font(run, size=8, color=(180, 180, 180))

def set_page_margins(doc):
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

def add_meta_table(doc, rows):
    """문서 헤더 정보 2컬럼 표"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    col_widths = [Cm(3.5), Cm(13)]
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, bold=True, size=9, bg="BDD7EE")
        set_cell_text(table.cell(i, 1), v, size=9)
    for row in table.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = w
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 문서 1: 분석 목표서
# ══════════════════════════════════════════════════════════════════════════════

def build_doc1():
    doc = Document()
    set_page_margins(doc)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("분  석  목  표  서")
    set_font(run, size=22, bold=True, color=(31, 73, 125))
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("MDTS — 선박 탑재형 엣지 AI 의료 지원 시스템")
    set_font(run2, size=12, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(10)

    add_divider(doc)

    # 메타 정보
    add_meta_table(doc, [
        ("프로젝트명", "MDTS (선박 탑재형 엣지 AI 의료 지원 시스템)"),
        ("담당 파트",  "빅데이터 분석"),
        ("작성일",     "2026-04-16"),
        ("작성자",     "권태향"),
        ("문서 버전",  "v1.0"),
    ])

    # 1. 분석 배경 및 목적
    add_heading(doc, "1. 분석 배경 및 목적", level=2)
    for line in [
        "선박은 사고 발생 시 육상 의료기관에서 수 시간~수일 이상 격리된 환경에 놓인다.",
        "승조원의 생체 신호 및 외상 이미지를 기반으로 AI가 위험 등급을 자동 판단하고,",
        "응급처치 가이드를 즉시 제공하는 것이 MDTS의 핵심 기능이다.",
        "",
        "본 분석의 목적은 DB에 축적되는 운용 데이터에서 의미 있는 패턴을 도출하여:",
    ]:
        add_body(doc, line)

    for line in [
        "시스템의 AI 판단 정확도와 실제 처치 간의 일치도를 검증한다.",
        "위험 등급별 생체 신호 기준 범위를 데이터 기반으로 수립한다.",
        "승조원 특성(직책, 혈액형, 연령대)과 부상/위험도 간 상관관계를 파악한다.",
        "육상 서버 동기화 패턴을 분석하여 통신 장애 리스크를 식별한다.",
    ]:
        add_bullet(doc, line)

    # 2. 분석 대상 테이블
    add_heading(doc, "2. 분석 대상 테이블", level=2)
    headers = ["테이블명", "한글명", "분석 역할"]
    rows = [
        ("tb_crew",     "승조원 정보",       "인적 특성 기준 변수 (독립변수 후보)"),
        ("tb_vital",    "생체 신호 데이터",   "핵심 수치형 측정값 (독립변수 주력)"),
        ("tb_analysis", "AI 분석 결과",      "예측 결과 및 위험 등급 (종속변수 주력)"),
        ("tb_firstaid", "응급 처치 기록",     "실제 처치 내용 (결과 검증 기준)"),
        ("tb_logs",     "동기화 로그",        "시스템 운용 패턴 분석"),
    ]
    t = doc.add_table(rows=1+len(rows), cols=3)
    t.style = "Table Grid"
    widths = [Cm(4), Cm(4.5), Cm(8)]
    for j, h in enumerate(headers):
        set_cell_text(t.cell(0, j), h, bold=True, size=9, bg="4472C4", center=True)
        t.cell(0, j)._tc.get_or_add_tcPr()
        # 헤더 흰 글씨
        run = t.cell(0, j).paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
    for i, (nm, kr, role) in enumerate(rows):
        bg = "D9E1F2" if i % 2 == 0 else "EBF0FA"
        set_cell_text(t.cell(i+1, 0), nm,   size=9, bg=bg)
        set_cell_text(t.cell(i+1, 1), kr,   size=9, bg=bg)
        set_cell_text(t.cell(i+1, 2), role, size=9, bg=bg)
    for row in t.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w
    doc.add_paragraph()

    # 3. 분석 목표 항목
    add_heading(doc, "3. 분석 목표 항목", level=2)

    goals = [
        {
            "title": "목표 A — 위험 등급 예측 모델 검증",
            "items": [
                ("분석 내용", "AI가 판단한 risk_level(1~4)이 실제 처치 내용(action_taken)과 일치하는지 확인"),
                ("사용 컬럼", "tb_analysis.risk_level, tb_firstaid.action_taken, tb_firstaid.guide_text"),
                ("분석 방법", "혼동 행렬(Confusion Matrix), 정밀도/재현율 계산"),
                ("기대 결과", "위험 등급별 처치 일치율 도출 → 모델 신뢰도 지표 확보"),
            ],
        },
        {
            "title": "목표 B — 생체 신호와 위험 등급 간 임계값 분석",
            "items": [
                ("분석 내용", "심박수(heart_rate), 산소포화도(spo2), 체온(temperature)의 수치 범위와 위험 등급의 관계"),
                ("사용 컬럼", "tb_vital.heart_rate, tb_vital.spo2, tb_vital.temperature, tb_analysis.risk_level"),
                ("분석 방법", "등급별 기술통계(평균/표준편차/IQR), 박스플롯, 상관분석"),
                ("기대 결과", "위험 등급별 생체 신호 임계값 범위표 → 현장 판단 기준선으로 활용"),
            ],
        },
        {
            "title": "목표 C — 승조원 특성과 위험도 연관성 분석",
            "items": [
                ("분석 내용", "직책, 연령대, 혈액형별 위험 등급 분포 차이 검토"),
                ("사용 컬럼", "tb_crew.position, tb_crew.birthdate, tb_crew.gender, tb_analysis.risk_level"),
                ("분석 방법", "그룹별 빈도 분석, 카이제곱 검정, 연령 파생 변수(나이) 계산"),
                ("기대 결과", "고위험군 직책 및 연령대 식별 → 사전 예방 모니터링 대상 선정"),
            ],
        },
        {
            "title": "목표 D — 처치 소요 시간 분석",
            "items": [
                ("분석 내용", "AI 분석 완료 시점(analyzed_at)부터 응급처치 완료(created_at)까지 소요 시간"),
                ("사용 컬럼", "tb_analysis.analyzed_at, tb_firstaid.created_at, tb_analysis.risk_level"),
                ("분석 방법", "소요 시간 파생 변수 계산, 위험 등급별 평균 비교"),
                ("기대 결과", "위험 등급별 처치 반응 시간 분포 → 시스템 응답 속도 SLA 기준 수립"),
            ],
        },
        {
            "title": "목표 E — 육상 동기화 장애 패턴 분석",
            "items": [
                ("분석 내용", "동기화 실패(sync_status = FALSE) 건의 시간대 및 발생 빈도 패턴"),
                ("사용 컬럼", "tb_logs.sync_status, tb_logs.synced_at, tb_logs.table_name"),
                ("분석 방법", "시계열 빈도 분석, 실패율 계산, 테이블별 오류 집계"),
                ("기대 결과", "동기화 취약 시간대 및 테이블 식별 → 인프라 개선 우선순위 도출"),
            ],
        },
    ]

    goal_colors = ["FFF2CC", "E2EFDA", "FCE4D6", "EBF0FA", "F4E6F1"]
    header_colors = ["FFD966", "A9D18E", "F4B183", "9DC3E6", "D5A6CC"]

    for idx, g in enumerate(goals):
        add_heading(doc, g["title"], level=3)
        tbl = doc.add_table(rows=len(g["items"]), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(g["items"]):
            set_cell_text(tbl.cell(i, 0), k, bold=True, size=9, bg=header_colors[idx], center=True)
            set_cell_text(tbl.cell(i, 1), v, size=9, bg=goal_colors[idx])
        for row in tbl.rows:
            row.cells[0].width = Cm(3.5)
            row.cells[1].width = Cm(13)
        doc.add_paragraph()

    # 4. 분석 제외 항목
    add_heading(doc, "4. 분석 제외 항목 및 사유", level=2)
    excl_headers = ["컬럼", "제외 사유"]
    excl_rows = [
        ("tb_analysis.file_name / file_size / file_ext",
         "이미지 파일 메타데이터로, 분석보다 스토리지 관리 목적 — 분석 단계에서는 보조 참조만 활용"),
        ("tb_crew.name",
         "개인 식별 정보(PII) — 분석 시 익명화 처리 후 crew_id로만 참조"),
        ("tb_analysis.analysis_result (전문)",
         "비정형 TEXT — 키워드 추출이 필요한 경우 별도 NLP 분석 단계에서 처리"),
    ]
    et = doc.add_table(rows=1+len(excl_rows), cols=2)
    et.style = "Table Grid"
    for j, h in enumerate(excl_headers):
        set_cell_text(et.cell(0, j), h, bold=True, size=9, bg="C00000", center=True)
        et.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, (col, reason) in enumerate(excl_rows):
        bg = "FCE4D6" if i % 2 == 0 else "FBE5D6"
        set_cell_text(et.cell(i+1, 0), col,    size=9, bg=bg)
        set_cell_text(et.cell(i+1, 1), reason, size=9, bg=bg)
    for row in et.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(11)
    doc.add_paragraph()

    # 5. 분석 결과물 계획
    add_heading(doc, "5. 분석 결과물 계획", level=2)
    res_headers = ["산출물", "형식", "활용처"]
    res_rows = [
        ("위험 등급별 생체 신호 임계값 범위표", "표/차트",         "발표 자료, 모델 튜닝 기준"),
        ("AI 판단 vs 실제 처치 일치율",         "혼동 행렬 이미지", "발표 자료, 시스템 신뢰도 근거"),
        ("고위험군 직책/연령대 분포",            "막대 그래프",      "발표 자료"),
        ("처치 반응 시간 분포",                  "박스플롯",         "발표 자료"),
        ("동기화 실패 패턴",                     "시계열 차트",      "발표 자료"),
    ]
    rt = doc.add_table(rows=1+len(res_rows), cols=3)
    rt.style = "Table Grid"
    for j, h in enumerate(res_headers):
        set_cell_text(rt.cell(0, j), h, bold=True, size=9, bg="404040", center=True)
        rt.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, row_data in enumerate(res_rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            set_cell_text(rt.cell(i+1, j), val, size=9, bg=bg)
    for row in rt.rows:
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(5.5)
    doc.add_paragraph()

    # 6. 전제 조건 및 제약사항
    add_heading(doc, "6. 전제 조건 및 제약사항", level=2)
    for line in [
        "현재 실제 운용 데이터 없음 → 분석 목표서/컬럼 사전은 실제 DB 운용 시작 후 데이터정의서로 확정",
        "생체 신호 정상 범위 기준은 한국 임상 표준 기준 적용 (심박수 60~100bpm, SpO2 ≥ 95%, 체온 36.1~37.2°C)",
        "tb_crew.birthdate에서 나이를 파생 변수로 계산하여 사용 (분석 시점 기준 만 나이)",
        "tb_logs.admin_id는 tb_crew.crew_id를 참조하므로 관리자 역할을 가진 승조원으로 한정",
    ]:
        add_bullet(doc, line)
    doc.add_paragraph()

    # 7. 다음 단계
    add_heading(doc, "7. 다음 단계", level=2)
    p = doc.add_paragraph()
    run = p.add_run("분석목표서 (완료)  →  분석용 컬럼 사전  →  데이터정의서  →  실제 분석 실행")
    set_font(run, size=11, bold=True, color=(31, 73, 125))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)

    path = os.path.join(OUT_DIR, "01_분석목표서.docx")
    doc.save(path)
    print(f"  저장: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 문서 2: 분석용 컬럼 사전
# ══════════════════════════════════════════════════════════════════════════════

def build_doc2():
    doc = Document()
    set_page_margins(doc)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("분석용  컬럼  사전")
    set_font(run, size=22, bold=True, color=(31, 73, 125))
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("MDTS — 선박 탑재형 엣지 AI 의료 지원 시스템")
    set_font(run2, size=12, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(10)

    add_divider(doc)

    # 메타 정보
    add_meta_table(doc, [
        ("프로젝트명", "MDTS (선박 탑재형 엣지 AI 의료 지원 시스템)"),
        ("담당 파트",  "빅데이터 분석"),
        ("작성일",     "2026-04-16"),
        ("작성자",     "권태향"),
        ("문서 버전",  "v1.0"),
        ("참조 문서",  "01_분석목표서.docx, DB멘토링(실전).SQL"),
    ])

    # 범례
    add_heading(doc, "범례", level=3)
    add_bullet(doc, "분석 역할: 독립변수(X) / 종속변수(Y) / 식별자 / 파생변수 / 참조전용")
    add_bullet(doc, "분석 우선순위: 상 / 중 / 하")
    doc.add_paragraph()

    # ── 테이블별 컬럼 사전 ────────────────────────────────────────────────────

    COL_HEADERS = ["#", "컬럼명", "한글명", "데이터 타입", "허용값/범위", "분석 역할", "우선순위", "분석 활용 방법", "비고"]
    COL_WIDTHS  = [Cm(0.8), Cm(2.8), Cm(2.2), Cm(3.0), Cm(3.5), Cm(2.5), Cm(1.5), Cm(5.5), Cm(4.0)]

    tables_data = [
        {
            "title": "1. tb_crew — 승조원 정보",
            "header_color": "1F497D",
            "row_colors": ["D9E1F2", "EBF0FA"],
            "cols": [
                ("1","crew_id",   "승조원 고유번호","INT (AUTO_INCREMENT)","1 이상 정수",         "식별자",        "중","tb_vital, tb_analysis, tb_firstaid, tb_logs 조인 키","PK, 직접 분석 대상 아님"),
                ("2","name",      "성명",           "VARCHAR(50)",         "한글/영문 문자열",    "참조전용",      "하","분석 시 익명화 — crew_id로만 참조",                 "PII, 결과 출력 시 마스킹"),
                ("3","birthdate", "생년월일",        "DATE",                "YYYY-MM-DD",          "파생변수",      "상","TIMESTAMPDIFF(YEAR, birthdate, NOW())로 만 나이 계산 → 연령대 그룹핑","나이 직접 저장 안 함"),
                ("4","gender",    "성별",            "CHAR(1)",             "'M' / 'F'",           "독립변수(X)",   "중","위험 등급 분포의 성별 차이 카이제곱 검정",            "값 표준화 필요"),
                ("5","bloodtype", "혈액형",          "VARCHAR(4)",          "'A'/'B'/'O'/'AB'",    "독립변수(X)",   "하","혈액형별 위험 등급 빈도 분석",                       "수혈 매칭 시나리오 활용 가능"),
                ("6","position",  "직책",            "VARCHAR(50)",         "선장/항해사/기관사 등","독립변수(X)",   "상","직책별 위험 등급 분포 — 고위험 직책 식별",           "입력값 정규화 필요"),
                ("7","joined_at", "가입 일시",        "DATETIME",            "CURRENT_TIMESTAMP",   "참조전용",      "하","승선 이력 기간 계산 시 참조",                        "직접 분석 변수 아님"),
            ],
        },
        {
            "title": "2. tb_vital — 생체 신호 데이터",
            "header_color": "375623",
            "row_colors": ["E2EFDA", "F0F7EC"],
            "cols": [
                ("1","vital_id",    "바이탈 고유번호","INT (AUTO_INCREMENT)","1 이상 정수",             "식별자",       "중","tb_analysis 조인 키",                              "PK"),
                ("2","crew_id",     "승조원 고유번호","INT",                 "tb_crew.crew_id 참조",    "식별자",       "중","승조원별 생체 신호 이력 집계 시 그룹 기준",          "FK"),
                ("3","heart_rate",  "심박수",         "INT",                 "30 ~ 250 (bpm)",          "독립변수(X)",  "상","위험 등급별 심박수 분포(박스플롯), 임계값 분석",     "정상: 60~100 / 이상: <40 or >150"),
                ("4","spo2",        "산소포화도",      "DECIMAL(5,2)",        "50.00 ~ 100.00 (%)",      "독립변수(X)",  "상","위험 등급과의 상관분석, 임계값 분석",               "정상: ≥95% / 위험: <90%"),
                ("5","temperature", "체온",            "DECIMAL(5,2)",        "34.00 ~ 42.00 (°C)",      "독립변수(X)",  "상","고열/저체온 기준 위험 등급 연관성 분석",            "정상: 36.1~37.2 / 고열: ≥38.5 / 저체온: ≤35.0"),
                ("6","measured_at", "측정 일시",        "DATETIME",            "CURRENT_TIMESTAMP",       "파생변수",     "중","시간대별 측정 빈도, analyzed_at과의 차이(반응 시간) 계산","목표 D 분석에 사용"),
            ],
        },
        {
            "title": "3. tb_analysis — AI 분석 결과",
            "header_color": "843C0C",
            "row_colors": ["FCE4D6", "FBF0EA"],
            "cols": [
                ("1","analysis_id",    "분석 고유번호","INT (AUTO_INCREMENT)","1 이상 정수",             "식별자",          "중","tb_firstaid 조인 키",                              "PK"),
                ("2","vital_id",       "바이탈 고유번호","INT",               "tb_vital.vital_id 참조",  "식별자",          "중","tb_vital 생체 신호와 분석 결과 연결",               "FK"),
                ("3","crew_id",        "승조원 고유번호","INT",               "tb_crew.crew_id 참조",    "식별자",          "중","승조원별 분석 이력 집계",                           "FK"),
                ("4","analysis_result","분석 내용",     "TEXT",               "AI 생성 자유 텍스트",     "참조전용",        "하","NLP 필요 시 별도 처리",                            "비정형, 직접 집계 불가"),
                ("5","diagnosis",      "진단 결과",      "VARCHAR(255)",       "AI 진단명 문자열",        "독립변수(X)",     "상","진단명 빈도 분석, 위험 등급과의 교차 분석",         "표준 코드화 여부 확인 필요"),
                ("6","file_name",      "파일 이름",       "VARCHAR(255)",       "UUID 포함 파일명",        "참조전용",        "하","이미지 파일 존재 여부 확인 시만 참조",              "분석 변수 아님"),
                ("7","file_size",      "파일 사이즈",     "INT",                "0 이상 정수 (bytes)",     "참조전용",        "하","스토리지 관리 목적 — 분석 제외",                   "분석 변수 아님"),
                ("8","file_ext",       "파일 확장자",     "VARCHAR(10)",        "'.jpg' / '.png' 등",      "참조전용",        "하","이미지 포맷 분포 확인 시만 참조",                  "분석 변수 아님"),
                ("9","risk_level",     "위험 등급",       "ENUM('1','2','3','4')","'1'(최저) ~ '4'(최고)", "종속변수(Y) ★",  "상","목표 A~C 전체의 핵심 종속변수",                    "1=GREEN/2=YELLOW/3=ORANGE/4=RED"),
                ("10","analyzed_at",   "분석 일시",        "DATETIME",           "CURRENT_TIMESTAMP",       "파생변수",        "중","tb_firstaid.created_at과의 차이로 처치 반응 시간 계산","목표 D 핵심 컬럼"),
            ],
        },
        {
            "title": "4. tb_firstaid — 응급 처치 기록",
            "header_color": "3F3F76",
            "row_colors": ["EBF0FA", "F5F8FD"],
            "cols": [
                ("1","firstaid_id",  "응급처치 고유번호","INT (AUTO_INCREMENT)","1 이상 정수",                "식별자",        "하","PK","직접 분석 변수 아님"),
                ("2","analysis_id",  "분석 고유번호",   "INT",                 "tb_analysis.analysis_id 참조","식별자",        "중","risk_level, analyzed_at 연결","FK"),
                ("3","crew_id",      "승조원 고유번호", "INT",                 "tb_crew.crew_id 참조",        "식별자",        "중","승조원별 처치 이력 집계","FK"),
                ("4","guide_text",   "가이드 내용",      "TEXT",                "AI 제공 처치 가이드",         "참조전용",      "하","AI 가이드 내용 확인 시 참조","비정형 데이터"),
                ("5","action_taken", "실제 조치 사항",   "TEXT",                "승조원이 수행한 처치 내용",   "독립변수(X)",   "상","guide_text와 의미 유사도 비교 → AI 가이드 준수율 측정 (목표 A)","비정형 — 키워드 매칭/NLP 필요"),
                ("6","created_at",   "처치 완료 일시",   "DATETIME",            "CURRENT_TIMESTAMP",           "파생변수",      "상","analyzed_at 기준 처치 반응 시간 계산 (목표 D)","목표 D 핵심 컬럼"),
            ],
        },
        {
            "title": "5. tb_logs — 육상 서버 동기화 로그",
            "header_color": "7030A0",
            "row_colors": ["F4E6F1", "FAF2FC"],
            "cols": [
                ("1","log_id",      "로그 고유번호",   "INT (AUTO_INCREMENT)","1 이상 정수",                     "식별자",          "하","PK","직접 분석 변수 아님"),
                ("2","admin_id",    "관리자 고유번호", "INT",                 "tb_crew.crew_id 참조",             "식별자",          "중","관리자별 동기화 처리 건수 집계","FK → tb_crew.crew_id"),
                ("3","table_name",  "대상 테이블명",   "VARCHAR(64)",         "tb_vital/tb_analysis/tb_firstaid 등","독립변수(X)",   "중","테이블별 동기화 실패 빈도 집계 (목표 E)","입력값 정규화 필요"),
                ("4","record_id",   "대상 레코드",     "VARCHAR(64)",         "각 테이블 PK 값 (문자열)",         "참조전용",        "하","특정 레코드 동기화 재시도 추적 시 참조","분석보다 운영 목적"),
                ("5","sync_status", "전송 상태",        "BOOLEAN",             "TRUE(성공) / FALSE(실패)",         "종속변수(Y) ★",  "상","동기화 실패율 계산, 시간대별/테이블별 실패 분포 (목표 E)","목표 E 핵심 종속변수"),
                ("6","synced_at",   "전송 일시",        "DATETIME",            "CURRENT_TIMESTAMP",               "파생변수",        "상","시간대별 실패율 시계열 분석 (목표 E)","목표 E 핵심 컬럼"),
            ],
        },
    ]

    for tbl_info in tables_data:
        add_heading(doc, tbl_info["title"], level=2)
        n_cols = len(COL_HEADERS)
        n_rows = len(tbl_info["cols"])
        t = doc.add_table(rows=1+n_rows, cols=n_cols)
        t.style = "Table Grid"

        # 헤더 행
        hc = tbl_info["header_color"]
        for j, h in enumerate(COL_HEADERS):
            set_cell_text(t.cell(0, j), h, bold=True, size=8, bg=hc, center=True)
            t.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        # 데이터 행
        rc0, rc1 = tbl_info["row_colors"]
        for i, row_vals in enumerate(tbl_info["cols"]):
            bg = rc0 if i % 2 == 0 else rc1
            for j, val in enumerate(row_vals):
                center = j in (0, 6)  # # 와 우선순위만 가운데
                set_cell_text(t.cell(i+1, j), val, size=8, bg=bg, center=center)
                # 종속변수 강조
                if j == 5 and "종속변수" in val:
                    t.cell(i+1, j).paragraphs[0].runs[0].font.bold = True
                    t.cell(i+1, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 0, 0)
                # 우선순위 상 강조
                if j == 6 and val == "상":
                    t.cell(i+1, j).paragraphs[0].runs[0].font.bold = True
                    t.cell(i+1, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(192, 0, 0)

        for row in t.rows:
            for j, w in enumerate(COL_WIDTHS):
                row.cells[j].width = w
        doc.add_paragraph()

    # ── 파생 변수 정의 ────────────────────────────────────────────────────────
    add_heading(doc, "파생 변수 정의 요약", level=2)
    dv_headers = ["파생 변수명", "계산 방법", "원본 컬럼", "활용 분석 목표"]
    dv_rows = [
        ("age",           "TIMESTAMPDIFF(YEAR, birthdate, NOW())",                   "tb_crew.birthdate",                               "목표 C"),
        ("age_group",     "age를 구간화 (20대 / 30대 / 40대+)",                      "age (파생)",                                      "목표 C"),
        ("response_time_min","TIMESTAMPDIFF(MINUTE, analyzed_at, created_at)",       "tb_analysis.analyzed_at, tb_firstaid.created_at", "목표 D"),
        ("risk_label",    "ENUM 숫자 → 문자 (1→GREEN / 2→YELLOW / 3→ORANGE / 4→RED)","tb_analysis.risk_level",                         "목표 A, B, C"),
        ("sync_hour",     "HOUR(synced_at)",                                         "tb_logs.synced_at",                               "목표 E"),
        ("sync_dayofweek","DAYOFWEEK(synced_at)",                                    "tb_logs.synced_at",                               "목표 E"),
    ]
    dt = doc.add_table(rows=1+len(dv_rows), cols=4)
    dt.style = "Table Grid"
    dv_widths = [Cm(3.5), Cm(6.5), Cm(5.5), Cm(2.5)]
    for j, h in enumerate(dv_headers):
        set_cell_text(dt.cell(0, j), h, bold=True, size=9, bg="404040", center=True)
        dt.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, row_vals in enumerate(dv_rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_vals):
            set_cell_text(dt.cell(i+1, j), val, size=9, bg=bg)
    for row in dt.rows:
        for j, w in enumerate(dv_widths):
            row.cells[j].width = w
    doc.add_paragraph()

    # ── 조인 관계 요약 ────────────────────────────────────────────────────────
    add_heading(doc, "테이블 간 조인 관계 요약", level=2)

    join_lines = [
        "tb_crew (crew_id)",
        "  ├── tb_vital      (crew_id FK)  ─→  tb_analysis (vital_id FK)  ─→  tb_firstaid (analysis_id FK)",
        "  ├── tb_firstaid   (crew_id FK)",
        "  └── tb_logs       (admin_id FK)",
    ]
    p_join = doc.add_paragraph()
    p_join.paragraph_format.left_indent = Cm(1)
    for line in join_lines:
        run = p_join.add_run(line + "\n")
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph()
    add_heading(doc, "핵심 분석 조인 SQL", level=3)

    sql1 = (
        "-- 목표 A, B, C, D 통합 분석용 기본 조인\n"
        "SELECT\n"
        "    c.crew_id,\n"
        "    TIMESTAMPDIFF(YEAR, c.birthdate, NOW())          AS age,\n"
        "    c.gender,\n"
        "    c.position,\n"
        "    v.heart_rate,\n"
        "    v.spo2,\n"
        "    v.temperature,\n"
        "    a.diagnosis,\n"
        "    a.risk_level,\n"
        "    a.analyzed_at,\n"
        "    f.action_taken,\n"
        "    f.created_at                                     AS treated_at,\n"
        "    TIMESTAMPDIFF(MINUTE, a.analyzed_at, f.created_at) AS response_time_min\n"
        "FROM tb_crew      c\n"
        "JOIN tb_vital     v ON c.crew_id     = v.crew_id\n"
        "JOIN tb_analysis  a ON v.vital_id    = a.vital_id\n"
        "JOIN tb_firstaid  f ON a.analysis_id = f.analysis_id;"
    )
    sql2 = (
        "-- 목표 E 동기화 분석용\n"
        "SELECT\n"
        "    l.sync_status,\n"
        "    l.table_name,\n"
        "    HOUR(l.synced_at)       AS sync_hour,\n"
        "    DAYOFWEEK(l.synced_at)  AS sync_dow,\n"
        "    COUNT(*)                AS cnt\n"
        "FROM tb_logs l\n"
        "GROUP BY l.sync_status, l.table_name, sync_hour, sync_dow;"
    )

    for sql in [sql1, sql2]:
        p_sql = doc.add_paragraph()
        p_sql.paragraph_format.left_indent = Cm(0.8)
        p_sql.paragraph_format.space_before = Pt(4)
        p_sql.paragraph_format.space_after  = Pt(8)
        run = p_sql.add_run(sql)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)

    # 다음 단계
    add_divider(doc)
    p = doc.add_paragraph()
    run = p.add_run("다음 단계:  분석목표서 (완료)  →  분석용 컬럼 사전 (완료)  →  데이터정의서  →  실제 분석 실행")
    set_font(run, size=11, bold=True, color=(31, 73, 125))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)

    path = os.path.join(OUT_DIR, "02_분석용_컬럼_사전.docx")
    doc.save(path)
    print(f"  저장: {path}")


if __name__ == "__main__":
    print("[MDTS] Word 문서 생성 중...")
    build_doc1()
    build_doc2()
    print("[완료]")
