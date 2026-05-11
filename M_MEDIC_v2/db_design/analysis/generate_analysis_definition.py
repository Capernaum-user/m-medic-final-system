"""
MDTS 데이터분석정의서 Word 문서 생성 (v2 - 정식 보고서 서식)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "04_데이터분석정의서.docx")

# ── 색상 팔레트 (전문 보고서 기준) ───────────────────────────────────────────
C = {
    "title_blue"   : (31,  73, 125),   # 제목 네이비
    "h2_blue"      : (23,  54,  93),   # 소제목 다크블루
    "h3_gray"      : (50,  50,  50),   # h3 다크그레이
    "body"         : (30,  30,  30),   # 본문
    "white"        : (255,255,255),
    "hdr_navy"     : "1F3864",         # 표 헤더 네이비
    "hdr_dark"     : "2E4057",         # 표 헤더 다크슬레이트
    "hdr_teal"     : "17496B",         # 표 헤더 틸
    "row_a"        : "EDF2F7",         # 짝수 행
    "row_b"        : "FFFFFF",         # 홀수 행
    "good"         : "E8F5E9",         # 양호 (연초록)
    "caution"      : "FFF9C4",         # 주의 (연노랑)
    "warn"         : "FFF3E0",         # 경고 (연주황)
    "danger"       : "FFEBEE",         # 위험 (연빨강)
    "good_txt"     : (27, 94, 32),
    "caution_txt"  : (130, 100, 0),
    "warn_txt"     : (183, 90, 0),
    "danger_txt"   : (183, 28, 28),
    "url_blue"     : (13,  71, 161),   # URL 파란색
    "note_bg"      : "FFF8E1",         # 노트 박스 배경
    "alert_bg"     : "FFEBEE",         # 경고 박스 배경
    "ok_green"     : (27,  94,  32),
    "ng_red"       : (183, 28,  28),
    "na_gray"      : (97,  97,  97),
    "accent"       : "E3F2FD",         # 강조 연파랑
}

# ── 공통 헬퍼 ────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, color=None, name="맑은 고딕", italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rf  = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), name)
    rPr.insert(0, rf)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    # 기존 shd 제거 후 재설정
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cw(cell, text, bold=False, size=9, bg=None, center=False,
        color=None, italic=False, wrap=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        shade_cell(cell, bg)

def set_col_widths(table, widths):
    for row in table.rows:
        for j, w in enumerate(widths):
            if j < len(row.cells):
                row.cells[j].width = w

def hdr_row(t, headers, bg=C["hdr_navy"], sizes=None):
    for j, h in enumerate(headers):
        sz = sizes[j] if sizes else 9
        cw(t.cell(0, j), h, bold=True, size=sz, bg=bg, center=True)
        t.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(*C["white"])

# ── 페이지 설정 ──────────────────────────────────────────────────────────────

def setup_page(doc):
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)
    _add_header_footer(doc)

def _add_header_footer(doc):
    from docx.oxml import OxmlElement
    sec = doc.sections[0]

    # ── 헤더 ──
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.clear()
    run = hp.add_run("MDTS  데이터분석정의서  |  v1.0  |  2026-04-16")
    set_font(run, size=8, color=C["na_gray"])
    # 헤더 아래 선
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── 푸터 (페이지 번호) ──
    ftr  = sec.footer
    ftr.is_linked_to_previous = False
    fp   = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run_l = fp.add_run("빅데이터분석  |  MDTS 프로젝트          ")
    set_font(run_l, size=8, color=C["na_gray"])
    # 페이지 번호 필드
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = ' PAGE '
    fldChar2  = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    r.append(fldChar1); r.append(instrText); r.append(fldChar2)
    run_pg = fp.add_run(" / ")
    set_font(run_pg, size=8, color=C["na_gray"])
    fp._p.append(r)
    # 전체 페이지
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText'); instrText2.text = ' NUMPAGES '
    fldChar4  = OxmlElement('w:fldChar'); fldChar4.set(qn('w:fldCharType'), 'end')
    r2 = OxmlElement('w:r')
    r2.append(fldChar3); r2.append(instrText2); r2.append(fldChar4)
    fp._p.append(r2)

# ── 텍스트 헬퍼 ──────────────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=C["title_blue"])
    # 아래 테두리선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),   'single'); b.set(qn('w:sz'),    '8')
    b.set(qn('w:space'), '2');      b.set(qn('w:color'), '1F3864')
    pBdr.append(b); pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=C["h2_blue"])
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run("▶  " + text)
    set_font(run, size=11, bold=True, color=C["h3_gray"])
    return p

def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10, color=C["body"])
    return p

def add_bullet(doc, text, indent=0.8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run("•  " + text)
    set_font(run, size=10, color=C["body"])

def add_note(doc, text, bg=C["note_bg"], label="[참고]"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade_cell(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(label + "  ")
    set_font(r1, size=9, bold=True, color=C["warn_txt"] if "주의" in label or "경고" in label else C["h2_blue"])
    r2 = p.add_run(text)
    set_font(r2, size=9, color=C["body"])
    t.cell(0,0).width = Cm(16.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def add_meta_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        cw(t.cell(i, 0), k, bold=True, size=9,
           bg="D6E4F0" if i % 2 == 0 else "EBF5FB")
        cw(t.cell(i, 1), v, size=9,
           bg="F8FBFD" if i % 2 == 0 else "FFFFFF")
    set_col_widths(t, [Cm(3.8), Cm(12.7)])
    doc.add_paragraph()

def add_spacer(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)

# ── 분석 항목 공통 블록 ──────────────────────────────────────────────────────

def analysis_block(doc, title, hdr_bg,
                   purpose, input_rows, method_rows,
                   output_rows, criteria_rows, visual_rows):

    add_h2(doc, title)
    add_h3(doc, "분석 목적")
    add_body(doc, purpose, indent=0.8)

    # 입력 데이터
    add_h3(doc, "입력 데이터 정의")
    t = doc.add_table(rows=1+len(input_rows), cols=4)
    t.style = "Table Grid"
    hdr_row(t, ["테이블", "컬럼명", "역할", "필터 / 조건"], bg=hdr_bg)
    for i, row in enumerate(input_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            cw(t.cell(i+1,j), val, size=9, bg=bg, center=(j==0))
    set_col_widths(t, [Cm(3.0), Cm(3.5), Cm(2.8), Cm(8.2)])
    add_spacer(doc)

    # 분석 방법
    add_h3(doc, "분석 방법")
    t2 = doc.add_table(rows=1+len(method_rows), cols=3)
    t2.style = "Table Grid"
    hdr_row(t2, ["단계", "기법 / 도구", "상세 설명"], bg=hdr_bg)
    for i, row in enumerate(method_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            cw(t2.cell(i+1,j), val, size=9, bg=bg, center=(j==0))
    set_col_widths(t2, [Cm(1.5), Cm(4.5), Cm(11.5)])
    add_spacer(doc)

    # 출력 지표
    add_h3(doc, "출력 지표 정의")
    t3 = doc.add_table(rows=1+len(output_rows), cols=4)
    t3.style = "Table Grid"
    hdr_row(t3, ["지표명", "계산 방법", "단위", "설명"], bg=hdr_bg)
    for i, row in enumerate(output_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            cw(t3.cell(i+1,j), val, size=9, bg=bg, center=(j==2))
    set_col_widths(t3, [Cm(3.8), Cm(6.0), Cm(1.5), Cm(6.2)])
    add_spacer(doc)

    # 결과 해석 기준
    add_h3(doc, "결과 해석 기준")
    judge_map = {
        "양호": (C["good"],    C["good_txt"]),
        "주의": (C["caution"], C["caution_txt"]),
        "경고": (C["warn"],    C["warn_txt"]),
        "위험": (C["danger"],  C["danger_txt"]),
    }
    t4 = doc.add_table(rows=1+len(criteria_rows), cols=3)
    t4.style = "Table Grid"
    hdr_row(t4, ["판정", "기준값 / 조건", "후속 조치"], bg=hdr_bg)
    for i, (judge, cond, action) in enumerate(criteria_rows):
        bg_hex, txt_rgb = judge_map.get(judge, (C["row_a"], C["body"]))
        cw(t4.cell(i+1,0), judge,  size=9, bg=bg_hex, bold=True, center=True, color=txt_rgb)
        cw(t4.cell(i+1,1), cond,   size=9, bg=bg_hex, color=C["body"])
        cw(t4.cell(i+1,2), action, size=9, bg=bg_hex, color=C["body"])
    set_col_widths(t4, [Cm(1.8), Cm(8.0), Cm(7.7)])
    add_spacer(doc)

    # 시각화 정의
    add_h3(doc, "시각화 정의")
    t5 = doc.add_table(rows=1+len(visual_rows), cols=3)
    t5.style = "Table Grid"
    hdr_row(t5, ["차트 유형", "X축 / 카테고리", "Y축 / 값 / 목적"], bg=hdr_bg)
    for i, row in enumerate(visual_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            cw(t5.cell(i+1,j), val, size=9, bg=bg)
    set_col_widths(t5, [Cm(3.5), Cm(6.5), Cm(7.5)])
    add_spacer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 메인 빌드
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    setup_page(doc)

    # ── 표지 ─────────────────────────────────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph()

    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_main.add_run("데  이  터  분  석  정  의  서")
    set_font(run, size=24, bold=True, color=C["title_blue"])
    p_main.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2  = p_sub.add_run("MDTS  ―  선박 탑재형 엣지 AI 의료 지원 시스템")
    set_font(run2, size=13, color=C["na_gray"])
    p_sub.paragraph_format.space_after = Pt(4)

    p_eng = doc.add_paragraph()
    p_eng.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3  = p_eng.add_run("Maritime Digital Treatment System")
    set_font(run3, size=10, italic=True, color=C["na_gray"])

    for _ in range(8):
        doc.add_paragraph()

    add_meta_table(doc, [
        ("프로젝트명",  "MDTS (선박 탑재형 엣지 AI 의료 지원 시스템)"),
        ("담당 파트",   "빅데이터 분석"),
        ("작성일",      "2026-04-16"),
        ("작성자",      "권태향"),
        ("문서 버전",   "v1.1  (GICOMS 데이터 확보 반영)"),
        ("참조 문서",   "01_분석목표서  /  02_분석용_컬럼_사전  /  03_데이터정의서"),
        ("변경 이력",   "v1.0 최초 작성 (2026-04-16)  |  v1.1 GICOMS 확보 완료 반영, 참조 데이터 URL 추가"),
    ])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. 문서 개요
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "1.  문서 개요")

    add_h2(doc, "1.1  목적")
    add_body(doc,
        "본 문서는 MDTS DB에 축적되는 운용 데이터를 대상으로 수행할 분석 항목 5개(A~E)를 "
        "입력 데이터 · 분석 방법 · 출력 지표 · 결과 해석 기준 · 시각화 방법 단위로 구체적으로 정의한다. "
        "분석 담당자는 본 문서를 단일 기준으로 삼아 일관된 분석을 수행하며, "
        "실제 데이터 수집 후에는 본 문서를 근거로 분석을 즉시 개시할 수 있다.")

    add_h2(doc, "1.2  분석 항목 요약")
    t_sum = doc.add_table(rows=6, cols=4)
    t_sum.style = "Table Grid"
    hdr_row(t_sum, ["목표", "분석 항목명", "핵심 분석 내용", "주요 출력 지표"], bg=C["hdr_dark"])
    summary = [
        ("A", "위험 등급 예측 모델 검증",    "AI risk_level vs 실제 처치 일치율 측정",          "일치율(%), Weighted F1, 혼동 행렬"),
        ("B", "생체 신호 임계값 분석",        "등급별 심박수·SpO2·체온 분포 비교",               "등급별 Mean/Q1/Q3, ANOVA p-value, AUC"),
        ("C", "승조원 특성 × 위험도 연관성",  "직책·연령대·성별 × 위험 등급 교차 분석",          "카이제곱 p-value, 상대 위험도(RR)"),
        ("D", "처치 반응 시간 분석",           "AI 분석 → 처치 완료 소요 시간 측정",              "등급별 평균·P90·SLA 초과율"),
        ("E", "육상 동기화 장애 패턴 분석",   "sync_status=FALSE 시간대·테이블 분포",            "전체 실패율, 시간대별 히트맵"),
    ]
    for i, (code, name, content, output) in enumerate(summary):
        bg = C["accent"] if i%2==0 else C["row_b"]
        cw(t_sum.cell(i+1,0), code,    size=10, bold=True, bg=bg,
           color=C["title_blue"], center=True)
        cw(t_sum.cell(i+1,1), name,    size=9,  bold=True, bg=bg, color=C["body"])
        cw(t_sum.cell(i+1,2), content, size=9,  bg=bg, color=C["body"])
        cw(t_sum.cell(i+1,3), output,  size=9,  bg=bg, color=C["body"])
    set_col_widths(t_sum, [Cm(1.3), Cm(4.5), Cm(6.2), Cm(5.5)])
    add_spacer(doc)

    add_h2(doc, "1.3  공통 전처리 기준")
    preproc = [
        ("결측값 처리",      "heart_rate=0, spo2=0.00, temperature=0.00 인 레코드 → 센서 미측정으로 간주, 분석 제외"),
        ("이상값 제거",      "heart_rate <30 or >250 / spo2 >100 / temperature <30 or >45 → 측정 오류로 제외"),
        ("연령 파생 변수",   "age = TIMESTAMPDIFF(YEAR, birthdate, NOW()) / age_group: 20대이하/<30, 30대/30~39, 40대이상/40+"),
        ("risk_label 변환", "risk_level '1'→GREEN / '2'→YELLOW / '3'→ORANGE / '4'→RED"),
        ("반응 시간 계산",   "response_time_min = TIMESTAMPDIFF(MINUTE, analyzed_at, created_at) — 음수·>1440분 제외"),
        ("동기화 파생",      "sync_hour = HOUR(synced_at) / sync_dow = DAYOFWEEK(synced_at) [1=일~7=토]"),
        ("PII 처리",         "name 컬럼 분석 제외 — crew_id로만 식별. 결과 출력 시 마스킹 권장"),
    ]
    t_pre = doc.add_table(rows=1+len(preproc), cols=2)
    t_pre.style = "Table Grid"
    hdr_row(t_pre, ["전처리 항목", "처리 기준"], bg=C["hdr_dark"])
    for i, (k, v) in enumerate(preproc):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        cw(t_pre.cell(i+1,0), k, bold=True, size=9, bg=bg, color=C["h2_blue"])
        cw(t_pre.cell(i+1,1), v, size=9, bg=bg)
    set_col_widths(t_pre, [Cm(3.5), Cm(14.0)])
    add_spacer(doc)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 2. 분석 데이터 현황 및 출처
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "2.  분석 데이터 현황 및 출처")
    add_body(doc,
        "본 프로젝트에서 사용되는 데이터는 세 가지로 구분된다: "
        "① MDTS 시스템이 운용 중 자체 생성하는 DB 데이터(분석 목표 A~E 주 데이터), "
        "② AI 모델 학습에 사용된 외부 데이터셋, "
        "③ 도메인 기준값 수립에 활용된 참조 자료.")

    # ── 2.1 MDTS DB 운용 데이터 ──────────────────────────────────────────────
    add_h2(doc, "2.1  MDTS DB 운용 데이터  (분석 목표 A~E 주 데이터)")
    t_db = doc.add_table(rows=6, cols=5)
    t_db.style = "Table Grid"
    hdr_row(t_db, ["테이블명", "논리명", "데이터 생성 주체", "수집 방법", "확보 현황"], bg=C["hdr_navy"])
    db_rows = [
        ("tb_crew",     "승조원 정보",     "관리자 수동 입력",          "웹/앱 UI — 승선 시 1회 등록"),
        ("tb_vital",    "생체 신호",       "IoT 센서 자동 측정",        "심박수·SpO2·체온계 → 엣지 디바이스 자동 저장"),
        ("tb_analysis", "AI 분석 결과",    "MDTS AI 엔진 자동 생성",    "이미지 업로드 + 생체 신호 → AI 추론 후 저장"),
        ("tb_firstaid", "응급 처치 기록",  "승조원 수동 입력",           "처치 완료 후 앱에서 조치 사항 입력"),
        ("tb_logs",     "동기화 로그",     "시스템 자동 기록",           "육상 서버 전송 시도 시 결과 자동 저장"),
    ]
    for i, (tbl, logic, src, method) in enumerate(db_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        cw(t_db.cell(i+1,0), tbl,    size=9, bold=True, bg=bg, color=C["h2_blue"], center=True)
        cw(t_db.cell(i+1,1), logic,  size=9, bg=bg)
        cw(t_db.cell(i+1,2), src,    size=9, bg=bg)
        cw(t_db.cell(i+1,3), method, size=9, bg=bg)
        cw(t_db.cell(i+1,4), "미확보\n(시스템 운용 전)", size=9, bg="FFEBEE",
           bold=True, color=C["ng_red"], center=True)
    set_col_widths(t_db, [Cm(2.5), Cm(2.5), Cm(3.2), Cm(5.5), Cm(2.8)])
    add_spacer(doc)
    add_note(doc,
        "현재(2026-04-16) 기준 5개 테이블 모두 실제 운용 데이터 미확보 상태입니다. "
        "분석 목표 A~E는 MDTS 시스템 운용 개시 후 데이터 누적 시 즉시 실행 가능합니다.",
        bg=C["alert_bg"], label="[주의]")

    # ── 2.2 AI 모델 학습 데이터 ──────────────────────────────────────────────
    add_h2(doc, "2.2  AI 모델 학습 데이터")
    add_body(doc,
        "tb_analysis의 diagnosis 및 risk_level 값은 아래 데이터셋으로 학습된 AI 모델이 생성한다. "
        "분석 결과 해석 시 각 모델의 학습 데이터 특성을 함께 고려해야 한다.")
    add_spacer(doc, 4)

    ml_rows = [
        # (데이터셋, 사용모델, 출처, 로컬경로, 규모, 클래스/컬럼, 현황, 비고)
        ("GICOMS\n해양사고 통계\n(재구성)",
         "RandomForest /\nXGBoost\n(이진·다중 분류)",
         "중앙해양안전심판원(KMST)\n해양수산부 공식 통계 기반\nhttps://www.maritimeaccident.org",
         "M_MEDIC_v2/01_data/raw/\ngicoms_marine_accidents\n_2014_2024.csv",
         "48건\n(2014~2024\n11개 연도)",
         "연도·사고유형(8종)·선박종류(7종)\n사고원인(9종)·발생해역(4종)\n사상자수",
         "확보 완료\n(재구성 데이터)",
         "공식 통계 수치 기반 재구성.\n실제 GICOMS 원본 확보 시 교체 권장"),
        ("해양사고\n합성 데이터",
         "위와 동일\n(보조 학습용)",
         "해양안전심판원 통계 패턴 기반\n자체 합성 생성",
         "M_MEDIC_v2/01_data/processed/\nmarine_accidents\n_augmented.csv",
         "1,500건\n(합성)",
         "연도·사고유형·선박종류·사고원인\n발생해역 → 사상자발생(0/1)\n위험등급(4단계)",
         "확보 완료\n(합성 데이터)",
         "GICOMS 재구성 데이터와 병행 사용\n또는 단독 사용 가능"),
        ("Wound\nClassification\nDataset",
         "MobileNetV3-Small\n(외상 분류\n10클래스)",
         "Kaggle 공개 데이터셋\n(wound-classification-system)",
         "D:/GeminiUniverse/vscode-workspace/\nwound-classification-system/\nWound_dataset copy/",
         "1,946장\n10클래스",
         "Abrasion / Bruises / Burn / Cut\nIngrown_Nail / Laceration\nLaseration / Normal\nStab_Wound / Surgical_Wound",
         "확보 완료",
         "클래스 불균형 존재\nWeightedRandomSampler 적용"),
        ("WOUND_DATA\n피부 병변\n이미지",
         "EfficientNet-B3\n(피부질환 분류\n7클래스)",
         "ISIC Archive / Kaggle\nhttps://www.kaggle.com/datasets/\nkmader/skin-lesion-analysis\n-toward-melanoma-detection",
         "M_MEDIC_v2/01_data/raw/\nWOUND_DATA/\n(Kaggle 다운로드)",
         "10,015장\n7클래스",
         "mel / nv / bcc\nakiec / bkl / df / vasc",
         "확보 완료",
         "멀티모달 구조\n이미지 + 나이·성별·부위 메타데이터 융합"),
    ]
    t_ml = doc.add_table(rows=1+len(ml_rows), cols=8)
    t_ml.style = "Table Grid"
    hdr_row(t_ml,
            ["데이터셋명","사용 모델","출처 / 제공처","로컬 경로","규모","클래스 / 주요 컬럼","확보 현황","비고"],
            bg=C["hdr_teal"],
            sizes=[9,9,9,9,9,9,9,9])
    status_colors = {
        "확보 완료\n(재구성 데이터)": C["ok_green"],
        "확보 완료\n(합성 데이터)":   C["ok_green"],
        "확보 완료":                  C["ok_green"],
        "미확보\n(수동 다운로드 필요)": C["ng_red"],
    }
    for i, row in enumerate(ml_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            is_status = (j == 6)
            is_url    = (j == 2)
            is_path   = (j == 3)
            sc = status_colors.get(val, None) if is_status else None
            cw(t_ml.cell(i+1,j), val, size=8, bg=bg,
               color=sc if sc else (C["url_blue"] if is_url or is_path else None),
               italic=(is_url or is_path),
               bold=is_status)
    set_col_widths(t_ml, [Cm(2.5), Cm(2.5), Cm(3.2), Cm(3.5), Cm(1.5), Cm(3.8), Cm(2.0), Cm(3.0)])
    add_spacer(doc)
    add_note(doc,
        "GICOMS 데이터는 중앙해양안전심판원(KMST) 및 해양수산부 공식 통계 수치를 기반으로 재구성된 데이터입니다. "
        "원본 GICOMS 데이터 확보 시 M_MEDIC_v2/01_data/raw/ 경로에 저장하면 학습 파이프라인이 자동 감지합니다.",
        label="[참고]")

    # ── 2.3 참조 데이터 ──────────────────────────────────────────────────────
    add_h2(doc, "2.3  참조 데이터  (도메인 기준값 출처)")
    add_body(doc,
        "아래 자료는 직접 분석 대상이 아니라, 생체 신호 정상 범위·위험 등급 기준·응급처치 가이드의 "
        "근거 자료로 사용된다.")
    add_spacer(doc, 4)

    ref_rows = [
        ("심박수 정상 범위\n(60~100 bpm)",
         "2020 한국심폐소생술\n가이드라인\n(KACPR, 대한심폐소생술협회)",
         "https://www.kacpr.org\n→ 가이드라인 자료실\n→ 2020년판 PDF",
         "성인 정상: 60~100 bpm\n서맥: <60  /  빈맥: >100",
         "목표 B — heart_rate\n임계값 기준"),
        ("SpO2 정상 범위\n(≥ 95%)",
         "임상진료지침정보센터\n(대한의학회 운영)",
         "https://www.guideline.or.kr\n→ 호흡기내과 관련\n지침 참조",
         "정상: ≥ 95%\n주의: 90~94%\n위험: < 90% (저산소증)",
         "목표 B — spo2\n임계값 기준"),
        ("체온 정상 범위\n(36.1~37.2°C)",
         "서울아산병원 의학정보\n활력징후 기준",
         "https://www.amc.seoul.kr/asan/\nhealthinfo/easymediterm/\neasyMediTermDetail.do\n?dictId=4900",
         "정상: 36.1~37.2°C\n미열: 37.3~38.0\n고열: ≥ 38.1  /  저체온: ≤ 35.0",
         "목표 B — temperature\n임계값 기준"),
        ("선박의료 관련 법령\n(의약품 비치·의료관리)",
         "선원법 시행규칙 제52조\n(의료용품 비치 및 관리)\n※ 별표 5의5는 기초안전교육\n   규정으로 의료와 무관",
         "국가법령정보센터 (법제처)\nhttps://www.law.go.kr\n→ '선원법 시행규칙' 검색\n→ 제52조 확인",
         "선박 내 의약품 비치 기준\n의료 담당자 지정\n의료용품 관리 절차",
         "guide_text 작성 기준\n목표 A — 처치 일치율\n레이블링 근거"),
        ("해양사고 통계\n(공표 자료)",
         "중앙해양안전심판원\n연간 해양사고 통계\n(2014~2024 공표본)",
         "https://www.maritimeaccident.org\n→ 통계자료\n→ 연도별 해양사고 통계",
         "사고유형·선박종류·사고원인별\n연도별 건수 및 사상자 수",
         "GICOMS 재구성 데이터\n생성 근거\nML 도메인 규칙 설정"),
        ("위험 등급 기준",
         "MDTS 내부 설계 기준\n(도메인 전문가 협의)",
         "자체 설계\n(별도 외부 문서 없음)",
         "1=GREEN / 2=YELLOW\n3=ORANGE / 4=RED",
         "risk_level 기준값\n목표 A~D 전체"),
    ]
    t_ref = doc.add_table(rows=1+len(ref_rows), cols=5)
    t_ref.style = "Table Grid"
    hdr_row(t_ref, ["참조 데이터명","출처 / 제공처","접근 URL / 경로","주요 내용","활용 분석 목표"],
            bg=C["hdr_dark"])
    for i, row in enumerate(ref_rows):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            is_url = (j == 2)
            cw(t_ref.cell(i+1,j), val, size=8, bg=bg,
               color=C["url_blue"] if is_url else None,
               italic=is_url)
    set_col_widths(t_ref, [Cm(2.8), Cm(3.5), Cm(4.5), Cm(4.0), Cm(2.7)])
    add_spacer(doc)

    # ── 2.4 데이터 확보 현황 요약 ────────────────────────────────────────────
    add_h2(doc, "2.4  데이터 확보 현황 요약")
    status_rows = [
        ("MDTS DB 운용 데이터 (5개 테이블)",           "5개 테이블",    "미확보",     "시스템 운용 개시 후 자동 수집"),
        ("GICOMS 해양사고 통계 (재구성)",               "48건 / 6컬럼",  "확보 완료",  "01_data/raw/gicoms_marine_accidents_2014_2024.csv"),
        ("해양사고 합성 데이터",                         "1,500건",       "확보 완료",  "01_data/processed/marine_accidents_augmented.csv"),
        ("Wound Classification Dataset",               "1,946장",       "확보 완료",  "wound-classification-system/Wound_dataset copy/"),
        ("WOUND_DATA 피부 병변 이미지",                   "10,015장",      "확보 완료",  "Kaggle 다운로드 완료 → 01_data/raw/WOUND_DATA/"),
        ("심박수·SpO2·체온 정상 범위 기준",              "-",             "URL 확인",   "kacpr.org / guideline.or.kr / amc.seoul.kr"),
        ("선원법 시행규칙 제52조 (선박의료 관련)",        "-",             "URL 확인",   "law.go.kr → '선원법 시행규칙' 검색 → 제52조"),
    ]
    t_st = doc.add_table(rows=1+len(status_rows), cols=4)
    t_st.style = "Table Grid"
    hdr_row(t_st, ["데이터셋","규모","확보 현황","비고 / 경로"], bg=C["hdr_dark"])
    status_color_map = {
        "확보 완료": (C["good"],    C["ok_green"]),
        "미확보":    (C["danger"],  C["ng_red"]),
        "URL 확인":  (C["caution"], C["caution_txt"]),
    }
    for i, (name, size, status, note) in enumerate(status_rows):
        bg_row = C["row_a"] if i%2==0 else C["row_b"]
        bg_s, txt_s = status_color_map.get(status, (bg_row, C["body"]))
        cw(t_st.cell(i+1,0), name,   size=9, bold=True, bg=bg_row)
        cw(t_st.cell(i+1,1), size,   size=9, bg=bg_row, center=True)
        cw(t_st.cell(i+1,2), status, size=9, bold=True,
           bg=bg_s, color=txt_s, center=True)
        cw(t_st.cell(i+1,3), note,   size=9, bg=bg_row,
           color=C["url_blue"], italic=True)
    set_col_widths(t_st, [Cm(5.5), Cm(2.2), Cm(2.5), Cm(7.3)])
    add_spacer(doc)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 3. 분석 항목 상세 정의
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "3.  분석 항목 상세 정의")

    # ── 목표 A ────────────────────────────────────────────────────────────────
    analysis_block(
        doc,
        title     = "3.1  목표 A  ―  위험 등급 예측 모델 검증",
        hdr_bg    = "1F497D",
        purpose   = (
            "AI가 산출한 위험 등급(risk_level)이 실제 수행된 처치 내용(action_taken)과 "
            "얼마나 일치하는지 정량적으로 측정한다. 모델 신뢰도 지표를 확보하고, "
            "등급별 과대·과소 예측 패턴을 식별하여 AI 모델 개선 방향을 도출한다."),
        input_rows = [
            ("tb_analysis", "risk_level",   "종속변수(Y) — AI 예측 위험 등급",  "IN ('1','2','3','4')"),
            ("tb_analysis", "analyzed_at",  "분석 일시 — 기간 필터 기준",        "기간 설정 시 활용"),
            ("tb_firstaid", "action_taken", "실제 처치 내용 — 검증 기준",        "NOT NULL AND != ''"),
            ("tb_firstaid", "analysis_id",  "tb_analysis 조인 키",              "INNER JOIN ON analysis_id"),
        ],
        method_rows = [
            ("1", "데이터 조인",               "tb_analysis ↔ tb_firstaid INNER JOIN. 처치 기록 없는 분석 건 제외"),
            ("2", "처치 적절성 레이블링",       "action_taken 키워드(지혈·심폐소생·산소·냉각·봉합 등) 추출 → 처치 수준 1~4 분류"),
            ("3", "혼동 행렬 산출",             "AI 예측 등급 vs 처치 수준 등급의 4×4 Confusion Matrix"),
            ("4", "Precision/Recall/F1 계산", "등급별 Precision, Recall, F1-Score (weighted average)"),
            ("5", "과대·과소 예측 분석",        "혼동 행렬 상삼각(과대) vs 하삼각(과소) 비중 — 과소 예측이 안전 위협"),
        ],
        output_rows = [
            ("전체 일치율",      "일치 건수 / 전체 × 100",                "%",  "AI 예측 = 실제 처치 수준인 비율"),
            ("Weighted F1",      "Σ(등급별 F1 × 등급 비중)",              "-",  "전체 모델 대표 성능 지표"),
            ("과소 예측 비율",   "하삼각 합계 / 전체 × 100",              "%",  "실제보다 낮은 등급 예측 비율 — 안전 위협 지표"),
            ("4×4 혼동 행렬",    "행=실제, 열=예측, 값=건수",             "건", "등급 간 오분류 전체 패턴"),
        ],
        criteria_rows = [
            ("양호", "일치율 ≥ 80%  AND  과소 예측 비율 < 10%",    "현 AI 모델 유지. 정기 재검증 (분기 1회)"),
            ("주의", "일치율 60~79%  OR  과소 예측 10~19%",         "모델 재학습 검토. 과소 예측 등급 집중 분석"),
            ("경고", "일치율 40~59%  OR  과소 예측 20~29%",         "긴급 재학습. 등급 상향 보정 로직 임시 적용"),
            ("위험", "일치율 < 40%  OR  과소 예측 ≥ 30%",          "AI 판단 신뢰 불가. 수동 판단 전환 및 즉시 보고"),
        ],
        visual_rows = [
            ("히트맵 (혼동 행렬)",  "예측 등급 (1~4)",            "실제 등급 (1~4) — 대각선=정답"),
            ("막대 그래프",          "위험 등급 (GREEN~RED)",       "Precision / Recall / F1 비교"),
            ("원형 차트",            "-",                          "일치 / 과대 예측 / 과소 예측 비율"),
        ],
    )

    add_page_break(doc)

    # ── 목표 B ────────────────────────────────────────────────────────────────
    analysis_block(
        doc,
        title     = "3.2  목표 B  ―  생체 신호와 위험 등급 간 임계값 분석",
        hdr_bg    = "375623",
        purpose   = (
            "위험 등급(GREEN~RED)별로 심박수·산소포화도·체온의 수치 분포를 비교하여 "
            "각 등급을 구분하는 임계값 범위를 데이터 기반으로 도출한다. "
            "결과는 현장 판단 기준표 및 AI 모델 튜닝 근거로 활용된다."),
        input_rows = [
            ("tb_vital",    "heart_rate",  "독립변수(X) — 심박수",    "0 제외, 30~250"),
            ("tb_vital",    "spo2",        "독립변수(X) — 산소포화도", "0.00 제외, ≤100.00"),
            ("tb_vital",    "temperature", "독립변수(X) — 체온",       "0.00 제외, 30~45"),
            ("tb_analysis", "risk_level",  "그룹 기준 — 위험 등급",    "JOIN ON vital_id"),
        ],
        method_rows = [
            ("1", "데이터 조인",            "tb_vital ↔ tb_analysis (vital_id). risk_label 파생 추가"),
            ("2", "등급별 기술통계",         "risk_label 그룹별 mean/std/Q1/Q3/max 산출 (3지표 × 4등급)"),
            ("3", "박스플롯 시각화",         "등급별 분포 및 이상값 시각 확인"),
            ("4", "ANOVA / Kruskal-Wallis", "4개 등급 간 평균 차이 통계적 유의성 검정"),
            ("5", "임계값 도출",             "인접 등급 Q1~Q3 경계값 비교로 구분선 후보 설정"),
            ("6", "ROC/AUC (이진 변환)",    "RED(4) vs 기타 이진 분류 기준 각 지표 AUC 산출"),
        ],
        output_rows = [
            ("등급별 기술통계표",  "GROUP BY risk_level → MEAN/STD/Q1/Q3",  "-",  "3지표 × 4등급 = 12개 세트"),
            ("임계값 범위표",      "인접 등급 Q1~Q3 경계 비교",               "-",  "현장 판단 기준 (잠정)"),
            ("ANOVA p-value",     "등급 간 분산 분석",                        "-",  "< 0.05: 유의미한 등급 간 차이"),
            ("AUC (RED vs 기타)", "ROC 곡선 아래 면적",                       "-",  "0.8 이상: 해당 지표 RED 구분 유효"),
        ],
        criteria_rows = [
            ("양호", "ANOVA p < 0.05  AND  모든 지표 AUC ≥ 0.75",   "임계값을 AI 모델 기준으로 채택"),
            ("주의", "일부 지표 p ≥ 0.05 또는 AUC 0.60~0.74",       "해당 지표 단독 사용 지양. 복합 조건 검토"),
            ("경고", "2개 이상 지표 p ≥ 0.05 또는 AUC < 0.60",      "추가 feature 탐색 필요"),
        ],
        visual_rows = [
            ("박스플롯 (3종)", "위험 등급 (GREEN~RED)",     "심박수·SpO2·체온 각 분포"),
            ("히트맵",         "위험 등급 × 지표",           "셀 값 = 평균값"),
            ("ROC 곡선",       "False Positive Rate",       "True Positive Rate — AUC 표시"),
        ],
    )

    add_page_break(doc)

    # ── 목표 C ────────────────────────────────────────────────────────────────
    analysis_block(
        doc,
        title     = "3.3  목표 C  ―  승조원 특성과 위험도 연관성 분석",
        hdr_bg    = "843C0C",
        purpose   = (
            "직책·연령대·성별 등 승조원의 인적 특성과 위험 등급 분포 간의 통계적 연관성을 검토한다. "
            "고위험군 직책·연령대를 식별하여 사전 예방 모니터링 대상을 선정한다."),
        input_rows = [
            ("tb_crew",     "position",  "독립변수(X) — 직책",          "NULL 제외, 정규화 후 사용"),
            ("tb_crew",     "birthdate", "파생변수 기준 — 생년월일",     "→ age, age_group"),
            ("tb_crew",     "gender",    "독립변수(X) — 성별",           "'M' / 'F'"),
            ("tb_analysis", "risk_level","종속변수(Y) — 위험 등급",       "JOIN ON crew_id"),
        ],
        method_rows = [
            ("1", "파생 변수 생성",       "age = TIMESTAMPDIFF / age_group 구간화 (20대이하/30대/40대이상)"),
            ("2", "직책 정규화",          "선장/항해사/기관사/조리원/일반선원/기타 6개 코드로 매핑"),
            ("3", "그룹별 빈도 분석",     "직책/연령대/성별 × risk_level 교차표(Crosstab) 산출"),
            ("4", "카이제곱 검정",         "각 특성 변수별 독립성 검정 — p < 0.05: 위험 등급에 유의한 영향"),
            ("5", "상대 위험도(RR) 계산", "그룹 RED 비율 / 전체 RED 비율 — 1.5 이상 시 고위험군"),
        ],
        output_rows = [
            ("그룹별 위험 등급 분포표", "GROUP BY 특성 × risk_level",      "%/건","직책·연령대·성별 × 4등급 교차표"),
            ("카이제곱 p-value",        "독립성 검정",                       "-",  "< 0.05: 해당 특성이 위험 등급과 연관"),
            ("상대 위험도(RR)",          "그룹 RED 비율 / 전체 RED 비율",    "배", "1.5 이상: 고위험군 지정"),
            ("고위험군 목록",            "RR ≥ 1.5 그룹 추출",               "-",  "사전 예방 모니터링 대상 목록"),
        ],
        criteria_rows = [
            ("양호", "모든 특성 변수 카이제곱 p ≥ 0.05",         "특성 무관 균등 위험 분포 — 개인별 생체 신호 중심 관리"),
            ("주의", "1개 특성 p < 0.05  AND  RR < 2.0",         "해당 그룹 모니터링 주기 단축"),
            ("경고", "2개 이상 특성 p < 0.05  OR  RR 2.0~2.9",   "고위험군 대상 승선 전 추가 건강검진 권고"),
            ("위험", "RR ≥ 3.0 그룹 존재",                       "해당 직책·연령대 운항 전 필수 의료 점검 정책 수립 요청"),
        ],
        visual_rows = [
            ("스택 막대 그래프 (3종)", "직책 / 연령대 / 성별",   "각 등급 비율 누적 — 그룹별 위험 분포"),
            ("버블 차트",              "연령 (X) × 직책 (Y)",    "버블 크기 = RED 비율"),
            ("히트맵",                 "직책 × 연령대",           "셀 값 = RED 비율(%)"),
        ],
    )

    add_page_break(doc)

    # ── 목표 D ────────────────────────────────────────────────────────────────
    analysis_block(
        doc,
        title     = "3.4  목표 D  ―  처치 반응 시간 분석",
        hdr_bg    = "3F3F76",
        purpose   = (
            "AI 분석 완료 시점(analyzed_at)부터 응급처치 완료 시점(created_at)까지의 "
            "소요 시간(response_time_min)을 위험 등급별로 분석한다. "
            "등급별 처치 반응 시간 목표(SLA)를 수립하고, 초과 케이스를 파악하여 "
            "시스템 응답 속도 및 현장 교육 효과를 검증한다."),
        input_rows = [
            ("tb_analysis", "analyzed_at",  "파생변수 기준 — AI 분석 완료 시각",  "NOT NULL"),
            ("tb_analysis", "risk_level",   "그룹 기준 — 위험 등급",              "1~4"),
            ("tb_firstaid", "created_at",   "파생변수 기준 — 처치 완료 시각",      "NOT NULL"),
            ("tb_firstaid", "analysis_id",  "조인 키",                             "JOIN ON analysis_id"),
        ],
        method_rows = [
            ("1", "반응 시간 파생",     "response_time_min = TIMESTAMPDIFF(MINUTE, analyzed_at, created_at)"),
            ("2", "이상값 제거",         "< 0 (오류) 또는 > 1440분 (24h 초과) 제외"),
            ("3", "등급별 기술통계",     "risk_label 그룹별 mean / median / std / P90 계산"),
            ("4", "SLA 초과율 계산",     "등급별 목표 시간 초과 건수 / 전체 × 100"),
            ("5", "시계열 추세 분석",    "월별/분기별 평균 반응 시간 — 시스템 개선 효과 측정"),
        ],
        output_rows = [
            ("response_time_min",    "TIMESTAMPDIFF(MINUTE, analyzed_at, created_at)",  "분", "핵심 파생 변수"),
            ("등급별 평균 반응 시간", "GROUP BY risk_level → AVG",                       "분", "4등급 × 평균/중앙값/P90"),
            ("SLA 초과율",            "초과 건수 / 전체 × 100",                          "%",  "목표 시간 대비 초과 비율"),
            ("P90 반응 시간",         "90번째 백분위값",                                  "분", "극단값 제외 '거의 최대' 시간"),
        ],
        criteria_rows = [
            ("양호", "RED ≤ 3분 / ORANGE ≤ 5분 / YELLOW ≤ 10분 / GREEN ≤ 30분",  "현 시스템 유지"),
            ("주의", "1개 등급 SLA 초과율 10~19%",                                  "처치 가이드 명확화 및 재교육"),
            ("경고", "1개 등급 초과율 20~29%  OR  RED 평균 > 5분",                  "현장 프로세스 재설계. 알림 강화"),
            ("위험", "RED 초과율 ≥ 30%  OR  RED 평균 > 10분",                       "긴급 대응 체계 재검토. 육상팀 즉시 보고"),
        ],
        visual_rows = [
            ("박스플롯",         "위험 등급 (GREEN~RED)",  "response_time_min 분포 + SLA 기준선"),
            ("꺾은선 (추세)",    "월 / 분기",              "등급별 평균 반응 시간 시계열"),
            ("누적 분포(CDF)",   "반응 시간 (분)",          "누적 비율 — SLA 달성 비율 확인"),
        ],
    )

    add_page_break(doc)

    # ── 목표 E ────────────────────────────────────────────────────────────────
    analysis_block(
        doc,
        title     = "3.5  목표 E  ―  육상 동기화 장애 패턴 분석",
        hdr_bg    = "7030A0",
        purpose   = (
            "육상 서버로의 데이터 전송 실패(sync_status=FALSE) 패턴을 "
            "시간대·요일·테이블 단위로 분석한다. 동기화 취약 시간대와 자주 실패하는 테이블을 식별하여 "
            "인프라 개선 우선순위를 도출하고, 통신 장애 예측 모델의 기초 데이터를 마련한다."),
        input_rows = [
            ("tb_logs", "sync_status", "종속변수(Y) — 전송 성공/실패",  "TRUE=1, FALSE=0 수치화"),
            ("tb_logs", "synced_at",   "파생변수 기준 — 전송 일시",      "→ sync_hour, sync_dow"),
            ("tb_logs", "table_name",  "독립변수(X) — 대상 테이블",      "허용값 정규화 후 사용"),
        ],
        method_rows = [
            ("1", "파생 변수 생성",    "sync_hour = HOUR(synced_at) / sync_dow = DAYOFWEEK(synced_at)"),
            ("2", "전체 실패율",        "SUM(sync_status=0) / COUNT(*) × 100 → 기준선"),
            ("3", "시간대별 실패율",    "GROUP BY sync_hour → 실패율. 피크 시간대 식별"),
            ("4", "요일별 실패율",      "GROUP BY sync_dow → 실패율. 주말 vs 평일 비교"),
            ("5", "테이블별 실패율",    "GROUP BY table_name → 실패율"),
            ("6", "복합 분석",          "sync_hour × table_name 3차원 집계 → 최악 조합 도출"),
        ],
        output_rows = [
            ("전체 실패율",       "FALSE 건수 / 전체 × 100",                "%",  "시스템 동기화 신뢰도 지표"),
            ("시간대별 실패율",   "GROUP BY sync_hour → 실패율",             "%",  "24개 시간대 × 실패율"),
            ("테이블별 실패율",   "GROUP BY table_name → 실패율",            "%",  "테이블별 취약도 비교"),
            ("최악 조합 TOP5",    "시간대 × 테이블 교차 실패율 상위 5",      "%",  "인프라 개선 최우선 대상"),
        ],
        criteria_rows = [
            ("양호", "전체 실패율 < 5%  AND  모든 시간대 < 15%",          "현 인프라 유지. 월 1회 모니터링"),
            ("주의", "전체 실패율 5~9%  OR  일부 시간대 15~24%",          "취약 시간대 재전송 로직 추가 검토"),
            ("경고", "전체 실패율 10~19%  OR  일부 시간대 25~39%",        "통신 장비 점검 및 백업 전송 경로 마련"),
            ("위험", "전체 실패율 ≥ 20%  OR  일부 시간대 ≥ 40%",         "서버 연결 아키텍처 재설계. 오프라인 버퍼 전략 즉시 수립"),
        ],
        visual_rows = [
            ("히트맵",         "요일 (X) × 시간대 (Y)",   "셀 값 = 실패율(%) — 취약 구간 식별"),
            ("막대 그래프",    "테이블명",                  "테이블별 실패율 비교"),
            ("꺾은선 (추세)", "날짜",                      "일별 실패율 — 이상 급증 구간 감지"),
        ],
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 4. 파생 변수 상세 정의
    # ══════════════════════════════════════════════════════════════════════════
    add_h1(doc, "4.  파생 변수 상세 정의")

    deriv = [
        ("age",              "TIMESTAMPDIFF(YEAR, birthdate, NOW())",
         "tb_crew.birthdate", "INT",  "0 이상 정수 (예: 34)",  "목표 C", "분석 시점 기준 만 나이"),
        ("age_group",        "CASE age <20→'20대이하' <30→'30대' ELSE '40대이상'",
         "age (파생)",        "VARCHAR","20대이하/30대/40대이상","목표 C", "연령 구간화"),
        ("risk_label",       "1→GREEN / 2→YELLOW / 3→ORANGE / 4→RED",
         "tb_analysis.risk_level","VARCHAR","GREEN/YELLOW/ORANGE/RED","목표 A~D","숫자 ENUM → 의미 문자열"),
        ("response_time_min","TIMESTAMPDIFF(MINUTE, analyzed_at, created_at)",
         "analyzed_at, created_at","INT","0 이상. 제외: <0 or >1440","목표 D","음수 시 데이터 오류로 제외"),
        ("sync_hour",        "HOUR(synced_at)",
         "tb_logs.synced_at","INT","0 ~ 23","목표 E","시간대별 실패율 기준"),
        ("sync_dow",         "DAYOFWEEK(synced_at)",
         "tb_logs.synced_at","INT","1(일) ~ 7(토)","목표 E","MySQL DAYOFWEEK 기준"),
        ("sync_fail",        "CASE sync_status=FALSE THEN 1 ELSE 0",
         "tb_logs.sync_status","INT","0 또는 1","목표 E","집계 편의를 위해 수치화"),
    ]
    t_d = doc.add_table(rows=1+len(deriv), cols=7)
    t_d.style = "Table Grid"
    hdr_row(t_d, ["파생 변수명","계산식","원본 컬럼","타입","허용값 / 예시","사용 목표","비고"],
            bg=C["hdr_dark"])
    for i, row in enumerate(deriv):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            cw(t_d.cell(i+1,j), val, size=8, bg=bg, center=(j in (3,5)))
    set_col_widths(t_d, [Cm(2.5), Cm(5.5), Cm(3.5), Cm(1.3), Cm(2.8), Cm(1.8), Cm(3.1)])
    add_spacer(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 5. 분석 결과물 명세
    # ══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_h1(doc, "5.  분석 결과물 명세")

    deliverables = [
        ("A", "혼동 행렬 이미지",            "result_A_confusion_matrix.png",  "PNG", "4×4 히트맵 — AI 예측 vs 실제 처치","O"),
        ("A", "모델 성능 지표 표",            "result_A_metrics.csv",           "CSV", "Precision/Recall/F1, 전체 일치율", "O"),
        ("B", "생체 신호 임계값 범위표",      "result_B_threshold_table.csv",   "CSV", "등급별 mean/std/Q1/Q3 (3지표×4등급)","O"),
        ("B", "박스플롯 이미지 (3종)",        "result_B_boxplot_{지표}.png",    "PNG", "heart_rate / spo2 / temperature",  "O"),
        ("B", "ANOVA 검정 결과",              "result_B_anova.csv",             "CSV", "3개 지표별 F통계량, p-value",        "-"),
        ("C", "그룹별 위험 등급 분포표",      "result_C_crosstab.csv",          "CSV", "직책/연령대/성별 × 4등급 교차표",    "O"),
        ("C", "스택 막대 그래프 (3종)",       "result_C_stacked_{특성}.png",    "PNG", "직책 / 연령대 / 성별 각 1개",        "O"),
        ("C", "고위험군 목록",                "result_C_high_risk_groups.csv",  "CSV", "RR ≥ 1.5 그룹 목록",                "O"),
        ("D", "처치 반응 시간 통계표",        "result_D_response_time.csv",     "CSV", "등급별 mean/median/P90/SLA초과율",   "O"),
        ("D", "반응 시간 박스플롯",           "result_D_boxplot.png",           "PNG", "4등급 × 반응 시간 분포 + SLA 기준선","O"),
        ("D", "월별 추세 꺾은선",             "result_D_trend.png",             "PNG", "등급별 월 평균 반응 시간 시계열",     "-"),
        ("E", "동기화 실패율 히트맵",         "result_E_heatmap.png",           "PNG", "요일 × 시간대 실패율 히트맵",         "O"),
        ("E", "테이블별 실패율 표",           "result_E_table_fail.csv",        "CSV", "테이블별 실패율 및 건수",              "O"),
        ("E", "최악 조합 TOP5",               "result_E_worst_combo.csv",       "CSV", "시간대 × 테이블 교차 실패율 상위 5",  "O"),
    ]
    t_del = doc.add_table(rows=1+len(deliverables), cols=6)
    t_del.style = "Table Grid"
    hdr_row(t_del, ["목표","산출물명","권장 파일명","형식","주요 내용","발표\n활용"], bg=C["hdr_dark"])
    goal_bgs = {"A":C["accent"],"B":"E8F5E9","C":"FFF3E0","D":"EDE7F6","E":"F3E5F5"}
    for i, row in enumerate(deliverables):
        goal = row[0]
        bg   = goal_bgs.get(goal, C["row_a"])
        for j, val in enumerate(row):
            is_pres = (j == 5)
            pres_color = C["ok_green"] if val == "O" else C["na_gray"]
            cw(t_del.cell(i+1,j), val, size=8, bg=bg,
               center=(j in (0,3,5)),
               color=pres_color if is_pres else None,
               bold=(is_pres and val=="O"))
    set_col_widths(t_del, [Cm(1.3), Cm(4.0), Cm(4.8), Cm(1.3), Cm(6.0), Cm(1.5)])
    add_spacer(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 6. 변경 이력
    # ══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_h1(doc, "6.  변경 이력")

    t_hist = doc.add_table(rows=3, cols=5)
    t_hist.style = "Table Grid"
    hdr_row(t_hist, ["버전","작성일","작성자","변경 내용","비고"], bg=C["hdr_dark"])
    history = [
        ("v1.0","2026-04-16","권태향",
         "최초 작성 — 분석 항목 A~E 전체 입력·방법·지표·기준·시각화 정의","-"),
        ("v1.1","2026-04-16","권태향",
         "GICOMS 해양사고 통계 재구성 데이터 확보 반영 / 참조 데이터 URL 추가 / 선원법 시행규칙 조항 오류 정정 (별표 5의5 → 제52조)","GICOMS 확보"),
    ]
    for i, row in enumerate(history):
        bg = C["row_a"] if i%2==0 else C["row_b"]
        for j, val in enumerate(row):
            cw(t_hist.cell(i+1,j), val, size=9, bg=bg, center=(j in (0,1,2,4)))
    set_col_widths(t_hist, [Cm(1.5), Cm(2.8), Cm(2.5), Cm(9.5), Cm(2.0)])

    add_spacer(doc, 10)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f  = p_foot.add_run(
        "분석목표서  →  컬럼 사전  →  데이터정의서  →  데이터분석정의서  →  실제 분석 실행")
    set_font(run_f, size=10, bold=True, color=C["title_blue"])

    doc.save(OUT_PATH)
    print(f"  저장: {OUT_PATH}")


if __name__ == "__main__":
    print("[MDTS] 데이터분석정의서 v1.1 생성 중...")
    build()
    print("[완료]")
