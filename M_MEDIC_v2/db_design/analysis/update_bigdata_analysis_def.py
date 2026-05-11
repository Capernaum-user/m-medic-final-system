"""
빅데이터 분석 정의서 업데이트 (Kaggle 추가 데이터 반영)
프로젝트: M-MEDIC v2
업데이트 일자: 2026-04-20
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path("db_design/analysis/05_빅데이터분석정의서.docx")

# ─── 페이지 및 컬럼 너비 ────────────────────────────────────────────────────
CONTENT_W = 9639
COL1 = 1843
COL2 = CONTENT_W - COL1
TOTAL = COL1 + COL2

# ─── 배경색 ──────────────────────────────────────────────────────────────────
BG_SECTION = "D9D9D9"
BG_LABEL   = "F2F2F2"
BG_WHITE   = "FFFFFF"
BG_INNER_H = "E8E8E8"

# ─── 폰트 ─────────────────────────────────────────────────────────────────────
KR_FONT = "맑은 고딕"
EN_FONT = "맑은 고딕"

def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_width(cell, dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def cell_pad(cell, top=80, bottom=80, left=113, right=113):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:cantSplit")):
        trPr.remove(old)
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)

def row_header(row):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:tblHeader")):
        trPr.remove(old)
    th = OxmlElement("w:tblHeader")
    trPr.append(th)

def para_line_bottom(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single")
    btm.set(qn("w:sz"), "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "000000")
    pBdr.append(btm)
    pPr.append(pBdr)

def tbl_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def tbl_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

def tbl_width_dxa(table, dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(dxa))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)

def tbl_grid(table, widths_dxa):
    tbl = table._tbl
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

def set_run_font(run, size=10, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = EN_FONT
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    rFonts.set(qn("w:eastAsia"), KR_FONT)
    rFonts.set(qn("w:cs"), EN_FONT)

def R(para, text, bold=False, size=10):
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run

def row_funcname(table, name):
    row = table.add_row()
    row_cant_split(row)
    row_header(row)
    lc, rc = row.cells[0], row.cells[1]
    shade(lc, BG_SECTION); shade(rc, BG_WHITE)
    cell_width(lc, COL1);  cell_width(rc, COL2)
    cell_pad(lc, top=100, bottom=100, left=80, right=80)
    cell_pad(rc, top=90, bottom=90, left=130, right=113)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(lp, "기능명", bold=True, size=10)
    rp = rc.paragraphs[0]
    R(rp, name, bold=False, size=10)

def row_section(table, text):
    row = table.add_row()
    row_cant_split(row)
    cell = row.cells[0].merge(row.cells[1])
    shade(cell, BG_SECTION)
    cell_width(cell, TOTAL)
    cell_pad(cell, top=90, bottom=90)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, text, bold=True, size=10)

def row_data(table, label, split=False):
    row = table.add_row()
    if not split:
        row_cant_split(row)
    lc, rc = row.cells[0], row.cells[1]
    shade(lc, BG_LABEL); shade(rc, BG_WHITE)
    cell_width(lc, COL1); cell_width(rc, COL2)
    cell_pad(lc, top=110, bottom=110, left=80, right=80)
    cell_pad(rc, top=90, bottom=90, left=130, right=113)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(lp, label, bold=True, size=10)
    return lc, rc

def fill(rc, lines):
    first = True
    for item in lines:
        if isinstance(item, str):
            text, bold, ind = item, False, 0.0
        else:
            text = item[0]
            bold = item[1] if len(item) > 1 else False
            ind  = item[2] if len(item) > 2 else 0.0
        if first:
            p = rc.paragraphs[0]; first = False
        else:
            p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if ind:
            p.paragraph_format.left_indent = Cm(ind)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        R(p, text, bold=bold, size=10)

def append_to(rc, lines):
    for item in lines:
        if isinstance(item, str):
            text, bold, ind = item, False, 0.0
        else:
            text = item[0]
            bold = item[1] if len(item) > 1 else False
            ind  = item[2] if len(item) > 2 else 0.0
        p = rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if ind:
            p.paragraph_format.left_indent = Cm(ind)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        R(p, text, bold=bold, size=10)

def inner_table(rc, headers, rows_data, col_widths=None):
    inner_total = COL2 - 260
    if col_widths is None:
        each = inner_total // len(headers)
        col_widths = [each] * (len(headers) - 1) + [inner_total - each * (len(headers) - 1)]
    else:
        s = sum(col_widths)
        col_widths = [int(w * inner_total / s) for w in col_widths]
        diff = inner_total - sum(col_widths)
        col_widths[-1] += diff

    tbl = rc.add_table(rows=1, cols=len(headers))
    tbl_borders(tbl)
    tbl_layout_fixed(tbl)
    tbl_width_dxa(tbl, inner_total)
    tbl_grid(tbl, col_widths)

    hrow = tbl.rows[0]
    row_cant_split(hrow)
    for cell, h, cw in zip(hrow.cells, headers, col_widths):
        shade(cell, BG_INNER_H)
        cell_width(cell, cw)
        cell_pad(cell, top=50, bottom=50, left=60, right=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, h, bold=True, size=9)
    for rd in rows_data:
        row = tbl.add_row()
        row_cant_split(row)
        for cell, val, cw in zip(row.cells, rd, col_widths):
            cell_width(cell, cw)
            cell_pad(cell, top=40, bottom=40, left=60, right=60)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            R(p, val, size=9)
    return tbl

def setup_document(doc):
    section = doc.sections[0]
    section.page_width        = Cm(21.0)
    section.page_height       = Cm(29.7)
    section.top_margin        = Cm(2.0)
    section.bottom_margin     = Cm(2.0)
    section.left_margin       = Cm(2.0)
    section.right_margin      = Cm(2.0)
    section.different_first_page_header_footer = True

    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    rFonts.set(qn("w:eastAsia"), KR_FONT)
    rFonts.set(qn("w:cs"), EN_FONT)

    header = section.header
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    R(hp, "MDTS 프로젝트  빅데이터 분석 정의서", size=9)
    para_line_bottom(hp)

def build_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    R(p, "MDTS 프로젝트 성과발표회", size=9)
    para_line_bottom(p)
    for _ in range(6): doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    tbl_borders(t); tbl_layout_fixed(t); tbl_width_dxa(t, CONTENT_W); tbl_grid(t, [CONTENT_W])
    tc = t.rows[0].cells[0]; row_cant_split(t.rows[0]); shade(tc, BG_WHITE); cell_pad(tc, top=400, bottom=400)
    p = tc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("빅데이터  분석  정의서"); set_run_font(r, size=30, bold=True)
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(2.5)
    r1 = p.add_run("과제명 :  "); set_run_font(r1, size=13, bold=True)
    r2 = p.add_run('MobileNetV3 기반 선박 외상 분류 시스템,  "MDTS"'); set_run_font(r2, size=13, bold=True)
    for _ in range(6): doc.add_paragraph()
    for text in ("2026.04.20.", "팀명:  MDTS (빅데이터분석팀)"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); set_run_font(r, size=12, bold=True)
    doc.add_page_break()

def build_overview(doc):
    p = doc.add_paragraph(); R(p, "Ⅰ.  개요", bold=True, size=13)
    def item(num, title, detail=None):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); R(p, f"{num}. {title}", size=11)
        if detail:
            p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(1.5); R(p2, f": {detail}", size=11)
    item("1", "아이디어 주제", "MobileNetV3을 활용한 선박 외상 이미지 자동 분류 및 응급처치 가이드 제공 시스템")
    item("2", "개발 목표", "선박 응급처치 지원을 위한 외상 분류 AI 모델 개발 및 처치 가이드 서비스 구현")
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); R(p, "3. 개발 내용", size=11)
    for detail in (": MobileNetV3-Small을 활용한 6종 외상 이미지 분류 모델 개발", ": 외상 분류 결과 기반 선박 내 응급처치 가이드 자동 제공", ": 해양사고 통계 기반 위험도 예측 모델 구현"):
        p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Cm(1.5); R(p2, detail, size=11)
    doc.add_paragraph()

def build_section2_heading(doc):
    p = doc.add_paragraph(); R(p, "Ⅱ.  기능별 빅데이터 분석 명세서", bold=True, size=13); doc.add_paragraph()

def create_main_table(doc):
    table = doc.add_table(rows=0, cols=2)
    tbl_borders(table); tbl_layout_fixed(table); tbl_width_dxa(table, TOTAL); tbl_grid(table, [COL1, COL2])
    return table

def build_func1(doc):
    table = create_main_table(doc)
    row_funcname(table, "외상 이미지 분류 및 정밀 진단 (CNN/Segmentation)")
    row_section(table, "1. 데이터 준비")
    _, rc = row_data(table, "데이터 정의", split=True)
    fill(rc, [
        "선박 내 외상 및 화상 이미지 데이터셋 (추가 확보 완료)",
        ("- Abrasions/Bruises/Burns/Cut/Laceration 등 외상 5종", False, 0.3),
        ("- Burns 정밀 진단용 추가 데이터셋 (YOLO 어노테이션 포함)", False, 0.3),
        ("- Wound Segmentation 데이터셋 (2,760세트)", False, 0.3),
    ])
    _, rc = row_data(table, "데이터\n획득 방법", split=True)
    fill(rc, [
        "1. Kaggle - ibrahimfateen/wound-classification (외상 분류)",
        "2. Kaggle - shubhambaid/skin-burn-dataset (화상 특화)",
        "3. Kaggle - leoscode/wound-segmentation-images (세그멘테이션)",
        "※ 당뇨병성 상처 및 만성 질환 데이터는 선박 환경 무관하여 폐기 처리함",
    ])
    _, rc = row_data(table, "수집 데이터", split=True)
    inner_table(rc,
        ["데이터셋 명칭", "장수/규모", "주요 용도"],
        [
            ("Wound Classification", "약 2,000장", "외상 5종 분류 학습"),
            ("Skin Burn Extra", "600장+", "화상 심도 정밀 분석"),
            ("Wound Segmentation", "2,760세트", "상처 범위(%) 측정"),
        ], col_widths=[4, 3, 3]
    )
    row_section(table, "2. 전처리")
    _, rc = row_data(table, "전처리 과정", split=True)
    fill(rc, ["1) 이미지 리사이징 (224x224)", "2) Data Augmentation (Flip, Rotation, Jitter)", "3) Segmentation Mask 정규화", "4) 화상 부위 객체 탐지용 좌표 변환"])
    row_section(table, "3. 데이터 분석")
    _, rc = row_data(table, "데이터\n분석 목표"); fill(rc, ["외상 5종 자동 분류 및 상처 면적 측정을 통한 중증도 판정"])
    _, rc = row_data(table, "데이터\n분석\n알고리즘"); fill(rc, ["MobileNetV3-Small (Classification), UNet (Segmentation), YOLOv8 (Burn Detection)"])
    row_section(table, "4. 모델 생성 및 학습")
    _, rc = row_data(table, "학습 모델"); fill(rc, ["MobileNetV3-Small, YOLOv8-Nano"])
    row_section(table, "5. 검증")
    _, rc = row_data(table, "모델링\n검증 방안"); fill(rc, ["Accuracy, mAP(Mean Average Precision), Dice Coefficient (Segmentation)"])

def build_func2(doc):
    doc.add_paragraph()
    table = create_main_table(doc)
    row_funcname(table, "해양사고 위험도 및 인명피해 예측 (ML)")
    row_section(table, "1. 데이터 준비")
    _, rc = row_data(table, "데이터 정의", split=True)
    fill(rc, ["해상 환경 및 사고 유형별 인명피해 통계 데이터", "- 사고유형, 선박종류, 사고원인, 발생해역, 사상자수 등"])
    _, rc = row_data(table, "데이터\n획득 방법", split=True)
    fill(rc, ["1. GICOMS 해양사고 정보시스템 재구성 데이터", "2. Kaggle - UK MAIB Maritime Accidents (2020-2024)", "3. 합성 해양사고 시뮬레이션 데이터 (1,500건)"])
    _, rc = row_data(table, "수집 데이터", split=True)
    inner_table(rc, ["출처", "건수", "특징"], [("GICOMS", "48건", "국내 사고 특성"), ("UK MAIB", "11,000건", "국제 표준 사고 데이터"), ("Simulated", "1,500건", "데이터 불균형 해소용")], col_widths=[3, 3, 4])
    row_section(table, "2. 전처리")
    _, rc = row_data(table, "전처리 과정"); fill(rc, ["1) 범주형 변수 원-핫 인코딩", "2) 사상자수 기반 위험도(H/M/L) 라벨링", "3) SMOTE를 이용한 클래스 불균형 완화"])
    row_section(table, "3. 데이터 분석")
    _, rc = row_data(table, "데이터\n분석 목표"); fill(rc, ["해역 및 환경 변수에 따른 사고 위험 등급 실시간 예측"])
    _, rc = row_data(table, "데이터\n분석\n알고리즘"); fill(rc, ["Random Forest Classifier, XGBoost"])
    row_section(table, "4. 모델 생성 및 학습")
    _, rc = row_data(table, "학습 모델"); fill(rc, ["Random Forest (주 모델)"])
    row_section(table, "5. 검증")
    _, rc = row_data(table, "모델링\n검증 방안"); fill(rc, ["F1-Score, AUC-ROC, Feature Importance 분석"])

def main():
    doc = Document()
    setup_document(doc)
    build_cover(doc)
    build_overview(doc)
    build_section2_heading(doc)
    build_func1(doc)
    build_func2(doc)
    doc.save(str(OUT))
    print(f"업데이트 완료: {OUT}")

if __name__ == "__main__":
    main()
