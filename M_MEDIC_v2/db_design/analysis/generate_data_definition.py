"""
MDTS 데이터정의서 Word 문서 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "03_데이터정의서.docx")

# ── 공통 헬퍼 ────────────────────────────────────────────────────────────────

def set_font(run, size=10, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_write(cell, text, bold=False, size=9, bg=None, center=False,
               color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color)
    run.font.italic = italic
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        shade_cell(cell, bg)

def add_heading(doc, text, level=1):
    STYLES = {
        1: (18, (31,  73, 125), 16, 6),
        2: (14, (68, 114, 196), 12, 4),
        3: (11, ( 0,   0,   0),  8, 2),
    }
    fs, color, sb, sa = STYLES[level]
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=fs, bold=True, color=color)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p

def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def add_bullet(doc, text, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run("• " + text)
    set_font(run, size=10)

def add_divider(doc, light=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run("─" * 80)
    set_font(run, size=7, color=(200,200,200) if light else (150,150,150))

def set_col_widths(table, widths):
    for row in table.rows:
        for j, w in enumerate(widths):
            if j < len(row.cells):
                row.cells[j].width = w

def add_meta_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        cell_write(t.cell(i,0), k, bold=True, size=9, bg="BDD7EE")
        cell_write(t.cell(i,1), v, size=9)
    set_col_widths(t, [Cm(3.5), Cm(13)])
    doc.add_paragraph()

def add_page_break(doc):
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

# ── 섹션: 테이블 상세 정의 ───────────────────────────────────────────────────

def table_header_row(t, headers, widths, hdr_color, txt_color=(255,255,255)):
    for j, h in enumerate(headers):
        cell_write(t.cell(0,j), h, bold=True, size=9,
                   bg=hdr_color, center=True,
                   color=txt_color if txt_color != (255,255,255) else None)
        if txt_color == (255,255,255):
            t.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    set_col_widths(t, widths)

def build_table_info(doc, tbl_name, tbl_comment, engine="InnoDB",
                     charset="utf8mb4", collation="utf8mb4_general_ci"):
    """테이블 기본 정보 박스"""
    t = doc.add_table(rows=2, cols=4)
    t.style = "Table Grid"
    labels = [("테이블명(물리)", tbl_name),
              ("테이블명(논리)", tbl_comment),
              ("스토리지 엔진", engine),
              ("문자셋 / Collation", f"{charset} / {collation}")]
    for i, (lbl, val) in enumerate(labels):
        r, c = divmod(i, 2)
        # 각 행은 레이블 + 값 쌍으로
        pass

    # 단순 2행 4열: [물리명 레이블][물리명값][논리명 레이블][논리명값]
    t2 = doc.add_table(rows=2, cols=4)
    t2.style = "Table Grid"
    row0 = [("테이블명 (물리)", tbl_name),    ("테이블명 (논리)", tbl_comment)]
    row1 = [("스토리지 엔진",   engine),       ("문자셋 / Collation", f"{charset} / {collation}")]
    for ri, row_data in enumerate([row0, row1]):
        for ci, (lbl, val) in enumerate(row_data):
            cell_write(t2.cell(ri, ci*2),   lbl, bold=True, size=9, bg="D9E1F2")
            cell_write(t2.cell(ri, ci*2+1), val, size=9)
    set_col_widths(t2, [Cm(3.8), Cm(4.7), Cm(3.8), Cm(4.7)])
    doc.add_paragraph()
    # 첫번째 빈 table 삭제
    t._element.getparent().remove(t._element)

def build_column_table(doc, col_rows, hdr_color):
    """컬럼 정의 표"""
    headers = ["No", "컬럼명\n(물리)", "컬럼명\n(논리)", "데이터\n타입", "길이/\n정밀도",
               "NULL\n허용", "PK", "FK", "기본값", "도메인 규칙 / 설명"]
    widths  = [Cm(0.7), Cm(2.8), Cm(2.5), Cm(2.4), Cm(1.8),
               Cm(1.3), Cm(0.9), Cm(0.9), Cm(2.2), Cm(6.5)]

    t = doc.add_table(rows=1+len(col_rows), cols=len(headers))
    t.style = "Table Grid"
    table_header_row(t, headers, widths, hdr_color)

    for i, row in enumerate(col_rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        no, phys, logic, dtype, length, nullable, pk, fk, default, domain = row

        # PK 행은 연한 노란색
        if pk == "PK":
            bg = "FFFACD"

        cell_write(t.cell(i+1, 0),  str(no),     size=8, bg=bg, center=True)
        cell_write(t.cell(i+1, 1),  phys,        size=8, bg=bg,
                   bold=(pk=="PK"), color=(31,73,125) if pk=="PK" else None)
        cell_write(t.cell(i+1, 2),  logic,       size=8, bg=bg)
        cell_write(t.cell(i+1, 3),  dtype,       size=8, bg=bg, center=True)
        cell_write(t.cell(i+1, 4),  length,      size=8, bg=bg, center=True)
        # NULL 허용: N → 빨강
        null_color = (192,0,0) if nullable == "N" else None
        cell_write(t.cell(i+1, 5),  nullable,    size=8, bg=bg, center=True,
                   color=null_color, bold=(nullable=="N"))
        # PK/FK 강조
        pk_color = (31,73,125) if pk=="PK" else None
        fk_color = (68,114,196) if fk=="FK" else None
        cell_write(t.cell(i+1, 6),  pk,          size=8, bg=bg, center=True,
                   color=pk_color, bold=(pk=="PK"))
        cell_write(t.cell(i+1, 7),  fk,          size=8, bg=bg, center=True,
                   color=fk_color, bold=(fk=="FK"))
        cell_write(t.cell(i+1, 8),  default,     size=8, bg=bg, center=True)
        cell_write(t.cell(i+1, 9),  domain,      size=8, bg=bg)

    set_col_widths(t, widths)
    doc.add_paragraph()

def build_index_table(doc, idx_rows, hdr_color):
    """인덱스 정의 표"""
    headers = ["인덱스명", "유형", "대상 컬럼", "정렬", "설명"]
    widths  = [Cm(4.5), Cm(2.0), Cm(4.5), Cm(1.5), Cm(5.5)]
    t = doc.add_table(rows=1+len(idx_rows), cols=5)
    t.style = "Table Grid"
    table_header_row(t, headers, widths, hdr_color)
    for i, row in enumerate(idx_rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell_write(t.cell(i+1, j), val, size=8, bg=bg,
                       center=(j in (1,3)))
    set_col_widths(t, widths)
    doc.add_paragraph()

def build_constraint_table(doc, con_rows, hdr_color):
    """제약조건 정의 표"""
    headers = ["제약조건명", "유형", "컬럼", "참조 테이블.컬럼", "ON DELETE", "ON UPDATE"]
    widths  = [Cm(5.5), Cm(1.8), Cm(2.5), Cm(4.5), Cm(2.2), Cm(1.5)]
    t = doc.add_table(rows=1+len(con_rows), cols=6)
    t.style = "Table Grid"
    table_header_row(t, headers, widths, hdr_color)
    for i, row in enumerate(con_rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell_write(t.cell(i+1, j), val, size=8, bg=bg,
                       center=(j in (1,4,5)))
    set_col_widths(t, widths)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 메인 문서 빌드
# ══════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # 페이지 여백
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.2)

    # ── 표지 ─────────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("데  이  터  정  의  서")
    set_font(run, size=26, bold=True, color=(31, 73, 125))
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("MDTS — 선박 탑재형 엣지 AI 의료 지원 시스템")
    set_font(run2, size=13, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Database : mdts_db  |  DBMS : MySQL 8.x")
    set_font(run3, size=11, color=(120, 120, 120))

    for _ in range(6):
        doc.add_paragraph()

    cover_meta = [
        ("프로젝트명", "MDTS (선박 탑재형 엣지 AI 의료 지원 시스템)"),
        ("담당 파트",  "빅데이터 분석"),
        ("작성일",     "2026-04-16"),
        ("작성자",     "권태향"),
        ("문서 버전",  "v1.0"),
        ("참조 문서",  "01_분석목표서.docx, 02_분석용_컬럼_사전.docx, DB멘토링(실전).SQL"),
        ("변경 이력",  "v1.0 최초 작성 (2026-04-16)"),
    ]
    add_meta_table(doc, cover_meta)

    add_page_break(doc)

    # ── 1. 문서 개요 ──────────────────────────────────────────────────────────
    add_heading(doc, "1. 문서 개요", level=1)
    add_divider(doc)

    add_heading(doc, "1.1 목적", level=2)
    add_body(doc,
        "본 문서는 MDTS(선박 탑재형 엣지 AI 의료 지원 시스템)의 MySQL 데이터베이스를 구성하는 "
        "모든 테이블, 컬럼, 인덱스, 제약조건을 공식적으로 정의한다. "
        "개발팀·분석팀·운영팀이 동일한 기준으로 DB를 이해하고 활용하기 위한 단일 기준 문서(Single Source of Truth)로 사용된다.")

    add_heading(doc, "1.2 적용 범위", level=2)
    for line in [
        "데이터베이스명: mdts_db",
        "DBMS: MySQL 8.x / MariaDB 10.6+",
        "적용 테이블: tb_crew, tb_vital, tb_analysis, tb_logs, tb_firstaid (총 5개)",
        "문자셋: utf8mb4 / Collation: utf8mb4_general_ci",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "1.3 용어 정의", level=2)
    terms = [
        ("PK",           "Primary Key. 테이블 내 행을 유일하게 식별하는 기본키"),
        ("FK",           "Foreign Key. 다른 테이블의 PK를 참조하는 외래키"),
        ("AUTO_INCREMENT","삽입 시 자동으로 1씩 증가하는 정수 시퀀스"),
        ("NOT NULL",     "NULL 값을 허용하지 않음 — 반드시 값이 입력되어야 함"),
        ("CURRENT_TIMESTAMP","레코드 삽입 시점의 서버 시각을 자동 기록"),
        ("ENUM",         "미리 정의된 문자열 집합 중 하나만 허용하는 타입"),
        ("DECIMAL(p,s)", "전체 p자리, 소수점 이하 s자리의 고정 소수점 실수"),
        ("BOOLEAN",      "TRUE(1) / FALSE(0) 논리값. MySQL에서 TINYINT(1)로 처리"),
        ("ON DELETE RESTRICT","부모 레코드 삭제 시 자식 레코드가 있으면 삭제 차단"),
        ("ON UPDATE RESTRICT","부모 PK 변경 시 자식 레코드가 있으면 변경 차단"),
    ]
    t_term = doc.add_table(rows=1+len(terms), cols=2)
    t_term.style = "Table Grid"
    for j, h in enumerate(["용어", "정의"]):
        cell_write(t_term.cell(0,j), h, bold=True, size=9, bg="404040", center=True)
        t_term.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for i, (term, definition) in enumerate(terms):
        bg = "F2F2F2" if i%2==0 else "FFFFFF"
        cell_write(t_term.cell(i+1,0), term,       size=9, bold=True, bg=bg)
        cell_write(t_term.cell(i+1,1), definition, size=9, bg=bg)
    set_col_widths(t_term, [Cm(4.5), Cm(13)])
    doc.add_paragraph()

    add_page_break(doc)

    # ── 2. 데이터베이스 구조 개요 ────────────────────────────────────────────
    add_heading(doc, "2. 데이터베이스 구조 개요", level=1)
    add_divider(doc)

    add_heading(doc, "2.1 테이블 목록", level=2)
    tbl_list = [
        ("tb_crew",     "승조원 정보",       "선박에 승선한 승조원의 인적 정보를 관리",          "tb_vital, tb_analysis, tb_firstaid, tb_logs"),
        ("tb_vital",    "생체 신호 데이터",   "승조원의 심박수·산소포화도·체온 측정 이력",          "tb_crew"),
        ("tb_analysis", "AI 분석 결과",      "AI가 이미지·생체 신호를 분석한 결과 및 위험 등급",   "tb_vital, tb_crew"),
        ("tb_firstaid", "응급 처치 기록",     "AI 가이드에 따라 실제 수행한 응급처치 내용",         "tb_analysis, tb_crew"),
        ("tb_logs",     "동기화 로그",        "육상 서버로의 데이터 전송 성공/실패 이력",           "tb_crew"),
    ]
    t_list = doc.add_table(rows=1+len(tbl_list), cols=4)
    t_list.style = "Table Grid"
    for j, h in enumerate(["테이블명", "논리명", "설명", "참조하는 테이블"]):
        cell_write(t_list.cell(0,j), h, bold=True, size=9, bg="1F497D", center=True)
        t_list.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for i, row in enumerate(tbl_list):
        bg = "D9E1F2" if i%2==0 else "EBF0FA"
        for j, val in enumerate(row):
            cell_write(t_list.cell(i+1,j), val, size=9, bg=bg)
    set_col_widths(t_list, [Cm(3.2), Cm(3.2), Cm(7.5), Cm(4.1)])
    doc.add_paragraph()

    add_heading(doc, "2.2 테이블 간 관계 (FK 참조 구조)", level=2)
    rel_lines = [
        "tb_crew (crew_id  ← PK)",
        "   │",
        "   ├─[FK crew_id]──► tb_vital      (vital_id ← PK)",
        "   │                      │",
        "   │                      └─[FK vital_id]──► tb_analysis (analysis_id ← PK)",
        "   │                                               │",
        "   ├─[FK crew_id]──► tb_analysis                  └─[FK analysis_id]──► tb_firstaid",
        "   ├─[FK crew_id]──► tb_firstaid",
        "   └─[FK admin_id]─► tb_logs",
    ]
    p_rel = doc.add_paragraph()
    p_rel.paragraph_format.left_indent = Cm(1.2)
    for line in rel_lines:
        run = p_rel.add_run(line + "\n")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 73, 125)
    doc.add_paragraph()

    add_heading(doc, "2.3 외래키 전체 목록", level=2)
    fk_all = [
        ("FK_tb_vital_crew",         "tb_vital.crew_id",         "tb_crew.crew_id",          "RESTRICT", "RESTRICT"),
        ("FK_tb_analysis_vital",     "tb_analysis.vital_id",     "tb_vital.vital_id",         "RESTRICT", "RESTRICT"),
        ("FK_tb_firstaid_analysis",  "tb_firstaid.analysis_id",  "tb_analysis.analysis_id",   "RESTRICT", "RESTRICT"),
        ("FK_tb_firstaid_crew",      "tb_firstaid.crew_id",      "tb_crew.crew_id",           "RESTRICT", "RESTRICT"),
        ("FK_tb_logs_admin",         "tb_logs.admin_id",         "tb_crew.crew_id",           "RESTRICT", "RESTRICT"),
    ]
    t_fk = doc.add_table(rows=1+len(fk_all), cols=5)
    t_fk.style = "Table Grid"
    for j, h in enumerate(["제약조건명", "자식 컬럼", "부모 컬럼", "ON DELETE", "ON UPDATE"]):
        cell_write(t_fk.cell(0,j), h, bold=True, size=9, bg="404040", center=True)
        t_fk.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for i, row in enumerate(fk_all):
        bg = "F2F2F2" if i%2==0 else "FFFFFF"
        for j, val in enumerate(row):
            cell_write(t_fk.cell(i+1,j), val, size=9, bg=bg, center=(j in (3,4)))
    set_col_widths(t_fk, [Cm(5.2), Cm(4.0), Cm(4.0), Cm(2.2), Cm(2.2)])
    doc.add_paragraph()

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 3. 테이블 상세 정의
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. 테이블 상세 정의", level=1)
    add_divider(doc)

    # ── 3.1 tb_crew ──────────────────────────────────────────────────────────
    add_heading(doc, "3.1  tb_crew  —  승조원 정보", level=2)
    build_table_info(doc, "tb_crew", "승조원 정보")

    crew_cols = [
        # No, 물리명, 논리명, 타입, 길이/정밀도, NULL허용, PK, FK, 기본값, 도메인/설명
        (1,"crew_id",   "승조원 고유번호","INT",          "11",    "N","PK","",  "AUTO_INCREMENT","1 이상 정수. 자동 증가. 전 테이블의 조인 기준키"),
        (2,"name",      "성명",          "VARCHAR",      "50",    "N","",  "",  "-",             "한글/영문 최대 50자. 분석 시 crew_id로 대체(PII)"),
        (3,"birthdate", "생년월일",       "DATE",         "YYYY-MM-DD","N","","","-",            "만 나이 파생 변수 계산 기준. 1900-01-01 이상"),
        (4,"gender",    "성별",           "CHAR",         "1",     "N","",  "",  "-",             "허용값: 'M'(남성) / 'F'(여성). 대문자 통일 권장"),
        (5,"bloodtype", "혈액형",         "VARCHAR",      "4",     "N","",  "",  "-",             "허용값: 'A' / 'B' / 'O' / 'AB'"),
        (6,"position",  "직책",           "VARCHAR",      "50",    "N","",  "",  "-",             "예: 선장, 항해사, 기관사, 조리원, 일반선원. 자유 입력 → 코드화 권장"),
        (7,"joined_at", "가입 일시",       "DATETIME",     "-",     "N","",  "",  "CURRENT_TIMESTAMP","레코드 생성 시각 자동 기록. 수정 불가 권장"),
    ]
    build_column_table(doc, crew_cols, "1F497D")

    add_heading(doc, "인덱스", level=3)
    build_index_table(doc, [
        ("PRIMARY",      "PRIMARY KEY", "crew_id",           "ASC", "기본키 인덱스. 자동 생성"),
        ("IX_tb_crew_1", "NORMAL",      "name, joined_at",   "ASC", "성명 + 가입일시 복합 조회 성능 향상"),
    ], "1F497D")

    add_heading(doc, "제약조건", level=3)
    build_constraint_table(doc, [
        ("PRIMARY KEY (crew_id)", "PK", "crew_id", "-", "-", "-"),
    ], "1F497D")

    add_page_break(doc)

    # ── 3.2 tb_vital ─────────────────────────────────────────────────────────
    add_heading(doc, "3.2  tb_vital  —  생체 신호 데이터", level=2)
    build_table_info(doc, "tb_vital", "생체 신호 데이터")

    vital_cols = [
        (1,"vital_id",    "바이탈 고유번호","INT",        "11",    "N","PK","",        "AUTO_INCREMENT","1 이상 정수. 자동 증가"),
        (2,"crew_id",     "승조원 고유번호","INT",        "11",    "N","",  "FK",      "-",             "tb_crew.crew_id 참조. 측정 대상 승조원 식별"),
        (3,"heart_rate",  "심박수",        "INT",        "11",    "N","",  "",        "0",             "단위: bpm. 정상 60~100. 이상값: <30 or >250 시 검증 필요"),
        (4,"spo2",        "산소포화도",     "DECIMAL",    "5,2",   "N","",  "",        "0.00",          "단위: %. 허용값 0.00~100.00. 정상: ≥95.00. 위험: <90.00"),
        (5,"temperature", "체온",           "DECIMAL",    "5,2",   "N","",  "",        "0.00",          "단위: °C. 허용값 30.00~45.00. 정상: 36.10~37.20"),
        (6,"measured_at", "측정 일시",       "DATETIME",   "-",     "N","",  "",        "CURRENT_TIMESTAMP","측정 시각 자동 기록. 분석 목표 D의 반응 시간 계산 기준"),
    ]
    build_column_table(doc, vital_cols, "375623")

    add_heading(doc, "인덱스", level=3)
    build_index_table(doc, [
        ("PRIMARY",       "PRIMARY KEY", "vital_id",    "ASC", "기본키 인덱스"),
        ("IX_tb_vital_1", "NORMAL",      "measured_at", "ASC", "측정 일시 기준 시계열 조회 성능 향상"),
    ], "375623")

    add_heading(doc, "제약조건", level=3)
    build_constraint_table(doc, [
        ("PRIMARY KEY (vital_id)",                  "PK", "vital_id", "-",                   "-",       "-"),
        ("FK_tb_vital_crew_id_tb_crew_crew_id",     "FK", "crew_id",  "tb_crew.crew_id",     "RESTRICT","RESTRICT"),
    ], "375623")

    add_page_break(doc)

    # ── 3.3 tb_analysis ──────────────────────────────────────────────────────
    add_heading(doc, "3.3  tb_analysis  —  AI 분석 결과", level=2)
    build_table_info(doc, "tb_analysis", "AI 분석 결과")

    analysis_cols = [
        (1, "analysis_id",     "분석 고유번호",  "INT",              "11",    "N","PK","",   "AUTO_INCREMENT","1 이상 정수. 자동 증가"),
        (2, "vital_id",        "바이탈 고유번호","INT",              "11",    "N","",  "FK", "-",             "tb_vital.vital_id 참조. 분석 대상 생체 신호 연결"),
        (3, "crew_id",         "승조원 고유번호","INT",              "11",    "N","",  "FK", "-",             "tb_crew.crew_id 참조. vital_id의 crew_id와 일치해야 함"),
        (4, "analysis_result", "분석 내용",      "TEXT",             "-",     "N","",  "",   "-",             "AI 생성 자유 텍스트. 비정형. 직접 집계 불가. NLP 필요 시 별도 처리"),
        (5, "diagnosis",       "진단 결과",       "VARCHAR",          "255",   "N","",  "",   "-",             "AI 진단명. 예: 'Burn_2nd_degree', 'Cut_Deep'. 표준 코드화 권장"),
        (6, "file_name",       "파일 이름",        "VARCHAR",          "255",   "N","",  "",   "-",             "형식: {원본명}_{UUID}.{ext}. 예: image_2f48f241-9d64.jpg"),
        (7, "file_size",       "파일 사이즈",      "INT",              "11",    "N","",  "",   "-",             "단위: bytes. 0 이상 정수"),
        (8, "file_ext",        "파일 확장자",      "VARCHAR",          "10",    "N","",  "",   "-",             "허용값: '.jpg' / '.jpeg' / '.png'. 점(.) 포함"),
        (9, "risk_level",      "위험 등급",        "ENUM",             "'1'~'4'","N","","",   "-",             "1=GREEN(안전) / 2=YELLOW(주의) / 3=ORANGE(경고) / 4=RED(위험). 분석 핵심 종속변수"),
        (10,"analyzed_at",     "분석 일시",        "DATETIME",         "-",     "N","",  "",   "CURRENT_TIMESTAMP","분석 완료 시각. tb_firstaid.created_at과의 차이로 반응 시간 계산"),
    ]
    build_column_table(doc, analysis_cols, "843C0C")

    add_heading(doc, "인덱스", level=3)
    build_index_table(doc, [
        ("PRIMARY",           "PRIMARY KEY", "analysis_id", "ASC", "기본키 인덱스"),
        ("IX_tb_analysis_1",  "NORMAL",      "analyzed_at", "ASC", "분석 일시 기준 시계열 조회 성능 향상"),
    ], "843C0C")

    add_heading(doc, "제약조건", level=3)
    build_constraint_table(doc, [
        ("PRIMARY KEY (analysis_id)",                          "PK", "analysis_id", "-",                     "-",       "-"),
        ("FK_tb_analysis_vital_id_tb_vital_vital_id",          "FK", "vital_id",    "tb_vital.vital_id",      "RESTRICT","RESTRICT"),
    ], "843C0C")

    add_page_break(doc)

    # ── 3.4 tb_logs ──────────────────────────────────────────────────────────
    add_heading(doc, "3.4  tb_logs  —  육상 서버 동기화 로그", level=2)
    build_table_info(doc, "tb_logs", "육상 서버 동기화 로그")

    logs_cols = [
        (1,"log_id",      "로그 고유번호",  "INT",       "11",  "N","PK","",   "AUTO_INCREMENT","1 이상 정수. 자동 증가"),
        (2,"admin_id",    "관리자 고유번호","INT",       "11",  "N","",  "FK", "-",             "tb_crew.crew_id 참조. 동기화를 수행한 관리자 역할 승조원"),
        (3,"table_name",  "대상 테이블명",  "VARCHAR",   "64",  "N","",  "",   "-",             "허용값: 'tb_vital' / 'tb_analysis' / 'tb_firstaid' / 'tb_crew'. 입력값 정규화 필요"),
        (4,"record_id",   "대상 레코드",    "VARCHAR",   "64",  "N","",  "",   "-",             "동기화 대상 레코드의 PK 값. 문자열로 저장"),
        (5,"sync_status", "전송 상태",       "BOOLEAN",   "1",   "N","",  "",   "-",             "TRUE(1)=성공 / FALSE(0)=실패. 분석 목표 E 핵심 종속변수"),
        (6,"synced_at",   "전송 일시",        "DATETIME",  "-",   "N","",  "",   "CURRENT_TIMESTAMP","전송 시각 자동 기록. 시간대별 실패율 분석 기준"),
    ]
    build_column_table(doc, logs_cols, "3F3F76")

    add_heading(doc, "인덱스", level=3)
    build_index_table(doc, [
        ("PRIMARY",       "PRIMARY KEY", "log_id",    "ASC", "기본키 인덱스"),
        ("IX_tb_logs_1",  "NORMAL",      "synced_at", "ASC", "전송 일시 기준 시계열 조회 성능 향상"),
    ], "3F3F76")

    add_heading(doc, "제약조건", level=3)
    build_constraint_table(doc, [
        ("PRIMARY KEY (log_id)",                    "PK", "log_id",   "-",               "-",       "-"),
        ("FK_tb_logs_admin_id_tb_crew_crew_id",     "FK", "admin_id", "tb_crew.crew_id", "RESTRICT","RESTRICT"),
    ], "3F3F76")

    add_page_break(doc)

    # ── 3.5 tb_firstaid ──────────────────────────────────────────────────────
    add_heading(doc, "3.5  tb_firstaid  —  응급 처치 기록", level=2)
    build_table_info(doc, "tb_firstaid", "응급 처치 기록")

    firstaid_cols = [
        (1,"firstaid_id",  "응급처치 고유번호","INT",      "11",  "N","PK","",   "AUTO_INCREMENT","1 이상 정수. 자동 증가"),
        (2,"analysis_id",  "분석 고유번호",   "INT",      "11",  "N","",  "FK", "-",             "tb_analysis.analysis_id 참조. 어떤 분석 결과에 대한 처치인지 연결"),
        (3,"crew_id",      "승조원 고유번호", "INT",      "11",  "N","",  "FK", "-",             "tb_crew.crew_id 참조. 처치 대상 승조원"),
        (4,"guide_text",   "가이드 내용",      "TEXT",     "-",   "N","",  "",   "-",             "AI가 제공한 응급처치 가이드 전문. 비정형. 참조 전용"),
        (5,"action_taken", "실제 조치 사항",   "TEXT",     "-",   "N","",  "",   "-",             "승조원이 실제 수행한 처치 내용. guide_text 준수율 분석(목표 A)에 활용"),
        (6,"created_at",   "처치 완료 일시",   "DATETIME", "-",   "N","",  "",   "CURRENT_TIMESTAMP","처치 완료 시각. tb_analysis.analyzed_at과의 차이로 반응 시간 계산(목표 D)"),
    ]
    build_column_table(doc, firstaid_cols, "7030A0")

    add_heading(doc, "인덱스", level=3)
    build_index_table(doc, [
        ("PRIMARY",   "PRIMARY KEY", "firstaid_id", "ASC", "기본키 인덱스"),
    ], "7030A0")

    add_heading(doc, "제약조건", level=3)
    build_constraint_table(doc, [
        ("PRIMARY KEY (firstaid_id)",                                  "PK", "firstaid_id", "-",                          "-",       "-"),
        ("FK_tb_firstaid_analysis_id_tb_analysis_analysis_id",         "FK", "analysis_id", "tb_analysis.analysis_id",    "RESTRICT","RESTRICT"),
        ("FK_tb_firstaid_crew_id_tb_crew_crew_id",                     "FK", "crew_id",     "tb_crew.crew_id",            "RESTRICT","RESTRICT"),
    ], "7030A0")

    add_page_break(doc)

    # ── 4. 도메인 규칙 및 데이터 품질 기준 ──────────────────────────────────
    add_heading(doc, "4. 도메인 규칙 및 데이터 품질 기준", level=1)
    add_divider(doc)

    add_heading(doc, "4.1 생체 신호 정상/이상 범위", level=2)
    vital_range = [
        ("heart_rate",  "심박수",    "bpm",  "60 ~ 100",  "< 40  또는  > 150",  "< 30  또는  > 250",  "정상 미만: 서맥 / 초과: 빈맥"),
        ("spo2",        "산소포화도","%%",    "≥ 95.00",   "90.00 ~ 94.99",      "< 85.00",            "< 90: 저산소증 의심, < 85: 즉각 처치 필요"),
        ("temperature", "체온",      "°C",   "36.10~37.20","37.50~38.49 또는 35.00~36.09","≥ 38.50 또는 ≤ 34.99","고열: ≥38.5 / 저체온: ≤35.0"),
    ]
    t_vr = doc.add_table(rows=1+len(vital_range), cols=7)
    t_vr.style = "Table Grid"
    vr_headers = ["컬럼명", "측정항목", "단위", "정상 범위", "주의 범위", "위험 범위", "비고"]
    for j, h in enumerate(vr_headers):
        cell_write(t_vr.cell(0,j), h, bold=True, size=9, bg="375623", center=True)
        t_vr.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    range_bgs = [("E2EFDA","A9D18E","FF9999")]
    for i, row in enumerate(vital_range):
        for j, val in enumerate(row):
            if j == 4:  bg = "FFF2CC"
            elif j == 5: bg = "FCE4D6"
            else:        bg = "E2EFDA" if i%2==0 else "F0F7EC"
            cell_write(t_vr.cell(i+1,j), val, size=8, bg=bg, center=(j in (2,3,4,5)))
    set_col_widths(t_vr, [Cm(2.5), Cm(2.5), Cm(1.3), Cm(2.8), Cm(3.0), Cm(2.8), Cm(3.1)])
    doc.add_paragraph()

    add_heading(doc, "4.2 위험 등급(risk_level) 정의", level=2)
    risk_rows = [
        ("1", "GREEN",  "안전",    "정상 범위 내 생체 신호. 즉각 처치 불필요",        "일반 관찰 유지"),
        ("2", "YELLOW", "주의",    "생체 신호 경미 이상 또는 경상. 관찰 강화 필요",   "30분 이내 재측정"),
        ("3", "ORANGE", "경고",    "생체 신호 중등도 이상 또는 중등상. 처치 권고",    "즉시 응급처치 시행"),
        ("4", "RED",    "위험",    "생체 신호 심각 이상 또는 중증. 긴급 처치 필수",   "무선 의료 지도 즉시 요청"),
    ]
    t_risk = doc.add_table(rows=1+len(risk_rows), cols=5)
    t_risk.style = "Table Grid"
    for j, h in enumerate(["등급값", "코드명", "한글명", "판단 기준", "권고 조치"]):
        cell_write(t_risk.cell(0,j), h, bold=True, size=9, bg="843C0C", center=True)
        t_risk.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    risk_bgs = ["E2EFDA", "FFF2CC", "FCE4D6", "FFCCCC"]
    for i, (val, code, kr, criteria, action) in enumerate(risk_rows):
        bg = risk_bgs[i]
        cell_write(t_risk.cell(i+1,0), val,      size=9, bg=bg, center=True, bold=True)
        cell_write(t_risk.cell(i+1,1), code,     size=9, bg=bg, center=True, bold=True)
        cell_write(t_risk.cell(i+1,2), kr,       size=9, bg=bg, center=True)
        cell_write(t_risk.cell(i+1,3), criteria, size=9, bg=bg)
        cell_write(t_risk.cell(i+1,4), action,   size=9, bg=bg)
    set_col_widths(t_risk, [Cm(1.6), Cm(2.2), Cm(1.8), Cm(7.5), Cm(5.0)])
    doc.add_paragraph()

    add_heading(doc, "4.3 파일명 명명 규칙", level=2)
    for line in [
        "형식: {원본파일명}_{UUID}.{확장자}",
        "예:   image0001_2f48f241-9d64-4d16-bf56-70b9d4e0e79a.jpg",
        "UUID: Java UUID.randomUUID() 또는 MySQL uuid() 함수 사용 (초당 100만 회 이상 중복 회피)",
        "허용 확장자: .jpg, .jpeg, .png",
        "file_ext 컬럼에는 점(.) 포함하여 저장 — 예: '.jpg'",
    ]:
        add_bullet(doc, line)
    doc.add_paragraph()

    add_heading(doc, "4.4 tb_crew.gender 값 정규화", level=2)
    for line in [
        "허용값: 'M' (남성), 'F' (여성) — 반드시 대문자",
        "입력 시 소문자('m', 'f') 또는 한글('남', '여') 입력 금지",
        "애플리케이션 레이어에서 UPPER() 처리 또는 CHECK 제약 추가 권장",
    ]:
        add_bullet(doc, line)
    doc.add_paragraph()

    add_page_break(doc)

    # ── 5. DDL 전체 ──────────────────────────────────────────────────────────
    add_heading(doc, "5. DDL (전체 테이블 생성 SQL)", level=1)
    add_divider(doc)
    add_body(doc, "아래 SQL은 테이블 생성 순서가 FK 의존성을 고려하여 정렬되었습니다. 순서대로 실행하면 오류 없이 생성됩니다.")
    doc.add_paragraph()

    ddl = """-- ================================================================
-- MDTS Database : mdts_db
-- 생성 순서: tb_crew → tb_vital → tb_analysis → tb_logs → tb_firstaid
-- ================================================================

CREATE TABLE tb_crew
(
    `crew_id`    INT            NOT NULL AUTO_INCREMENT COMMENT '승조원 고유번호',
    `name`       VARCHAR(50)    NOT NULL COMMENT '성명',
    `birthdate`  DATE           NOT NULL COMMENT '생년월일',
    `gender`     CHAR(1)        NOT NULL COMMENT '성별 (M/F)',
    `bloodtype`  VARCHAR(4)     NOT NULL COMMENT '혈액형 (A/B/O/AB)',
    `position`   VARCHAR(50)    NOT NULL COMMENT '직책',
    `joined_at`  DATETIME       NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '가입 일시',
    PRIMARY KEY (crew_id)
) COMMENT = '승조원 정보';

CREATE INDEX IX_tb_crew_1 ON tb_crew(name, joined_at);

-- ----------------------------------------------------------------

CREATE TABLE tb_vital
(
    `vital_id`     INT           NOT NULL AUTO_INCREMENT COMMENT '바이탈 고유번호',
    `crew_id`      INT           NOT NULL COMMENT '승조원 고유번호',
    `heart_rate`   INT           NOT NULL DEFAULT 0   COMMENT '심박수 (bpm)',
    `spo2`         DECIMAL(5,2)  NOT NULL DEFAULT 0.0 COMMENT '산소포화도 (%)',
    `temperature`  DECIMAL(5,2)  NOT NULL DEFAULT 0.0 COMMENT '체온 (°C)',
    `measured_at`  DATETIME      NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '측정 일시',
    PRIMARY KEY (vital_id),
    CONSTRAINT FK_tb_vital_crew_id_tb_crew_crew_id
        FOREIGN KEY (crew_id) REFERENCES tb_crew(crew_id)
        ON DELETE RESTRICT ON UPDATE RESTRICT
) COMMENT = '생체 신호 데이터';

CREATE INDEX IX_tb_vital_1 ON tb_vital(measured_at);

-- ----------------------------------------------------------------

CREATE TABLE tb_analysis
(
    `analysis_id`      INT                   NOT NULL AUTO_INCREMENT COMMENT '분석 고유번호',
    `vital_id`         INT                   NOT NULL COMMENT '바이탈 고유번호',
    `crew_id`          INT                   NOT NULL COMMENT '승조원 고유번호',
    `analysis_result`  TEXT                  NOT NULL COMMENT '분석 내용 (AI 생성 텍스트)',
    `diagnosis`        VARCHAR(255)          NOT NULL COMMENT '진단 결과',
    `file_name`        VARCHAR(255)          NOT NULL COMMENT '파일 이름 ({원본}_{UUID}.{ext})',
    `file_size`        INT                   NOT NULL COMMENT '파일 사이즈 (bytes)',
    `file_ext`         VARCHAR(10)           NOT NULL COMMENT '파일 확장자 (.jpg/.png)',
    `risk_level`       ENUM('1','2','3','4') NOT NULL COMMENT '위험 등급 (1=GREEN~4=RED)',
    `analyzed_at`      DATETIME              NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '분석 일시',
    PRIMARY KEY (analysis_id),
    CONSTRAINT FK_tb_analysis_vital_id_tb_vital_vital_id
        FOREIGN KEY (vital_id) REFERENCES tb_vital(vital_id)
        ON DELETE RESTRICT ON UPDATE RESTRICT
) COMMENT = 'AI 분석 결과';

CREATE INDEX IX_tb_analysis_1 ON tb_analysis(analyzed_at);

-- ----------------------------------------------------------------

CREATE TABLE tb_logs
(
    `log_id`       INT          NOT NULL AUTO_INCREMENT COMMENT '로그 고유번호',
    `admin_id`     INT          NOT NULL COMMENT '관리자 고유번호 (tb_crew.crew_id)',
    `table_name`   VARCHAR(64)  NOT NULL COMMENT '대상 테이블명',
    `record_id`    VARCHAR(64)  NOT NULL COMMENT '대상 레코드 PK',
    `sync_status`  BOOLEAN      NOT NULL COMMENT '전송 상태 (TRUE=성공/FALSE=실패)',
    `synced_at`    DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '전송 일시',
    PRIMARY KEY (log_id),
    CONSTRAINT FK_tb_logs_admin_id_tb_crew_crew_id
        FOREIGN KEY (admin_id) REFERENCES tb_crew(crew_id)
        ON DELETE RESTRICT ON UPDATE RESTRICT
) COMMENT = '육상 서버 동기화 로그';

CREATE INDEX IX_tb_logs_1 ON tb_logs(synced_at);

-- ----------------------------------------------------------------

CREATE TABLE tb_firstaid
(
    `firstaid_id`   INT      NOT NULL AUTO_INCREMENT COMMENT '응급처치 고유번호',
    `analysis_id`   INT      NOT NULL COMMENT '분석 고유번호',
    `crew_id`       INT      NOT NULL COMMENT '승조원 고유번호',
    `guide_text`    TEXT     NOT NULL COMMENT 'AI 응급처치 가이드 내용',
    `action_taken`  TEXT     NOT NULL COMMENT '실제 조치 사항',
    `created_at`    DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP COMMENT '처치 완료 일시',
    PRIMARY KEY (firstaid_id),
    CONSTRAINT FK_tb_firstaid_analysis_id_tb_analysis_analysis_id
        FOREIGN KEY (analysis_id) REFERENCES tb_analysis(analysis_id)
        ON DELETE RESTRICT ON UPDATE RESTRICT,
    CONSTRAINT FK_tb_firstaid_crew_id_tb_crew_crew_id
        FOREIGN KEY (crew_id) REFERENCES tb_crew(crew_id)
        ON DELETE RESTRICT ON UPDATE RESTRICT
) COMMENT = '응급 처치 기록';"""

    p_ddl = doc.add_paragraph()
    p_ddl.paragraph_format.left_indent = Cm(0.5)
    run = p_ddl.add_run(ddl)
    run.font.name = "Courier New"
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(40, 40, 40)
    doc.add_paragraph()

    # ── 6. 변경 이력 ─────────────────────────────────────────────────────────
    add_page_break(doc)
    add_heading(doc, "6. 변경 이력", level=1)
    add_divider(doc)

    t_hist = doc.add_table(rows=2, cols=5)
    t_hist.style = "Table Grid"
    for j, h in enumerate(["버전", "작성일", "작성자", "변경 내용", "비고"]):
        cell_write(t_hist.cell(0,j), h, bold=True, size=9, bg="404040", center=True)
        t_hist.cell(0,j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for j, val in enumerate(["v1.0", "2026-04-16", "권태향", "최초 작성 — 5개 테이블 전체 정의 (분석목표서·컬럼사전 기반)", "-"]):
        cell_write(t_hist.cell(1,j), val, size=9, bg="F2F2F2", center=(j in (0,1,2,4)))
    set_col_widths(t_hist, [Cm(1.5), Cm(2.8), Cm(2.5), Cm(9.5), Cm(2.0)])

    doc.add_paragraph()
    add_divider(doc, light=True)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run(
        "다음 단계:  분석목표서 (완료)  →  분석용 컬럼 사전 (완료)  →  데이터정의서 (완료)  →  실제 분석 실행")
    set_font(run_foot, size=10, bold=True, color=(31, 73, 125))
    p_foot.paragraph_format.space_before = Pt(6)

    doc.save(OUT_PATH)
    print(f"  저장: {OUT_PATH}")


if __name__ == "__main__":
    print("[MDTS] 데이터정의서 생성 중...")
    build()
    print("[완료]")
