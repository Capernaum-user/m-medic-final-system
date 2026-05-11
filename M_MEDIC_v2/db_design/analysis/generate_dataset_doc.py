import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_dataset_doc():
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # 제목
    title = doc.add_heading('M-Medic v2 추가 데이터셋 정의서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 개요
    doc.add_heading('1. 개요', level=1)
    doc.add_paragraph(
        '본 문서는 M-Medic v2 프로젝트의 빅데이터 분석 및 모델 학습을 위해 Kaggle에서 추가로 확보한 '
        '데이터셋의 출처와 구성을 정의한다. 선박 및 해상 환경의 특수성을 고려하여 일반적인 피부 질병 데이터를 배제하고, '
        '외상(Trauma) 및 해양 사고 통계 데이터를 중심으로 구성하였다.'
    )

    # 데이터셋 1: 외상 분류 데이터
    doc.add_heading('2. 외상 분류 데이터셋 (Wound Classification)', level=1)
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '항목'
    hdr_cells[1].text = '내용'

    data1 = [
        ('출처 (Kaggle)', 'ibrahimfateen/wound-classification'),
        ('데이터 유형', '이미지 (.jpg)'),
        ('핵심 클래스', 'Abrasions(찰과상), Bruises(타박상), Burns(화상), Cut(절상), Laceration(열상)'),
        ('정리 내용', '당뇨병성 상처, 욕창 등 질병성 데이터는 프로젝트 원칙에 따라 폐기 완료'),
        ('활용 방안', '선박 내 부상 부위 자동 분류 및 응급 처치 가이드 매칭')
    ]
    for item, val in data1:
        row_cells = table1.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val

    # 데이터셋 2: 화상 특화 데이터
    doc.add_heading('3. 화상 정밀 진단 데이터셋 (Skin Burn)', level=1)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '항목'
    hdr_cells[1].text = '내용'

    data2 = [
        ('출처 (Kaggle)', 'shubhambaid/skin-burn-dataset'),
        ('데이터 구성', '화상 이미지 및 YOLO 기반 객체 검출용 텍스트(.txt) 어노테이션'),
        ('활용 방안', '해상 사고 시 빈번한 화상(Burn)의 심도 및 범위 정밀 탐지')
    ]
    for item, val in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val

    # 데이터셋 3: 상처 세그멘테이션 데이터
    doc.add_heading('4. 상처 범위 측정 데이터셋 (Wound Segmentation)', level=1)
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '항목'
    hdr_cells[1].text = '내용'

    data3 = [
        ('출처 (Kaggle)', 'leoscode/wound-segmentation-images'),
        ('샘플 규모', '2,760개 이미지 및 세그멘테이션 마스크(Mask)'),
        ('활용 방안', '상처 부위의 픽셀 단위 면적 측정 -> 화상 면적(%) 계산을 통한 위급 상황 판단')
    ]
    for item, val in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val

    # 데이터셋 4: 해양 사고 통계 데이터
    doc.add_heading('5. 해양 사고 분석 데이터셋 (Maritime Accidents)', level=1)
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = '항목'
    hdr_cells[1].text = '내용'

    data4 = [
        ('출처 (Kaggle)', 'jeleeladekunlefijabi/uk-maib-maritime-accidents-20202024'),
        ('데이터 유형', '정형 데이터 (영국 MAIB 해양사고 통계)'),
        ('포함 기간', '2020년 ~ 2024년'),
        ('활용 방안', '해상 환경 변수(날씨, 사고 유형)와 부상 발생 확률 간의 상관관계 분석 모델링')
    ]
    for item, val in data4:
        row_cells = table4.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val

    # 저장 경로
    output_path = 'db_design/analysis/06_추가_데이터셋_정의서.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == "__main__":
    create_dataset_doc()
