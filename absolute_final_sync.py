# -*- coding: utf-8 -*-
from docx import Document
import os

def the_absolute_final_sync():
    file_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx"
    doc = Document(file_path)

    # 1. 실측 기반 6종 이미지 세부 개수 (총합 4,212)
    class_counts = {
        "Abrasions   (찰과상) :": "Abrasions   (찰과상) : 668개",
        "Bruises     (타박상) :": "Bruises     (타박상) : 972개",
        "Burns                      (화상)   :": "Burns                      (화상)   : 504개",
        "Cut         (절상)   :": "Cut         (절상)   : 600개",
        "Laceration  (열상)   :": "Laceration  (열상)   : 732개",
        "Stab_wound  (자창)   :": "Stab_wound  (자창)   : 736개"
    }

    # 2. 핵심 유닛 수치 업데이트 맵
    replace_map = {
        "2,326개": "4,212개",
        "9,350개": "11,236개",
        "1,977": "3,580", # 4,212 * 0.85
        "349": "632",   # 4,212 * 0.15
        "전체 학습 데이터 유닛 | 9,350": "전체 학습 데이터 유닛 | 11,236",
    }

    # 모든 문단 및 표 스캔
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # 클래스별 실측치 적용
                    for key, val in class_counts.items():
                        if key in p.text:
                            p.text = val
                    
                    # 총계 및 학습 유닛 적용
                    for old, new in replace_map.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)
                    
                    # '합계 : 4,212'를 넘어가는 잘못된 합계 텍스트 정정
                    if "합계" in p.text and "2,326" in p.text:
                        p.text = p.text.replace("2,326", "4,212")

    for p in doc.paragraphs:
        for old, new in replace_map.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    doc.save(file_path)
    print("SUCCESS: 3-3 Big Data Definition finalized with 11,236 Units.")

if __name__ == '__main__':
    the_absolute_final_sync()
