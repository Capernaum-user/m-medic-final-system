# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def update_to_9350_units():
    file_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx"
    doc = Document(file_path)

    def set_font(run, font_name='Malgun Gothic', size=10, bold=False, color=None):
        run.font.name = font_name
        if bold: run.bold = True
        if size: run.font.size = Pt(size)
        if color: run.font.color.rgb = color
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 1. 문서 최상단(혹은 적절한 위치)에 '데이터 처리 규모' 신설
    # 첫 번째 섹션 뒤에 삽입
    p = doc.paragraphs[3] # 적절한 위치 선정
    new_p = p.insert_paragraph_before("")
    run = new_p.add_run("■ [M-Medic] 전체 빅데이터 학습 및 처리 규모")
    set_font(run, size=12, bold=True, color=RGBColor(0, 51, 153))

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    hdr = summary_table.rows[0].cells
    headers = ['구분', '항목', '유닛 개수', '비고']
    for i, h in enumerate(headers):
        set_font(hdr[i].paragraphs[0].add_run(h), bold=True)

    data_rows = [
        ('머신러닝(ML)', '해상 사고 이력 데이터', '1,504개', 'GICOMS 10개년 데이터'),
        ('딥러닝(DL)', '외상 이미지 판독 데이터', '2,326개', '6종 정밀 분류 학습셋'),
        ('세그멘테이션', '상처 윤곽 추출 데이터', '5,520개', '이미지+마스크 페어 구성'),
        ('합계', '전체 학습 데이터 유닛', '9,350개', '프로젝트 총 분석 자산')
    ]
    
    for c1, c2, c3, c4 in data_rows:
        row = summary_table.add_row().cells
        set_font(row[0].paragraphs[0].add_run(c1))
        set_font(row[1].paragraphs[0].add_run(c2))
        set_font(row[2].paragraphs[0].add_run(c3), bold=(c1=='합계'))
        set_font(row[3].paragraphs[0].add_run(c4))

    # 2. 문서 내 4,212개 언급이 있다면 "원천 데이터 4,212개"로 명확히 함
    for p in doc.paragraphs:
        if "4,212" in p.text and "원천" not in p.text:
            p.text = p.text.replace("4,212개", "4,212개 (원천 데이터 자산 기준)")

    doc.save(file_path)
    print("SUCCESS: 3-3 Definition updated with 9,350 Total Learning Units.")

if __name__ == '__main__':
    update_to_9350_units()
