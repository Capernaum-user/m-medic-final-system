import os
from docx import Document
import markdown
from bs4 import BeautifulSoup

def md_to_docx(md_path, docx_path):
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    # Convert markdown to HTML (easier to parse for docx)
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create docx document
    doc = Document()
    doc.add_heading(os.path.basename(md_path), 0)
    
    # Simple conversion from HTML to docx
    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'pre':
            # Code block
            doc.add_paragraph(element.get_text(), style='No Spacing')
        elif element.name is not None:
            # Other elements as plain text
            doc.add_paragraph(element.get_text())
            
    doc.save(docx_path)

# Source and target directories
source_root = r'D:\GeminiUniverse\vscode-workspace\maritime-medic'
target_root = os.path.join(source_root, 'hancom_docs')

if not os.path.exists(target_root):
    os.makedirs(target_root)

# Find and convert all .md files (excluding ones already in export/hancom folders to avoid duplicates)
for root, dirs, files in os.walk(source_root):
    # Skip export folders to avoid circular conversion or redundancy
    if 'big_data_export' in root or 'hancom_docs' in root:
        continue
        
    for file in files:
        if file.endswith('.md'):
            md_file_path = os.path.join(root, file)
            # Create relative path to maintain structure if needed, or just flatten
            rel_path = os.path.relpath(md_file_path, source_root)
            docx_file_name = rel_path.replace(os.sep, '_').replace('.md', '.docx')
            docx_file_path = os.path.join(target_root, docx_file_name)
            
            print(f"Converting: {md_file_path} -> {docx_file_path}")
            try:
                md_to_docx(md_file_path, docx_file_path)
            except Exception as e:
                print(f"Failed to convert {file}: {e}")

print("Conversion complete.")
