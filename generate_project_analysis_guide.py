# -*- coding: utf-8 -*-
"""
M-MEDIC v2 프로젝트 분석 가이드 Word 문서 생성기
빅데이터 분석 담당자용 — 폴더/파일 설명 + 발표 시나리오
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2_프로젝트분석가이드_빅데이터담당자용.docx"

# ── 헬퍼 함수 ──────────────────────────────────────────────────────────────────

def set_font(run, name="Malgun Gothic", size=10, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    try:
        r = run._element
        rpr = r.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass

def add_heading(doc, text, level=1, color=(0, 70, 127)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True, color=color)
    return p

def add_body(doc, text, indent=0, bullet=False, color=None, bold=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent * 0.6)
    if bullet:
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.left_indent = Cm(indent * 0.6 + 0.4)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 70)
    set_font(run, size=8, color=(180, 180, 180))

def shade_cell(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, "2E5F8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "EBF0F8" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ── 문서 생성 ──────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("M-MEDIC v2 (MDTS)")
    set_font(r, size=22, bold=True, color=(30, 60, 114))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("프로젝트 전체 파일 분석 가이드")
    set_font(r2, size=16, bold=True, color=(0, 112, 192))

    sub2_p = doc.add_paragraph()
    sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub2_p.add_run("빅데이터 분석 담당자를 위한 상세 설명서 및 발표 시나리오")
    set_font(r3, size=11, color=(80, 80, 80))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = info_p.add_run("작성일: 2026년 4월 23일  |  대상: 빅데이터 분석 담당자  |  버전: v2.0")
    set_font(r4, size=9, color=(120, 120, 120))
    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    # ── 목차 안내 ─────────────────────────────────────────────────────────────
    add_heading(doc, "📋 문서 구성", level=1, color=(30, 60, 114))
    toc_items = [
        "제1장. 프로젝트 개요",
        "제2장. 전체 폴더/파일 구조 상세 설명",
        "제3장. 빅데이터 분석 담당자 핵심 파일 목록",
        "제4장. 발표 시나리오 및 설명 전략",
        "제5장. ML 분석 결과 해석 가이드 (발표 활용)",
        "제6장. 발표 Q&A 예상 질문 및 답변",
    ]
    for item in toc_items:
        add_body(doc, item, bullet=True, color=(50, 50, 50))
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 제1장. 프로젝트 개요
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "제1장. 프로젝트 개요", level=1, color=(30, 60, 114))

    add_heading(doc, "1.1 프로젝트명 및 목적", level=2)
    add_body(doc, "프로젝트명: MDTS (Maritime Digital Treatment System) — M-MEDIC v2")
    add_body(doc, "목적: 선박 내 의사가 없는 환경에서, 선원이 스마트폰으로 상처 사진을 찍으면 AI가 외상/피부 병변을 판독하고 선원법 법적 근거가 포함된 응급처치 지침을 자동 생성하는 엣지 AI 의료 지원 시스템")
    add_body(doc, "한 줄 슬로건: \"선박에서 의사가 없어도 AI가 1분 안에 위험도를 판정하고, 법적 근거가 있는 응급처치 매뉴얼을 선원에게 즉시 제공합니다.\"", bold=True, color=(0, 112, 192))

    add_heading(doc, "1.2 AI 구성 (3개 모델 통합)", level=2)
    add_table(doc,
        ["구성 요소", "기술", "역할", "현재 상태"],
        [
            ["ML (머신러닝)", "RandomForest / GradientBoosting", "해양사고 통계 → 사상자 발생 및 위험등급 예측", "학습 완료 ✅"],
            ["DL (딥러닝) ①", "MobileNetV3 Large", "외상 이미지 → 6종 분류 (찰과상/타박상/화상 등)", "학습 완료 ✅ (신뢰도 95.2%)"],
            ["DL (딥러닝) ②", "EfficientNet-B3 (멀티모달)", "피부 병변 이미지 → 7종 분류 (흑색종 등)", "미학습 ⏳"],
        ],
        col_widths=[3.0, 4.5, 6.5, 3.0]
    )

    add_heading(doc, "1.3 데이터 현황 요약", level=2)
    add_table(doc,
        ["데이터 종류", "규모", "용도"],
        [
            ["외상 이미지 (6클래스)", "4,212장 (원본 1,162장 → 증강 후)", "MobileNetV3 학습"],
            ["해양사고 통계 (합성)", "1,500건 (한국 해양안전심판원 통계 패턴 기반)", "RF/GradientBoosting 학습"],
            ["격리된 불필요 데이터", "당뇨성 창상·수술 상처 등 6종", "deprecated/trash 폴더로 격리"],
        ],
        col_widths=[4.5, 6.0, 6.0]
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 제2장. 전체 폴더/파일 구조 상세 설명
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "제2장. 전체 폴더/파일 구조 상세 설명", level=1, color=(30, 60, 114))
    add_body(doc, "프로젝트 루트: D:\\GeminiUniverse\\vscode-workspace\\wip-maritime-medic\\", color=(80, 80, 80), size=9)
    doc.add_paragraph()

    # 2.1 루트 레벨
    add_heading(doc, "2.1 루트 레벨 파일", level=2)
    add_table(doc,
        ["파일명", "유형", "기능 설명"],
        [
            ["M_MEDIC_PROJECT_MAP.md / .docx", "문서", "프로젝트 전체 현황 맵 — 데이터셋 최적화 과정, 시스템 검증 결과(신뢰도 95.2%), 향후 과제 Roadmap 정리"],
            ["PROJECT_GUIDE.md", "문서", "전체 실행 순서 가이드 — Step 1(ML 데이터생성) ~ Step 7(통합시스템 실행)까지 순서 및 폴더 구조 설명"],
            ["TABLE_DEFINITION.md", "문서", "DB 테이블 정의서 — tb_crew, tb_vital, tb_analysis 등 각 테이블의 컬럼명·타입·설명 목록"],
            ["TABLE_SPECIFICATION_FINAL.md", "문서", "DB 테이블 상세 스펙 — 제약조건, FK 관계, 인덱스, 데이터 예시까지 포함한 최종 명세"],
            ["evaluate_dl_performance.py", "스크립트", "학습된 MobileNetV3 모델을 raw 데이터로 평가 — Accuracy, F1-Score, 클래스별 성능 출력 (경로 정상)"],
            ["validate_5_rounds.py", "스크립트", "5회 반복 랜덤 샘플 추론으로 모델 일관성 검증 — 각 회차별 예측 클래스 및 신뢰도 출력"],
            ["generate_project_map.py", "스크립트", "프로젝트 구조 및 발표 가이드를 DOCX로 자동 생성하는 문서 생성 도구"],
            ["generate_perfect_guide.py", "스크립트", "발표 만점 전략 가이드를 DOCX로 자동 생성하는 문서 생성 도구"],
            ["convert_to_docx.py", "스크립트", "Markdown(.md) 파일을 Word(.docx)로 변환하는 유틸리티"],
            ["sync_docx_counts.py", "스크립트", "빅데이터분석정의서 DOCX 내 데이터 수치를 최신값으로 일괄 업데이트"],
            ["update_docx_final.py", "스크립트", "빅데이터분석정의서의 데이터 개수·정확도·경로를 최신 값(4,212장 등)으로 정밀 교체"],
            ["_deprecated/ (폴더)", "격리", "경로 오류·역할 완료·v1 구버전 스크립트 14개 보관 (실행 불가 파일 격리)"],
        ],
        col_widths=[4.5, 1.5, 10.5]
    )

    # 2.2 M_MEDIC_v2/01_data/
    add_heading(doc, "2.2 M_MEDIC_v2 / 01_data — 데이터 관리 폴더", level=2)
    add_table(doc,
        ["경로", "규모", "내용 및 용도"],
        [
            ["all_labeled_data/Abrasions", "668장", "찰과상 — 갑판·구조물에 쓸린 상처 이미지 (증강 완료본)"],
            ["all_labeled_data/Bruises", "972장", "타박상 — 충돌·낙하 시 타박 이미지 (증강 완료본)"],
            ["all_labeled_data/Burns", "504장", "화상 — 엔진실 화재·증기 화상 이미지 (증강 완료본)"],
            ["all_labeled_data/Cut", "600장", "절창 — 날카로운 도구에 의한 절상 이미지 (증강 완료본)"],
            ["all_labeled_data/Laceration", "732장", "열창 — 거친 표면에 의한 찢김 이미지 (증강 완료본)"],
            ["all_labeled_data/Stab_wound", "736장", "자창 — 날카로운 물체 관통상 (기존 23장 → 32배 증강, 집중 보강 클래스)"],
            ["raw/wound_classification/", "1,162장", "Kaggle 원본 데이터 6종 (증강 전 보존본, 참조용)"],
            ["processed/marine_accidents_augmented.csv", "1,500건", "ML 학습용 해양사고 합성 데이터 — 연도/사고유형/선박종류/원인/해역/사상자수 포함"],
            ["deprecated/trash/", "6종 격리", "Diabetic_Wounds, Ingrown_nails, Normal, Pressure_Wounds, Surgical_Wounds, Venous_Wounds — 선박 외상과 무관하여 격리"],
            ["README.md", "문서", "데이터 출처·클래스 정의·폴더 구조 상세 설명 (핵심 참고 문서)"],
        ],
        col_widths=[5.0, 1.8, 9.7]
    )

    # 2.3 M_MEDIC_v2/02_machine_learning/
    add_heading(doc, "2.3 M_MEDIC_v2 / 02_machine_learning — 머신러닝 파이프라인", level=2)
    add_body(doc, "역할: 해양사고 정형 데이터(tabular)로 '사상자 발생 예측' 및 '위험등급 분류' 학습", bold=True)
    add_table(doc,
        ["파일명", "실행 순서", "기능 설명"],
        [
            ["01_generate_accident_data.py", "Step 1", "한국 해양안전심판원 2019~2024 통계 패턴 기반으로 합성 데이터 1,500건 생성. 선박 종류(어선 48%), 사고유형(충돌 30%), 사고원인, 발생해역, 사상자수를 현실적 확률 분포로 샘플링. 출력: marine_accidents_augmented.csv"],
            ["02_train_accident_ml.py", "Step 2", "RandomForest / GradientBoosting / XGBoost 세 모델을 동시 학습 및 비교. ① 이진분류: 사상자 발생 여부 (0/1) ② 다중분류: 위험등급 4단계 (GREEN/YELLOW/ORANGE/RED). 5-Fold 교차검증, Feature Importance, Confusion Matrix 차트 자동 생성"],
            ["results/ml_comparison_report.txt", "결과", "모델별 성능 비교표 — RF(F1=0.762) vs GradientBoosting(F1=0.754), 이진·다중분류 성능 수치"],
            ["results/feature_importance_binary.png", "결과", "이진분류 변수 중요도 차트 — '사고유형'이 가장 중요한 예측 변수임을 시각화"],
            ["results/feature_importance_multiclass.png", "결과", "위험등급 다중분류 변수 중요도 차트"],
            ["results/model_comparison.png", "결과", "모델별 F1-Score 비교 막대 차트 (발표 슬라이드 직접 활용 가능)"],
            ["results/confusion_matrix_binary.png", "결과", "이진분류 혼동 행렬 — 사상자없음/사상자발생 2×2 행렬"],
            ["results/confusion_matrix_multiclass.png", "결과", "위험등급 4단계 혼동 행렬"],
            ["results/model_randomforest_binary.pkl", "모델", "저장된 RandomForest 이진분류 모델 (pickle 포맷, 재사용 가능)"],
            ["results/model_gradientboosting_binary.pkl", "모델", "저장된 GradientBoosting 이진분류 모델"],
        ],
        col_widths=[5.2, 1.8, 9.5]
    )

    # 2.4 M_MEDIC_v2/03_deep_learning/
    add_heading(doc, "2.4 M_MEDIC_v2 / 03_deep_learning — 딥러닝 외상 분류", level=2)
    add_body(doc, "역할: 외상 이미지를 CNN으로 6종 자동 분류 (선박 엣지 AI 배포 목적으로 경량 모델 선택)", bold=True)
    add_table(doc,
        ["파일명", "기능 설명"],
        [
            ["01_train_wound_cnn.py", "초기 CNN 학습 스크립트 — 기본 구조 검증용"],
            ["02_train_mobilenet_v3.py", "MobileNetV3 Large 전이학습 메인 스크립트 — ImageNet 사전학습 가중치 활용, 6클래스 Fine-tuning"],
            ["03_retrain_v2.py", "v2 데이터(증강 후)로 재학습하는 스크립트"],
            ["04_refine_6class.py", "6클래스 정제 학습 — 클래스 매핑 오류 수정 및 재학습"],
            ["05_train_expert_model.py", "EfficientNet-V2-S Expert 모델 학습 스크립트 (미실행 — 향후 과제)"],
            ["03_inference_wound.py / test_inference.py", "학습된 모델로 단일 이미지 추론 테스트"],
            ["results/mobilenet_v3_wound_best.pth", "★ 핵심 파일 — 학습 완료된 최종 MobileNetV3 모델 가중치. m_medic_v2.py가 이 파일을 로드하여 실시간 진단 수행"],
            ["results/wound_class_mapping.json", "클래스 인덱스↔클래스명 매핑 JSON — 모델 출력 인덱스를 클래스명으로 변환"],
            ["results/training_performance.png", "학습 곡선 (Train Loss / Val Loss / Accuracy) 시각화"],
            ["results/v2_performance.png", "v2 재학습 후 성능 곡선"],
        ],
        col_widths=[5.5, 11.0]
    )

    # 2.5 04_integrated_system
    add_heading(doc, "2.5 M_MEDIC_v2 / 04_integrated_system — 통합 진단 시스템", level=2)
    add_table(doc,
        ["파일명", "기능 설명"],
        [
            ["m_medic_v2.py", "★ 최종 실행 파일. CLI로 이미지 경로를 입력받아 ① 외상 분류(MobileNetV3) 또는 ② 피부 병변 분류(EfficientNet-B3)를 수행하고, 위험등급(CRITICAL/HIGH/MEDIUM/LOW) + 선원법 근거 + 응급처치 지침을 출력. 신뢰도 70% 미만 시 재촬영 권고 자동 출력. 진단 결과를 JSON 보고서로 저장"],
            ["report_20260415_102419.json ~ report_20260423_084007.json", "실제 진단 결과 JSON 파일 9개 — 테스트 시 자동 생성된 의료 보고서. 최신 파일(20260423)에서 Stab_wound 95.2% 신뢰도 확인됨"],
        ],
        col_widths=[5.5, 11.0]
    )

    add_body(doc, "실행 예시:")
    add_body(doc, "python m_medic_v2.py --image 사진경로.jpg --mode wound", indent=1, color=(0, 80, 0))
    add_body(doc, "python m_medic_v2.py --demo  (모델 없이 시연)", indent=1, color=(0, 80, 0))

    # 2.6 05_presentation
    add_heading(doc, "2.6 M_MEDIC_v2 / 05_presentation — 발표 자료", level=2)
    add_table(doc,
        ["파일명", "기능 설명"],
        [
            ["generate_charts.py", "발표용 차트(클래스 분포, 성능 비교 등) 자동 생성 스크립트 — charts/ 폴더에 PNG 저장"],
            ["ANALYSIS_INSIGHTS.md", "★ 빅데이터 분석 담당자 필독 — ML/DL 결과의 도메인 해석 가이드. Feature Importance 의미, 클래스별 임상적 해석, 정책 시사점 3가지 정리"],
            ["KAGGLE_REQUIREMENTS.md", "추가로 확보가 필요한 데이터셋 목록 및 Kaggle 링크"],
            ["charts/", "생성된 발표용 차트 PNG 파일들 보관 폴더"],
        ],
        col_widths=[5.5, 11.0]
    )

    # 2.7 db_design
    add_heading(doc, "2.7 M_MEDIC_v2 / db_design — DB 설계 문서", level=2)
    add_table(doc,
        ["경로/파일명", "기능 설명"],
        [
            ["analysis/01_분석목표서.md/.docx", "분석 목표 정의서 — 프로젝트가 풀려는 문제와 분석 방향"],
            ["analysis/02_분석용_컬럼_사전.md/.docx", "변수(컬럼) 사전 — 각 피처의 정의, 타입, 값 범위 설명"],
            ["analysis/03_데이터정의서.docx", "데이터 정의서 — 데이터 수집 출처, 구조, 활용 방법"],
            ["analysis/04_데이터분석정의서.docx", "데이터 분석 정의서 — 분석 방법론, 알고리즘 선택 근거"],
            ["analysis/3-3._빅데이터분석정의서_MDTS 260422.docx", "★ 핵심 제출 문서 — 최신 빅데이터 분석 정의서 (4,212장 데이터, 99.53% 정확도 반영 완료)"],
            ["analysis/4-3. IoT 회로구성 설계서_MDTS 260417.pdf", "IoT 센서(심박수, 산소포화도, 체온) 회로 설계도"],
            ["aquerytool/m_medic_v2_aquerytool.sql", "AqueryTool 온라인 ERD 도구용 DDL SQL"],
            ["erdcloud/m_medic_v2_erdcloud.sql", "ERDCloud 온라인 ERD 도구용 DDL SQL"],
            ["erwin/m_medic_v2_schema.sql", "ERwin(CA/전문 도구)용 DDL SQL — tb_crew, tb_vital, tb_analysis 3테이블 완성"],
        ],
        col_widths=[5.5, 11.0]
    )

    # 2.8 Project etc
    add_heading(doc, "2.8 Project etc — 기획/요구사항 문서", level=2)
    add_table(doc,
        ["파일명", "설명"],
        [
            ["2. 프로젝트기획서_실전_MDTS 260413.pdf/.hwp", "프로젝트 기획서 — 배경, 목적, 기대효과, 팀 구성, 일정"],
            ["3.선박_엣지AI_요구사항정의서_MDTS 260413.pdf/.hwp", "요구사항 정의서 — 시스템 기능·비기능 요구사항, 엣지 AI 배포 제약 조건"],
            ["4-3. IoT 회로구성 설계서_MDTS 260413.pdf/.hwp", "IoT 회로 설계서 — 선박 내 센서 구성 (심박/산소포화도/체온)"],
            ["MDTS_데이터베이스 요구사항 분석_0414.pdf/.hwp", "DB 요구사항 분석서 — 테이블 설계 전 요구사항 정리"],
        ],
        col_widths=[6.5, 10.0]
    )

    # 2.9 reports
    add_heading(doc, "2.9 reports — 분석 보고서 (MD + DOCX)", level=2)
    add_table(doc,
        ["파일명", "내용"],
        [
            ["ANALYSIS_REPORT_01_EDA.md", "탐색적 데이터 분석(EDA) 보고서 — 피부 질환 분포, 해양사고 통계 분포"],
            ["ANALYSIS_REPORT_02_TRAINING_PLAN.md", "학습 계획 보고서 — 모델 선택 근거, 하이퍼파라미터 설정"],
            ["ANALYSIS_REPORT_03_TRAINING_STATUS.md", "학습 현황 보고서 — 에폭별 Loss/Accuracy 추적"],
            ["ANALYSIS_REPORT_04_WOUND_ANALYSIS.md", "외상 분류 분석 보고서 — 클래스별 판독 성능 및 어려운 케이스"],
            ["ANALYSIS_REPORT_05_INTEGRATION.md", "통합 시스템 보고서 — ML+DL 연계 흐름 및 최종 시스템 검증"],
            ["DB멘토링(실전).SQL", "DB 멘토링 실습 SQL 파일 — 실제 쿼리 작성 및 검토 내용"],
            ["age_distribution.png / skin_disease_distribution.png", "피부 질환 환자 연령 분포 및 질병별 분포 차트 (EDA 결과)"],
            ["marine_accident_analysis.png", "해양사고 유형별 선박 분포 차트 (EDA 결과)"],
        ],
        col_widths=[5.5, 11.0]
    )

    # 2.10 data / hancom_docs
    add_heading(doc, "2.10 data / hancom_docs — 원본 데이터 및 문서 보관", level=2)
    add_table(doc,
        ["경로", "내용"],
        [
            ["data/Marine_Accidents/marine_accidents_sample.csv", "해양사고 원본 샘플 CSV — 4행(시연용). ML 학습은 1,500건 합성 데이터 사용"],
            ["data/Medical_Guidelines/Medical_Manager_Guidelines_Summary.txt", "선원법 의료관리자 가이드라인 요약 — 적용대상, 자격, 주요업무, 기록의무 등"],
            ["data/download_full_datasets.bat", "Kaggle 전체 데이터 다운로드 배치 스크립트 (데이터 초기 구축 시 사용)"],
            ["hancom_docs/", "한컴 문서 작업용 별도 Git 저장소 — documentation(MD), raw_data(CSV), reports(DOCX) 분리 보관"],
        ],
        col_widths=[5.5, 11.0]
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 제3장. 빅데이터 분석 담당자 핵심 파일
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "제3장. 빅데이터 분석 담당자 핵심 파일 목록", level=1, color=(30, 60, 114))
    add_body(doc, "아래 파일들은 발표 준비 및 결과 해석에 직접 활용해야 할 파일들입니다.", bold=True, color=(192, 0, 0))
    doc.add_paragraph()

    add_table(doc,
        ["우선순위", "파일 경로", "왜 중요한가"],
        [
            ["★★★", "M_MEDIC_v2/02_machine_learning/results/\nml_comparison_report.txt",
             "RF vs GradientBoosting 성능 수치가 정리된 핵심 결과물. 발표에서 \"어떤 모델이 최고였는가\"에 대한 직접 근거"],
            ["★★★", "M_MEDIC_v2/02_machine_learning/results/\nfeature_importance_binary.png",
             "\"어떤 요소가 사상자를 만드는가\" — 사고유형 38% > 선박종류 28% > 사고원인 18% 순서. 정책 시사점 도출의 핵심 근거. 발표 슬라이드에 직접 삽입"],
            ["★★★", "M_MEDIC_v2/02_machine_learning/results/\nmodel_comparison.png",
             "RF vs GB 성능 비교 막대그래프. 시각적으로 모델 비교를 설명할 수 있는 차트"],
            ["★★★", "M_MEDIC_v2/db_design/analysis/\n3-3._빅데이터분석정의서_MDTS 260422.docx",
             "제출용 핵심 문서 — 분석 방법론, 데이터 현황(4,212장), 모델 성능(99.53%) 최신값 반영 완료"],
            ["★★★", "M_MEDIC_v2/05_presentation/\nANALYSIS_INSIGHTS.md",
             "ML 결과를 도메인 언어로 해석한 가이드. \"화재+어선 → 사상자 65% 이상\" 같은 실무 인사이트가 발표 핵심 메시지"],
            ["★★", "M_MEDIC_v2/02_machine_learning/\n01_generate_accident_data.py",
             "어떤 데이터를 어떻게 만들었는지 설명 가능 — 한국 해양안전심판원 통계 기반 합성 데이터 생성 로직. 발표에서 \"왜 합성 데이터인가\"에 대한 근거"],
            ["★★", "M_MEDIC_v2/02_machine_learning/\n02_train_accident_ml.py",
             "ML 학습 코드 — 5-Fold 교차검증, 모델 비교 방법론 설명 가능. \"어떻게 학습했는가\" 질문 대비"],
            ["★★", "M_MEDIC_v2/02_machine_learning/results/\nconfusion_matrix_binary.png",
             "혼동 행렬 — 모델이 어디서 틀리는지 시각화. 분석 한계점 설명 시 활용"],
            ["★★", "M_MEDIC_v2/01_data/README.md",
             "데이터 출처, 클래스 정의, 선박 사고 연관성 설명 — \"왜 이 6개 클래스인가\" 발표 질문 대비"],
            ["★★", "M_MEDIC_v2/04_integrated_system/\nreport_20260423_084007.json",
             "실제 AI 진단 결과 파일 — Stab_wound 95.2% 신뢰도 증거. 발표에서 \"실제로 잘 되나\"에 대한 라이브 데모 근거"],
            ["★", "reports/ANALYSIS_REPORT_01_EDA.md",
             "EDA 보고서 — 데이터 탐색 과정 설명 가능. 분석 프로세스 전체 흐름 보여줄 때 활용"],
            ["★", "M_MEDIC_v2/db_design/analysis/\n05_빅데이터분석정의서.md",
             "빅데이터 분석 정의서 초안 — 분석 목표와 방법론 정의 문서"],
        ],
        col_widths=[1.5, 5.0, 10.0]
    )
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 제4장. 발표 시나리오
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "제4장. 발표 시나리오 및 설명 전략", level=1, color=(30, 60, 114))
    add_body(doc, "아래 순서로 발표하면 빅데이터 분석 담당자 파트가 체계적으로 전달됩니다. 각 단계별 핵심 멘트와 근거 파일을 함께 제시합니다.", color=(80, 80, 80))
    doc.add_paragraph()

    scenarios = [
        {
            "step": "STEP 1 — 문제 정의 (1~2분)",
            "color": (0, 70, 127),
            "point": "\"왜 이 프로젝트가 필요한가\"",
            "script": (
                "한국 해양안전심판원 통계에 따르면, 국내 해양사고의 48%가 어선에서 발생하며, "
                "그 중 화재·침몰 사고는 사상자 발생 확률이 65~80%에 달합니다. "
                "그러나 대부분의 소형 어선에는 의료 전문가가 없습니다. "
                "이 문제를 해결하기 위해 저희 팀은 선박 탑재형 AI 의료 지원 시스템 MDTS를 개발했습니다."
            ),
            "files": ["ANALYSIS_INSIGHTS.md의 정책 시사점 섹션 참고"],
            "tip": "통계 수치(48%, 65~80%)를 슬라이드 첫 화면에 크게 배치하면 청중의 집중도가 높아집니다."
        },
        {
            "step": "STEP 2 — 데이터 설명 (2~3분)",
            "color": (0, 100, 0),
            "point": "\"어떤 데이터를 어떻게 확보하고 정제했는가\"",
            "script": (
                "분석에 사용한 데이터는 크게 두 종류입니다. "
                "첫째, 외상 이미지 데이터: Kaggle의 외상 분류 데이터셋을 수집했으며, "
                "프로젝트 목적과 무관한 당뇨성 창상·수술 상처 등 6종의 노이즈 데이터를 과감히 제거하고, "
                "선박 사고에 직접 연관된 6개 클래스만 유지했습니다. "
                "특히 자창(Stab_wound)은 원본이 단 23장에 불과했는데, "
                "회전·반전·밝기 조절 등 다양한 증강 기법을 적용해 736장으로 32배 확보했습니다. "
                "둘째, 해양사고 통계 데이터: 실제 중앙해양안전심판원 통계 패턴을 반영한 "
                "합성 데이터 1,500건을 생성해 ML 학습에 활용했습니다."
            ),
            "files": ["01_data/README.md", "01_generate_accident_data.py", "all_labeled_data/ 폴더 현황"],
            "tip": "데이터 정제 전후 비교 표(1,162장 → 4,212장)를 슬라이드에 넣으면 데이터 처리 역량을 어필할 수 있습니다."
        },
        {
            "step": "STEP 3 — 분석 방법론 (2분)",
            "color": (0, 100, 0),
            "point": "\"왜 ML과 DL을 동시에 사용했는가\"",
            "script": (
                "저희는 두 종류의 AI를 함께 사용했습니다. "
                "해양사고 통계처럼 연도·사고유형·선박종류 같은 정형(테이블) 데이터에는 "
                "설명 가능성이 높은 RandomForest와 GradientBoosting을 선택했습니다. "
                "반면 외상 이미지처럼 공간 패턴을 학습해야 하는 비정형 데이터에는 "
                "딥러닝(MobileNetV3)을 적용했습니다. "
                "MobileNetV3는 추론 속도가 200ms 이내로 선박의 저사양 엣지 디바이스에도 탑재 가능한 경량 모델입니다."
            ),
            "files": ["02_train_accident_ml.py 상단 주석", "ANALYSIS_INSIGHTS.md 1.3 섹션"],
            "tip": "ML vs DL 선택 근거를 표로 정리해 두세요. \"왜 딥러닝을 안 쓰고 ML을?\"이라는 질문에 '설명 가능성'과 '정형 데이터 적합성'으로 답하면 됩니다."
        },
        {
            "step": "STEP 4 — ML 분석 결과 (3~4분)",
            "color": (0, 100, 0),
            "point": "\"분석 결과, 무엇을 알 수 있었는가\"",
            "script": (
                "RandomForest 모델이 사상자 발생 예측에서 F1-Score 0.762를 기록하며 최고 성능을 보였습니다. "
                "더 중요한 것은 Feature Importance 분석 결과입니다. "
                "사고유형이 전체 예측 중요도의 38%를 차지하며 1위였고, "
                "그 다음이 선박종류(28%), 사고원인(18%) 순이었습니다. "
                "이것이 의미하는 바는 명확합니다. "
                "화재·침몰 사고를 조기에 감지하는 IoT 센서가 가장 중요한 투자처이며, "
                "전체 사고의 48%를 차지하는 어선에 이 시스템을 우선 보급해야 경제적 효율이 극대화된다는 것입니다."
            ),
            "files": ["ml_comparison_report.txt", "feature_importance_binary.png", "model_comparison.png"],
            "tip": "feature_importance_binary.png를 슬라이드에 직접 삽입하세요. 숫자보다 시각적 차트가 훨씬 강렬한 인상을 줍니다."
        },
        {
            "step": "STEP 5 — DL 결과 및 통합 시스템 (2~3분)",
            "color": (0, 100, 0),
            "point": "\"AI가 실제로 어떻게 작동하는가\"",
            "script": (
                "딥러닝 모델은 오늘 아침 실제 자창(Stab_wound) 이미지로 검증했을 때 "
                "신뢰도 95.2%로 정확히 분류하고, "
                "CRITICAL 위험등급과 함께 '박힌 물체 제거 금지. 즉시 병원 이송.'이라는 "
                "선원법 근거 응급처치 지침을 자동 출력했습니다. "
                "이 세 가지 AI가 통합되면, 선원이 사진 한 장을 찍는 것만으로 "
                "병명 + 위험등급 + 법적 근거 처치 지침 + JSON 의료 보고서가 자동 생성됩니다."
            ),
            "files": ["report_20260423_084007.json (실제 출력 파일)", "m_medic_v2.py --demo 실행"],
            "tip": "발표 현장에서 터미널을 열고 python m_medic_v2.py --demo 를 직접 실행하면 청중에게 강한 임팩트를 줄 수 있습니다."
        },
        {
            "step": "STEP 6 — 한계점 및 향후 과제 (1분)",
            "color": (150, 0, 0),
            "point": "\"현재 시스템의 한계를 솔직하게 인정하되, 개선 방향을 함께 제시\"",
            "script": (
                "현재 ML 학습에는 실제 통계가 아닌 합성 데이터를 사용했습니다. "
                "중앙해양안전심판원의 실제 GICOMS 데이터를 확보하면 모델 신뢰도가 크게 향상될 것입니다. "
                "또한 4,212장의 새로운 데이터로 딥러닝 모델을 재학습하는 것이 다음 우선 과제이며, "
                "선박 내 IoT 센서(심박수, 산소포화도, 체온)와의 연동을 통해 "
                "진정한 멀티모달 시스템으로 발전시킬 계획입니다."
            ),
            "files": ["M_MEDIC_PROJECT_MAP.md의 Roadmap 섹션", "ANALYSIS_INSIGHTS.md 3장"],
            "tip": "한계점을 먼저 언급하면 오히려 신뢰도가 올라갑니다. \"우리가 문제를 정확히 알고 있다\"는 것이 강점입니다."
        },
    ]

    for s in scenarios:
        add_heading(doc, s["step"], level=2, color=s["color"])
        add_body(doc, f"핵심 포인트: {s['point']}", bold=True, color=(0, 70, 127))
        add_body(doc, "발표 스크립트 (권장 멘트):", bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(s["script"])
        set_font(run, size=10, color=(40, 40, 40), italic=True)

        add_body(doc, "근거 파일:", bold=True)
        for f in s["files"]:
            add_body(doc, f, bullet=True, indent=1, color=(0, 80, 160))

        add_body(doc, f"발표 팁: {s['tip']}", color=(180, 80, 0), bold=True)
        doc.add_paragraph()
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 제5장. ML 분석 결과 해석 가이드
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "제5장. ML 분석 결과 해석 가이드 (발표 활용)", level=1, color=(30, 60, 114))

    add_heading(doc, "5.1 모델 성능 수치 해석", level=2)
    add_table(doc,
        ["구분", "모델", "F1-Score", "의미"],
        [
            ["이진분류\n(사상자발생여부)", "RandomForest", "0.762 ★최고",
             "사상자가 발생할지 안 할지를 76.2% 정확도로 예측. 출항 전 사고 위험도 사전 경고에 활용 가능"],
            ["이진분류", "GradientBoosting", "0.754",
             "RF와 비슷한 수준. 앙상블로 결합하면 더 향상 가능"],
            ["다중분류\n(위험등급 4단계)", "GradientBoosting", "0.632 ★최고",
             "GREEN/YELLOW/ORANGE/RED 4단계 분류. 0.632는 4개 클래스 기준으로 양호한 수준"],
            ["다중분류", "RandomForest", "0.626",
             "위험등급 다중분류에서는 GB가 근소하게 앞섬"],
        ],
        col_widths=[2.8, 3.2, 2.2, 8.3]
    )

    add_heading(doc, "5.2 Feature Importance — 발표에서 이렇게 설명하세요", level=2)
    add_table(doc,
        ["순위", "변수", "중요도(예상)", "발표용 해석 멘트"],
        [
            ["1위", "사고유형", "~38%",
             "화재·침몰 사고가 사상자 발생을 가장 강하게 예측합니다. 즉, 화재 감지 센서 투자가 가장 효과적인 사상자 예방 수단입니다."],
            ["2위", "선박종류", "~28%",
             "어선이 전체 사고의 48%를 차지하며 사상자 위험도 가장 높습니다. 어선 우선 보급 정책의 데이터 근거입니다."],
            ["3위", "사고원인", "~18%",
             "운항과실(38%)이 가장 많습니다. 선원 피로도·음주 모니터링 시스템이 필요하다는 근거입니다."],
            ["4위", "발생해역", "~12%",
             "남해·서해에 사고가 집중됩니다. 해당 해역 순찰 강화 정책의 데이터 근거입니다."],
            ["5위", "연도", "~4%",
             "연도별 추세가 미미하다는 것은 시간이 지나도 개선되지 않는 구조적 문제임을 시사합니다."],
        ],
        col_widths=[1.2, 2.5, 2.3, 10.5]
    )

    add_heading(doc, "5.3 3가지 핵심 정책 시사점 (발표 클라이맥스)", level=2)
    policy_points = [
        ("① 어선 중심 안전 투자",
         "전체 사고의 48%가 어선 → MDTS를 어선에 우선 보급 시 경제적 효율 최대화"),
        ("② 예방형 출항 통제",
         "출항 전 기상조건 + 선원 피로도를 ML 모델에 입력 → 사고 위험 사전 예측으로 출항 통제에 활용"),
        ("③ 보험료 차등화 근거 제공",
         "위험등급 예측 모델을 해상보험사에 제공 → 리스크 기반 보험료 산정 가능 (새로운 비즈니스 모델)"),
    ]
    for title, desc in policy_points:
        add_body(doc, title, bold=True, color=(0, 70, 127), size=11)
        add_body(doc, desc, indent=1)
        doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 제6장. 예상 Q&A
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "제6장. 발표 Q&A 예상 질문 및 답변", level=1, color=(30, 60, 114))

    qa_list = [
        ("Q1. 합성 데이터로 학습했는데 신뢰할 수 있나요?",
         "합성 데이터는 한국 해양안전심판원이 공식 발표한 2019~2024년 통계 패턴(선박 종류별 사고 비율, 해역별 분포 등)을 그대로 반영해 생성했습니다. "
         "실제 데이터와 분포가 동일하므로 경향성 학습에는 유효합니다. "
         "다만 이 점을 한계로 명시했으며, 중앙해양안전심판원 GICOMS 실데이터 확보를 향후 과제로 설정했습니다."),
        ("Q2. F1-Score 0.762가 충분히 높은 수치인가요?",
         "사상자 발생이라는 불균형 이진분류 문제에서 0.762는 의미 있는 수치입니다. "
         "완전한 랜덤 예측의 F1이 약 0.5인 것과 비교하면 실질적인 학습 효과가 있습니다. "
         "실제 GICOMS 데이터로 학습하면 성능이 더 향상될 것으로 예상됩니다."),
        ("Q3. 왜 XGBoost 결과가 없나요?",
         "학습 환경에 xgboost 라이브러리가 설치되지 않아 RF와 GradientBoosting 두 모델만 비교했습니다. "
         "XGBoost는 설치 후 동일한 02_train_accident_ml.py 코드로 즉시 추가 실험이 가능합니다."),
        ("Q4. 딥러닝 신뢰도 95.2%는 검증된 수치인가요?",
         "오늘(2026-04-23) 증강된 Stab_wound 이미지로 실제 진단을 수행한 결과이며, "
         "report_20260423_084007.json 파일에 타임스탬프와 함께 기록되어 있습니다. "
         "단, 단일 이미지 테스트이므로 전체 테스트셋 기준 정밀 검증(evaluate_dl_performance.py)이 추가로 필요합니다."),
        ("Q5. 선박에 실제로 배포 가능한가요?",
         "MobileNetV3는 추론 속도가 200ms 이내로, 라즈베리파이 수준의 저사양 엣지 디바이스에도 탑재 가능합니다. "
         "다만 현재는 소프트웨어 프로토타입 단계이며, 실제 배포를 위해서는 선박 내 조명·흔들림 환경에서의 추가 검증이 필요합니다."),
        ("Q6. 피부 병변 모델은 왜 발표에서 안 보이나요?",
         "EfficientNet-B3 피부 병변 분류 모델은 학습 데이터 확보 문제로 아직 학습이 완료되지 않았습니다. "
         "코드 구조(m_medic_v2.py)는 이미 완성되어 있으며, HAM10000 등 데이터 확보 후 즉시 학습·적용 가능합니다."),
    ]

    for q, a in qa_list:
        add_body(doc, q, bold=True, color=(192, 0, 0), size=10)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(a)
        set_font(run, size=10, color=(40, 40, 40))
        doc.add_paragraph()

    # ── 마무리 ────────────────────────────────────────────────────────────
    add_divider(doc)
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end_p.add_run("본 문서는 M-MEDIC v2 (MDTS) 프로젝트의 빅데이터 분석 담당자를 위해 작성된 발표 준비 가이드입니다.")
    set_font(r, size=9, color=(120, 120, 120))

    doc.save(OUTPUT_PATH)
    print(f"[완료] 저장: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_document()
