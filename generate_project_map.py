# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def update_guide_with_presentation_strategy():
    doc = Document()
    
    def set_font(run, font_name='Malgun Gothic'):
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 제목
    heading = doc.add_heading('', 0)
    run = heading.add_run('M-Medic (해상 의료 AI) 프로젝트 구조 및 발표 가이드')
    set_font(run)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 폴더 구조 및 데이터 현황
    doc.add_heading('', level=1).add_run('1. 폴더 구조 및 데이터 현황').font.name = 'Malgun Gothic'
    doc.add_paragraph('• M_MEDIC_v2: AI 모델 및 통합 실행 시스템 (v2.0)').style = 'List Bullet'
    doc.add_paragraph('• 01_data: 총 4,212장의 상처 사진 및 사고 데이터 보유 (응급 외상 6종 집중)').style = 'List Bullet'

    # 2. AI 모델 성능 결과
    doc.add_heading('', level=1).add_run('2. AI 모델 성능 및 오차 분석 결과').font.name = 'Malgun Gothic'
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['분석 항목', '머신러닝 (사고 예측)', '딥러닝 (상처 분류)']):
        set_font(hdr[i].paragraphs[0].add_run(txt))
        
    metrics = [
        ('정확도 (Accuracy)', '76.4%', '93.4%'),
        ('일관성 (F1-Score)', '0.76 (안정적)', '0.93 (매우 우수)'),
        ('주요 특이사항', '사고 등급 예측 오차 존재', '열상(Laceration) 미세 혼동')
    ]
    for item, ml, dl in metrics:
        row = table.add_row().cells
        set_font(row[0].paragraphs[0].add_run(item))
        set_font(row[1].paragraphs[0].add_run(ml))
        set_font(row[2].paragraphs[0].add_run(dl))

    # 3. 발표 및 시연 전략 (신규 추가 섹션)
    doc.add_heading('', level=1).add_run('3. 필승! 발표 및 시연 전략').font.name = 'Malgun Gothic'
    
    # 3-1. 추천 모델
    p1 = doc.add_paragraph()
    run = p1.add_run('① 추천 시연 모델: 상처 분류 딥러닝 모델 (MobileNetV3 Large)')
    run.bold = True
    set_font(run)
    doc.add_paragraph('   - 이유: 시각적 효과가 크고 정확도가 93% 이상으로 시연 성공률이 매우 높음.').style = 'List Bullet'
    doc.add_paragraph('   - 강점: 저사양 기기에서도 작동하는 가벼운 고성능 모델임을 강조.').style = 'List Bullet'

    # 3-2. 시연 시나리오
    p2 = doc.add_paragraph()
    run = p2.add_run('② 시연 시나리오 (3단계)')
    run.bold = True
    set_font(run)
    doc.add_paragraph('   1단계 (상황): "의사가 없는 배 위에서 사고 발생! 선원이 상처 사진을 촬영합니다."').style = 'List Number'
    doc.add_paragraph('   2단계 (실행): m_medic_v2.py 실행하여 실시간 판독 및 선원법 조치사항 출력 장면 시연.').style = 'List Number'
    doc.add_paragraph('   3단계 (마무리): 분석 결과가 담긴 자동 생성 보고서(JSON/Word)를 보여주며 신뢰성 강조.').style = 'List Number'

    # 3-3. 핵심 발표 멘트
    p3 = doc.add_paragraph()
    run = p3.add_run('③ 핵심 발표 멘트 (Key Message)')
    run.bold = True
    set_font(run)
    doc.add_paragraph('   "우리는 4,212장의 빅데이터를 활용했습니다."').style = 'List Bullet 2'
    doc.add_paragraph('   "정확도 93.4%의 AI가 선원법 지침에 따라 즉각적인 조치 가이드를 제공합니다."').style = 'List Bullet 2'
    doc.add_paragraph('   "이것은 단순한 AI 기술을 넘어, 해상 생존권을 보장하는 지능형 의료 지원 시스템입니다."').style = 'List Bullet 2'

    # 4. 실행 가이드 요약
    doc.add_heading('', level=1).add_run('4. 실행 방법 요약').font.name = 'Malgun Gothic'
    doc.add_paragraph('• 시스템 시연: m_medic_v2.py --demo')
    doc.add_paragraph('• 실제 진단: m_medic_v2.py --mode wound --image "사진경로"')

    save_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_PROJECT_MAP.docx"
    doc.save(save_path)
    print("SUCCESS: Guide updated with presentation strategy.")

if __name__ == '__main__':
    update_guide_with_presentation_strategy()
