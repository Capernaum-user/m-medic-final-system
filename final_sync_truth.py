# -*- coding: utf-8 -*-
from docx import Document
import os

def refine_to_absolute_truth():
    file_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx"
    doc = Document(file_path)

    # 1. 정밀 타격 교체 맵 (중복 제거 및 수치 정정)
    replace_map = {
        "4,212개": "2,326개", # 이미지 개수 부분만 타겟팅 (나중에 전체 관리 수치는 별도 처리)
        "3,580": "1,977",     # 2,326의 85% Train
        "632": "349",         # 2,326의 15% Val
        "1,504개": "1,504개 (해상 사고 이력 데이터)",
        "23개": "215개",       # 자창 데이터 보강 수치
    }

    # 2. 클래스별 실제 개수 (2,326개 기반)
    class_counts = {
        "Abrasions   (찰과상) :": "Abrasions   (찰과상) : 363개",
        "Bruises     (타박상) :": "Bruises     (타박상) : 400개",
        "Burns                      (화상)   :": "Burns                      (화상)   : 385개",
        "Cut         (절상)   :": "Cut         (절상)   : 412개",
        "Laceration  (열상)   :": "Laceration  (열상)   : 551개",
        "Stab_wound  (자창)   :": "Stab_wound  (자창)   : 215개"
    }

    # 3. 문서 전체 스캔 및 수정
    for p in doc.paragraphs:
        # 이미지 총량 부분 (문맥 확인 후 교체)
        if "외상 이미지" in p.text and "4,212" in p.text:
            p.text = p.text.replace("4,212", "2,326")
        
        # 일반 수치 교체
        for old, new in replace_map.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # '합계 : 4,212'를 '합계 : 2,326'으로 수정 (이미지 섹션일 경우)
                    if "합계" in p.text and "4,212" in p.text:
                        p.text = p.text.replace("4,212", "2,326")
                    
                    # 클래스별 개수 정밀 교정
                    for key, val in class_counts.items():
                        if key in p.text:
                            p.text = val
                    
                    # 기타 수치 교체
                    for old, new in replace_map.items():
                        if old in p.text:
                            p.text = p.text.replace(old, new)

    # 4. 데이터 정의서 하단에 전체 자산 관리 정보 추가 (혼동 방지)
    # 이미 존재한다면 중복 방지 로직 필요하나, 요청하신 대로 핵심만 수정
    
    doc.save(file_path)
    print("SUCCESS: Data counts synchronized. Images: 2,326, Accident Records: 1,504.")

if __name__ == '__main__':
    refine_to_absolute_truth()
