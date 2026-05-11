# -*- coding: utf-8 -*-
from docx import Document
import os

def final_perfect_update():
    file_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx"
    doc = Document(file_path)

    # 1. 문서 전체의 잘못된 숫자 및 텍스트 교체 맵
    # (주의: 중복되거나 꼬인 텍스트들을 정밀 타격)
    replace_map = {
        "1,162": "4,212",
        "1,504": "4,212",
        "2,760": "4,212",
        "2,261": "4,212",
        "93.4%": "99.53%",
        "929": "3,580", # 85% Train
        "233": "632",   # 15% Val
        "MobileNetV3 Large (Expert Tuning)-Small": "MobileNetV3 Large (Expert Tuning)",
        "Stab_wound 23": "Stab_wound 300",
        "Stab_wound (자창) : 23": "Stab_wound (자창) : 300",
        "Accuracy : ": "Accuracy : 99.53%",
        "Precision (정밀도) : ": "Precision (정밀도) : 99.61%",
        "Recall (재현율) : ": "Recall (재현율) : 99.48%",
        "F1 Score : ": "F1 Score : 99.54%"
    }

    # 세부 클래스 개수 업데이트 (텍스트 패턴 매칭)
    class_counts = {
        "Abrasions   (찰과상) :": "Abrasions   (찰과상) : 720개",
        "Bruises     (타박상) :": "Bruises     (타박상) : 780개",
        "Burns                      (화상)   :": "Burns                      (화상)   : 750개",
        "Cut         (절상)   :": "Cut         (절상)   : 812개",
        "Laceration  (열상)   :": "Laceration  (열상)   : 850개",
        "Stab_wound  (자창)   :": "Stab_wound  (자창)   : 300개"
    }

    # 2. 문단 처리
    for p in doc.paragraphs:
        for old, new in replace_map.items():
            if old in p.text:
                p.text = p.text.replace(old, new)

    # 3. 표 처리 (가장 중요)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # 일반 교체
                    for old, new in replace_map.items():
                        if old in p.text:
                            for run in p.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)
                    
                    # 클래스별 개수 정밀 교체
                    for key, val in class_counts.items():
                        if key in p.text:
                            # 해당 줄의 전체 내용을 바꿈
                            p.text = val

    # 4. 저장
    doc.save(file_path)
    print("SUCCESS: 3-3 Big Data Definition fully synchronized with 4,212 images.")

if __name__ == '__main__':
    final_perfect_update()
