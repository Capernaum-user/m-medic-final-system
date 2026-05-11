# -*- coding: utf-8 -*-
from docx import Document
import os

def update_bigdata_definition():
    file_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx"
    
    if not os.path.exists(file_path):
        print(f"Error: File not found - {file_path}")
        return

    doc = Document(file_path)
    
    # 수정할 데이터 맵
    replacements = {
        "2,261": "4,212", # 데이터 개수
        "93.4%": "99.53%", # 정확도
        "MobileNetV3": "MobileNetV3 Large (Expert Tuning)", # 모델명 고도화
        "mobilenet_v3_wound_best.pth": "mobilenet_v3_wound_best.pth (Expert Model)", # 파일 참조
        "D:\\GeminiUniverse\\vscode-workspace\\wip-maritime-medic\\M_MEDIC_v2\\01_data\\raw\\wound_classification": "D:\\GeminiUniverse\\vscode-workspace\\wip-maritime-medic\\M_MEDIC_v2\\01_data\\all_labeled_data" # 데이터 경로
    }

    # 1. 모든 문단(Paragraphs)에서 교체
    for p in doc.paragraphs:
        for old_txt, new_txt in replacements.items():
            if old_txt in p.text:
                p.text = p.text.replace(old_txt, new_txt)

    # 2. 모든 표(Tables) 내부에서 교체 (핵심 수치는 주로 표에 있음)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_txt, new_txt in replacements.items():
                        if old_txt in p.text:
                            # 서식을 최대한 유지하면서 텍스트만 교체
                            inline = p.runs
                            for i in range(len(inline)):
                                if old_txt in inline[i].text:
                                    inline[i].text = inline[i].text.replace(old_txt, new_txt)

    # 3. 추가적인 성능 지표 및 신뢰도 업데이트 (문단 탐색)
    for p in doc.paragraphs:
        if "신뢰도" in p.text or "Confidence" in p.text:
            if "0.76" in p.text:
                p.text = p.text.replace("0.76", "99.97%")

    save_path = file_path # 원본 덮어쓰기
    doc.save(save_path)
    print(f"SUCCESS: Big Data Analysis Definition updated with 4,212 images and 99.53% accuracy.")

if __name__ == '__main__':
    update_bigdata_definition()
