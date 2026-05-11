# -*- coding: utf-8 -*-
"""
3-3._빅데이터분석정의서_MDTS 260422.docx 내용 정확도 수정 스크립트
- 양식(서식·색상·표 구조) 유지
- 실측값 기준으로 수치·오기·미완성 항목만 교체
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SRC = Path(r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx")
BAK = SRC.with_name(SRC.stem + "_원본백업.docx")

# ── 백업 ──────────────────────────────────────────────────────────────────────
shutil.copy2(SRC, BAK)
print(f"[백업] {BAK.name}")

doc = Document(str(SRC))

# ── 헬퍼: 셀 내용 전체 교체 (서식 보존) ──────────────────────────────────────
def rewrite_cell(cell, lines: list[str], font_name: str, font_pt: float, bold: bool = False):
    """
    cell의 기존 단락을 모두 제거하고 lines 목록으로 재작성.
    각 줄은 별도 단락. 서식(폰트·크기·볼드·동아시아폰트)은 인자로 지정.
    """
    # 기존 단락 텍스트 전체 제거 (단락 요소 자체는 첫 것만 유지)
    tc = cell._tc
    # 모든 단락 요소 수집
    paras = cell.paragraphs
    # 첫 단락 클리어
    first_para = paras[0]
    for run in list(first_para.runs):
        run._element.getparent().remove(run._element)
    # 나머지 단락 제거
    for para in list(paras[1:]):
        para._element.getparent().remove(para._element)

    def _add_run(para, text):
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_pt)
        run.bold = bold
        # 동아시아 폰트 설정 (한글 깨짐 방지)
        rpr = run._element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:eastAsia"), font_name)
        rf.set(qn("w:ascii"), font_name)
        rf.set(qn("w:hAnsi"), font_name)
        return run

    _add_run(first_para, lines[0])
    for line in lines[1:]:
        new_para = cell.add_paragraph()
        _add_run(new_para, line)


# ── 수정 대상 ─────────────────────────────────────────────────────────────────
t2 = doc.tables[2]   # 메인 분석 명세 테이블 (17행 × 2열)
t3 = doc.tables[3]   # 데이터 규모 요약 테이블 (5행 × 4열)
FONT = "맑은 고딕"
PT   = 11.0   # 139700 EMU ÷ 12700 = 11pt

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행02] 데이터 정의 — 클래스별 수치 수정 (혼용된 형식 통일 포함)
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[2].cells[1], [
    "선박 사고 시 발생하는 외상 이미지  (6종 클래스)",
    "- Abrasions   (찰과상) : 668장   ·  갑판·구조물에 쓸린 상처",
    "- Bruises     (타박상) : 972장   ·  충돌·낙하로 인한 타박",
    "- Burns       (화상)   : 504장   ·  엔진실 화재·증기 화상",
    "- Cut         (절창)   : 600장   ·  날카로운 도구에 의한 절상",
    "- Laceration  (열창)   : 732장   ·  거친 표면에 의한 찢김",
    "- Stab_wound  (자창)   : 736장   ·  날카로운 물체 관통상",
    "합계 : 4,212장",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행04] 수집 데이터 — 경로·수치 수정 (세그멘테이션 설명 유지)
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[4].cells[1], [
    "[CNN 분류 학습용]  통합 경로 :  01_data/all_labeled_data/",
    "Abrasions  (찰과상)",
    " - 데이터 수 : 668  /  위험도 : LOW",
    "Bruises  (타박상)",
    " - 데이터 수 : 972  /  위험도 : LOW",
    "Burns  (화상)",
    " - 데이터 수 : 504  /  위험도 : HIGH",
    "Cut  (절창)",
    " - 데이터 수 : 600  /  위험도 : MEDIUM",
    "Laceration  (열창)",
    " - 데이터 수 : 732  /  위험도 : HIGH",
    "Stab_wound  (자창)",
    " - 데이터 수 : 736  /  위험도 : CRITICAL",
    "데이터 수 합계  : 4,212",
    "",
    "[세그멘테이션 별도 모델용]  01_data/raw/wound_segmentation/",
    "- train_images 2,208장 + test_images 552장 = 이미지 2,760장  (마스크 동수 포함 시 총 5,520개)",
    "- CNN 분류기와 독립적인 별도 파이프라인  (상처 영역 검출용)",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행06] 전처리 — Train/Val 수치 및 Stab_wound 수치 수정
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[6].cells[1], [
    "1)  이미지 리사이즈 : 224 × 224 픽셀 통일",
    "2)  데이터 증강  (훈련셋 한정)",
    "- RandomHorizontalFlip,  RandomVerticalFlip",
    "- RandomRotation (±25°),  ColorJitter  (밝기·대비·채도·색조)",
    "- RandomAffine  (translate ±5%)",
    "3)  정규화 : ImageNet 기준",
    "mean = [0.485, 0.456, 0.406]   /   std = [0.229, 0.224, 0.225]",
    "4)  WeightedRandomSampler : 소수 클래스  (Stab_wound)  오버샘플링",
    "5)  Train / Val 분할 : 8 : 2  (Stratified Split,  random_state=42)",
    "→ 훈련 3,369장  /  검증 843장",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행09] 알고리즘 — 파라미터 수 및 모델 크기 수정
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[9].cells[1], [
    "MobileNetV3 Large (Expert Tuning)  (Transfer Learning,  ImageNet 사전학습 가중치 활용)",
    "- 파라미터 수  :  약 4.2 M  (4,209,718개, 실측)",
    "- 모델 크기    :  약 17 MB  →  선박 엣지 디바이스 배포 최적화",
    "- 추론 속도    :  < 200 ms  (CPU 기준)",
    "- 분류기 교체  :  Classifier 마지막 레이어 → Linear  (in_features → 6)",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행12] 학습 모델 — mobilenet_v3_small → mobilenet_v3_large 오기 수정
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[12].cells[1], [
    "MobileNetV3 Large (Expert Tuning)  (torchvision.models.mobilenet_v3_large)",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행15] 검증 방안 — 방법론 설명에서 수치(99.61%) 제거
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[15].cells[1], [
    "- Accuracy          :  전체 외상 분류 정확도",
    "- Precision (정밀도) :  해당 클래스로 분류한 것 중 실제 해당 클래스인 비율",
    "- Recall  (재현율)   :  실제 해당 클래스 중 모델이 올바르게 예측한 비율",
    "- F1 Score          :  Precision 과 Recall 의 조화평균",
    "- Confusion Matrix  :  클래스별 분류 오류 패턴 시각화",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T2 행16] 모델 평가 결과 — 실측값으로 전면 업데이트 (검증셋 843장 기준)
# ─────────────────────────────────────────────────────────────────────────────
rewrite_cell(t2.rows[16].cells[1], [
    "- Accuracy          :  99.76%",
    "- Precision (정밀도) :  99.76%",
    "- Recall  (재현율)   :  99.76%",
    "- F1 Score          :  99.76%",
    "",
    "[클래스별 상세 결과]  (검증셋 843장 기준, 전체 4,212장의 20%)",
    "- Abrasions  (찰과상)  :  Precision 1.00  /  Recall 0.99  /  F1 1.00",
    "- Bruises    (타박상)  :  Precision 1.00  /  Recall 1.00  /  F1 1.00",
    "- Burns      (화상)    :  Precision 0.99  /  Recall 1.00  /  F1 0.99",
    "- Cut        (절창)    :  Precision 0.99  /  Recall 1.00  /  F1 1.00",
    "- Laceration (열창)    :  Precision 1.00  /  Recall 1.00  /  F1 1.00",
    "- Stab_wound (자창)    :  Precision 1.00  /  Recall 1.00  /  F1 1.00",
], FONT, PT)

# ─────────────────────────────────────────────────────────────────────────────
# [T3] 데이터 규모 요약 테이블 — DL 수치 및 합계 수정
# ─────────────────────────────────────────────────────────────────────────────
FONT3 = "Malgun Gothic"
PT3   = 10.0   # 127000 EMU = 10pt

# 행2 열2: 2,326개 → 4,212개
rewrite_cell(t3.rows[2].cells[2], ["4,212개"], FONT3, PT3)

# 행4 열2: 9,350개 → 11,236개  (1,504 + 4,212 + 5,520)
rewrite_cell(t3.rows[4].cells[2], ["11,236개"], FONT3, PT3, bold=True)

# ─────────────────────────────────────────────────────────────────────────────
# 저장
# ─────────────────────────────────────────────────────────────────────────────
doc.save(str(SRC))
print(f"[완료] 저장: {SRC.name}")
print("[수정 항목 요약]")
print("  T2 행02 - 클래스별 이미지 수 및 형식 통일 (→ 4,212장)")
print("  T2 행04 - 경로(all_labeled_data), 수치 수정 (→ 4,212)")
print("  T2 행06 - Train 3,369장 / Val 843장 반영")
print("  T2 행09 - 파라미터 4.2M, 모델크기 17MB 반영")
print("  T2 행12 - mobilenet_v3_large 오기 수정")
print("  T2 행15 - 방법론 설명에서 수치 제거")
print("  T2 행16 - 실측 결과 (Accuracy·Precision·Recall·F1 = 99.76%) 기입")
print("  T3 행02 - DL 이미지 수 4,212개 반영")
print("  T3 행04 - 합계 11,236개 반영")
