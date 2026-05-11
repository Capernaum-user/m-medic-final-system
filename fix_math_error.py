# -*- coding: utf-8 -*-
from docx import Document

def fix_segmentation_math():
    file_path = r"D:\GeminiUniverse\vscode-workspace\wip-maritime-medic\M_MEDIC_v2\db_design\analysis\3-3._빅데이터분석정의서_MDTS 260422.docx"
    doc = Document(file_path)

    # 타겟 텍스트: "이미지 4,212장  (마스크 동수 포함)"
    # 정정 텍스트: "이미지 2,760장 (마스크 2,760장 포함 시 총 5,520개)"
    
    found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "2,208" in cell.text and "552" in cell.text:
                    for p in cell.paragraphs:
                        if "4,212" in p.text:
                            # 산식 오류 정정
                            p.text = p.text.replace("= 이미지 4,212장  (마스크 동수 포함)", "= 이미지 2,760장 (마스크 동수 포함 시 총 5,520개)")
                            found = True

    if not found:
        # 문단에서도 검색
        for p in doc.paragraphs:
            if "2,208" in p.text and "4,212" in p.text:
                p.text = p.text.replace("4,212", "2,760")
                if "마스크 동수 포함" in p.text:
                    p.text = p.text.replace("(마스크 동수 포함)", "(마스크 동수 포함 시 총 5,520개)")

    doc.save(file_path)
    print("SUCCESS: Segmentation math fixed. 2,208 + 552 = 2,760 (Total 5,520 with masks).")

if __name__ == '__main__':
    fix_segmentation_math()
